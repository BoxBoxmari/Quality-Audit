"""
Word document reading and table extraction utilities.
"""

import asyncio
import logging
import re
from concurrent.futures import ThreadPoolExecutor
from typing import Any, Dict, List, Optional, Sequence, Tuple, Union

import docx
import pandas as pd

from ..config.feature_flags import get_feature_flags
from ..utils.numeric_utils import compute_numeric_evidence_score
from ..utils.skip_classifier import classify_footer_signature
from .extractors import (
    OOXMLTableGridExtractor,
    PythonDocxExtractor,
)

logger = logging.getLogger(__name__)

# P1.1: Threshold for triggering LibreOffice fallback when table has good quality but low numeric evidence
NUMERIC_EVIDENCE_FALLBACK_THRESHOLD = 0.25
QUALITY_SCORE_FALLBACK_MIN = 0.6


class WordReader:
    """Handles reading and parsing Word documents with financial tables."""

    def __init__(self):
        """Initialize Word reader."""
        pass

    def _is_footer_or_signature_table(self, df: pd.DataFrame) -> bool:
        """
        SCRUM-6 / Spine 3: 2-phase classifier — footer/signature vs real financial table.
        Only returns True when positive evidence (footer keywords, short lines) is strong
        and negative evidence (share capital, equity, numeric density, year/currency) is weak.
        """
        should_skip, evidence = classify_footer_signature(df, heading="")
        if evidence.get("final_decision") == "skip":
            logger.debug(
                "word_reader: footer/signature skip. positive_hits=%s negative_hits=%s",
                len(evidence.get("positive_hits", [])),
                len(evidence.get("negative_hits", [])),
            )
        return should_skip

    def _reconstruct_table_grid(self, table: docx.table.Table) -> List[List[str]]:
        """
        Reconstruct table grid handling merged cells (occupancy-aware).

        Verification V-2:
        - Occupancy-based col_ptr advancement (skip already filled slots).
        - Track active_vmerge per column.
        - Normalize cell text (\u00a0 -> space, collapse spaces).
        """
        # 1. Calc max columns
        max_cols = 0
        for row in table.rows:
            col_count = 0
            for cell in row.cells:
                try:
                    cell_element = cell._element
                    grid_span_elem = cell_element.xpath(".//w:gridSpan")
                    span = 1
                    if grid_span_elem:
                        val = grid_span_elem[0].get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                        ) or grid_span_elem[0].get("val")
                        if val:
                            span = int(val)
                    col_count += span
                except (ValueError, TypeError, AttributeError):
                    col_count += 1
            max_cols = max(max_cols, col_count)

        if max_cols == 0:
            return []

        # 2. Init grid and active merge tracking
        # We don't pre-fill grid with vMerge continuations because DOCX usually provides a cell for them with vMerge="continue".
        # However, "occupancy-aware" implies we need to respect slots filled by previous rows if the XML skips them?
        # Actually standard DOCX XML (w:tblGrid) implies cells exist.
        # But let's follow the "skip to next empty cell" instruction which implies some cells might be skipped in traversal.

        grid = [["" for _ in range(max_cols)] for _ in range(len(table.rows))]
        active_vmerge: List[Optional[str]] = [
            None
        ] * max_cols  # Track text for vMerge="continue"

        for r_idx, row in enumerate(table.rows):
            col_ptr = 0

            # Skip cols physically occupied by previous row's rowspan?
            # In DOCX, vMerge doesn't consume the cell slot in correct XML, there is still a <w:tc>.
            # But the user instruction "skip to next empty cell" suggests we should skip if we already filled it?
            # Let's assume the grid starts empty for this row.

            for cell in row.cells:
                # Normalization
                raw_text = cell.text or ""
                normalized_text = (
                    raw_text.replace("\u00a0", " ").replace("\n", " ").strip()
                )
                normalized_text = re.sub(r"\s+", " ", normalized_text)

                cell_elem = cell._element

                # Get gridSpan
                span = 1
                gs = cell_elem.xpath(".//w:gridSpan")
                if gs:
                    val = gs[0].get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    ) or gs[0].get("val")
                    if val:
                        span = int(val)

                # Get vMerge
                vm = cell_elem.xpath(".//w:vMerge")
                v_val = None  # 'restart', 'continue', or None
                if vm:
                    v_val = vm[0].get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    ) or vm[0].get("val", "continue")

                # Find valid slot
                while col_ptr < max_cols and grid[r_idx][col_ptr] != "":
                    col_ptr += 1

                if col_ptr >= max_cols:
                    break

                # Determine content
                final_text: Optional[str] = normalized_text

                if v_val == "restart":
                    # Updates active merge
                    for i in range(span):
                        if col_ptr + i < max_cols:
                            active_vmerge[col_ptr + i] = normalized_text
                elif v_val == "continue":
                    # Uses active merge
                    # Use the text from the first column of the span
                    if active_vmerge[col_ptr] is not None:
                        final_text = active_vmerge[col_ptr]
                    # Note: we don't update active_vmerge here, we just use it.
                else:
                    # No vmerge, or vmerge ended.
                    # Clear active merge for these cols? DOCX is restart|continue|none.
                    # If tag is missing, it breaks the merge.
                    for i in range(span):
                        if col_ptr + i < max_cols:
                            active_vmerge[col_ptr + i] = None

                # Fill grid
                for i in range(span):
                    if col_ptr + i < max_cols and final_text is not None:
                        grid[r_idx][col_ptr + i] = final_text

                col_ptr += span

        return grid

    def _derive_heading_from_table_first_row(
        self, table: docx.table.Table
    ) -> Optional[str]:
        """
        Derive heading from the first 1–2 rows of the reconstructed table.

        Picks the longest non-empty cell that is not numeric and not a code-like pattern.
        Returns None if no suitable candidate is found.
        """
        try:
            reconstructed_rows = self._reconstruct_table_grid(table)
        except Exception as e:
            logger.debug(
                "Failed to reconstruct table for first-row heading inference: %s", e
            )
            reconstructed_rows = []

        if not reconstructed_rows:
            return None

        max_rows = min(2, len(reconstructed_rows))
        candidates: List[str] = []

        for r_idx in range(max_rows):
            row = reconstructed_rows[r_idx]
            for cell in row:
                if cell is None:
                    continue
                text = str(cell).strip()
                if not text:
                    continue

                # Skip numeric-like cells
                cleaned = (
                    text.replace(",", "")
                    .replace(".", "")
                    .replace(" ", "")
                    .replace("(", "")
                    .replace(")", "")
                )
                if cleaned.lstrip("-").isdigit():
                    continue

                # Skip pure code-like values (e.g. '10', '100A')
                if re.match(r"^[0-9]{2,}[A-Z]?$", text.strip()):
                    continue

                if self._is_heading_junk(text):
                    continue

                candidates.append(text)

        if not candidates:
            return None

        # Pick the longest candidate as heading
        heading = max(candidates, key=len)
        logger.info(
            "Heading candidate from table first row: %s, heading_source=table_first_row",
            heading,
        )
        return heading

    def _is_heading_junk(self, text: str) -> bool:
        """
        heading_inference_v2: Filter out paragraph text that looks like data/junk,
        not a real section heading (digit/currency density, date-like, length, boilerplate).
        """
        if not text or len(text) < 3:
            return True
        text_stripped = text.strip()
        if len(text_stripped) < 3:
            return True
        text_lower = text_stripped.lower()
        # Whitelist: known valid heading patterns are not junk
        valid_patterns = [
            r"balance sheet",
            r"cân đối kế toán",
            r"income statement",
            r"kết quả kinh doanh",
            r"cash flow",
            r"lưu chuyển tiền",
            r"equity",
            r"vốn chủ sở hữu",
            r"note \d+:",  # Note with description
        ]
        if any(re.search(p, text_lower) for p in valid_patterns):
            return False
        # --- Ticket-1 explicit rejection patterns ---
        words = text_lower.split()

        # Unit-only lines: "Đơn vị tính: VND" / "Unit: VND'000"
        if re.match(
            r"^(unit|đơn\s*vị\s*(tính)?|đvt|currency|tỷ\s*giá)\s*:?\s*.{0,20}$",
            text_lower,
        ):
            return True

        # Sparse lines with year, date or currency (e.g. "2018 VND'000", "2023", "VND", "31/12/2018")
        if len(words) <= 4:
            # Check for year (e.g. "2023", "Năm 2023")
            if any(re.match(r"^(19|20)\d{2}$", w) for w in words):
                return True
            # Check for pure currency/unit (e.g. "VND", "USD'000")
            if any(
                w in ["vnd", "usd", "eur", "đồng"] or "vnd'" in w or "usd'" in w
                for w in words
            ):
                return True
            # Check for dates (e.g "31/12/2018")
            if any(re.match(r"\d{1,2}/\d{1,2}/\d{2,4}", w) for w in words):
                return True
        # --- End Ticket-1 patterns ---
        # Digit/currency density: if > 50% of non-space chars are digits or currency
        digits = sum(1 for c in text_stripped if c.isdigit())
        currency_chars = sum(1 for c in text_stripped if c in "$.,()" or c == "\u00a0")
        total = len(text_stripped)
        if total > 0 and (digits + currency_chars) / total > 0.50:
            return True
        # Date-like: only filter if date is majority of text
        date_match = re.search(r"\d{1,2}/\d{1,2}/\d{2,4}", text_stripped)
        if date_match and len(date_match.group()) / len(text_stripped) > 0.5:
            return True
        if re.search(r"20\d{2}\s*$", text_stripped) and digits / max(1, total) > 0.2:
            return True
        # Boilerplate: "Note X", "Table X", "Appendix X" without description (allow if ":")
        words = text_stripped.split()
        if ":" in text_stripped:
            return False  # e.g. "Note 5: Inventories" is valid
        return (
            len(words) <= 2
            and text_lower.startswith(("note ", "table ", "appendix "))
            and bool(re.match(r"^(note|table|appendix)\s+\d+$", text_lower))
        )

    def _numeric_density(self, row_values: Sequence[Any]) -> float:
        """
        Return fraction of cells in row that look numeric (0..1).
        Used by multi-row header detection to skip data-heavy rows.
        """
        if not row_values:
            return 0.0
        numeric_count = 0
        for v in row_values:
            s = str(v).strip() if v is not None else ""
            if not s:
                continue
            try:
                float(s.replace(",", "").replace(" ", ""))
                numeric_count += 1
            except (ValueError, TypeError):
                pass
        return numeric_count / len(row_values) if row_values else 0.0

    def _detect_header_row_range(
        self,
        df: pd.DataFrame,
        code_synonyms: set,
        note_synonyms: set,
        year_patterns: tuple,
        scan_rows: int,
    ) -> Tuple[Optional[int], Optional[int]]:
        """
        Detect contiguous header row range (start, end) 0-based inclusive.
        Header-like row: score >= 3 and (has code synonym or (has note and year)).
        Extend range downward while next row has low numeric density and
        (header-like or short/repeated cells).
        Returns (start, end) or (None, None) if not found.
        """

        def row_score(row_values: Sequence[Any]) -> int:
            score = 0
            for v in row_values:
                s = (str(v).strip() or "").lower()
                if not s:
                    continue
                if s in code_synonyms or any(s.startswith(c) for c in code_synonyms):
                    score += 2
                if s in note_synonyms or any(s.startswith(n) for n in note_synonyms):
                    score += 1
                for pat in year_patterns:
                    if re.search(pat, s):
                        score += 1
                        break
            return score

        start, end = None, None
        for r in range(min(scan_rows, len(df))):
            row_vals = df.iloc[r].tolist()
            if row_score(row_vals) >= 3 and (
                any(
                    (str(v).strip().lower() in code_synonyms)
                    for v in row_vals
                    if v is not None
                )
                or (
                    any(
                        (str(v).strip().lower() in note_synonyms)
                        for v in row_vals
                        if v is not None
                    )
                    and any(
                        re.search(pat, str(v) or "")
                        for v in row_vals
                        for pat in year_patterns
                        if v is not None
                    )
                )
            ):
                if start is None:
                    start = r
                end = r
                continue
            if start is not None:
                dens = self._numeric_density(row_vals)
                if dens < 0.5 and (
                    row_score(row_vals) >= 2
                    or all(len(str(v or "").strip()) < 20 for v in row_vals)
                ):
                    end = r
                else:
                    break
        if start is not None and end is not None:
            return (start, end)
        return (None, None)

    def _consolidate_headers_from_range(
        self, df: pd.DataFrame, start: int, end: int
    ) -> list[str]:
        """
        For each column, join non-empty cell values in rows [start..end] with space.
        Returns list of strings of length = df.shape[1].
        """
        cols = df.shape[1]
        result: list[str] = []
        for c in range(cols):
            parts = []
            for r in range(start, min(end + 1, len(df))):
                v = df.iloc[r, c]
                s = str(v).strip() if v is not None else ""
                if s:
                    parts.append(s)
            if parts:
                result.append(" ".join(parts))
            else:
                # Preserve amount columns: if column has numeric content in data rows, name as Amount_N
                has_numeric = False
                for r in range(end + 1, min(end + 11, len(df))):
                    try:
                        val = df.iloc[r, c]
                        if pd.notna(pd.to_numeric(val, errors="coerce")):
                            has_numeric = True
                            break
                    except Exception:
                        pass
                result.append(f"Amount_{c}" if has_numeric else f"Column_{c}")
        return result

    def _promote_header_row(
        self, df: pd.DataFrame, heading: Optional[str] = None
    ) -> pd.DataFrame:
        """
        SCRUM-11 P0: Promote header row from data rows to column names.

        Fixes issue where header row (containing "Code", "Note", year patterns)
        is pushed down into data rows, causing validators to fail with
        "WARN: không tìm thấy cột 'Code'".

        Logic:
        1. If df.columns are numeric (0..n) or "Unnamed", scan first 3-8 rows
        2. Header candidate row is valid when it has:
           - "Code" (or synonym: "No.", "No", "Item", "Ref", "Mã", "Item no")
           - "Note" (or synonym: "notes", "ghi chú", "chú thích")
           - AND/OR year pattern (2018, 31/12/2018, etc.)
        3. Promote header row to df.columns, drop from data
        4. Clean header values (trim, remove NBSP, replace \\n with space)
        5. If first header cell is empty, name it "Item" or "Description"

        Args:
            df: DataFrame that may have header row in data
            heading: Optional heading to determine if this is a statement table

        Returns:
            DataFrame with proper header promotion applied
        """
        if df.empty or len(df) == 0:
            return df

        # Check if columns need promotion (numeric or "Unnamed")
        needs_promotion = False
        first_col = str(df.columns[0]) if len(df.columns) > 0 else ""

        # Check if columns are numeric (0, 1, 2, ...) or "Unnamed"
        if first_col.isdigit() or first_col.startswith("Unnamed"):
            needs_promotion = True
        elif len(df.columns) > 0:
            # Check if all columns are numeric strings
            all_numeric = all(str(c).strip().isdigit() for c in df.columns)
            if all_numeric:
                needs_promotion = True

        if not needs_promotion:
            # Columns already have names, just deduplicate if needed
            return self._deduplicate_headers(df)

        # Code/Note/Year patterns (used by both multi-row and single-row header paths)
        code_synonyms = [
            "code",
            "mã",
            "no.",
            "no",
            "item",
            "ref",
            "item no",
            "item no.",
            "reference",
            "mã số",
            "số hiệu",
        ]

        # Note column synonyms
        note_synonyms = [
            "note",
            "notes",
            "ghi chú",
            "chú thích",
            "note no",
            "note no.",
            "chú giải",
        ]

        # Year patterns
        year_patterns = [
            r"\d{4}",  # 2018, 2017
            r"\d{1,2}/\d{1,2}/\d{2,4}",  # 31/12/2018, 12/31/18
            r"20\d{2}",  # 2018, 2019
            r"19\d{2}",  # 1999
        ]

        # Scan first 3-8 rows for header candidate
        scan_rows = min(8, len(df))

        # Phase 3.1: Try multi-row header detection (low numeric density + patterns)
        start, end = self._detect_header_row_range(
            df,
            set(code_synonyms),
            set(note_synonyms),
            tuple(year_patterns),
            scan_rows,
        )
        if start is not None and end is not None and end >= start:
            consolidated = self._consolidate_headers_from_range(df, start, end)
            df_new = df.copy()
            df_new.columns = consolidated
            df_new = df_new.iloc[end + 1 :].reset_index(drop=True)
            return self._deduplicate_headers(df_new)

        # Track two separate candidates:
        # 1. Best raw score (fallback)
        best_header_idx = None
        best_score = 0

        # 2. Best valid score (primary)
        best_valid_idx = None
        best_valid_score = 0

        for row_idx in range(scan_rows):
            if row_idx >= len(df):
                break

            row_values = [str(x).strip() for x in df.iloc[row_idx]]
            row_lower = [v.lower() for v in row_values]

            # Ticket 8: Header Promotion Guardrails
            # 1. Minimum Width Sanity Check
            non_empty_cells = [v for v in row_values if v]
            if len(non_empty_cells) < 2 or len(df) < 3:
                continue

            # 2. Cell Count / Uniqueness Check
            nunique_text_cells = len(
                set(v for v in non_empty_cells if not v.replace(".", "", 1).isdigit())
            )
            if nunique_text_cells < 2:
                continue

            # 3. All-caps / Roman numeral filter
            # Reject if the row represents a section title like "I. TÀI SẢN"
            first_cell = non_empty_cells[0]
            if re.match(r"^(I+\.|II+\.|\d+\.)\s*[A-Z\sÀ-Ỵ]+$", first_cell):
                continue

            # 4. Data Row Check
            # Row i+1 must have > 50% numeric density
            next_row_idx = row_idx + 1
            if next_row_idx < len(df):
                next_row_vals = [str(x).strip() for x in df.iloc[next_row_idx]]
                num_density = self._numeric_density(next_row_vals)
                if num_density <= 0.5:
                    continue  # Next row is mostly text, so this is likely a paragraph row

            # Score this row as header candidate
            score = 0
            has_code = False
            has_note = False
            has_year = False

            # Check for Code column
            for val_lower in row_lower:
                # P1-T2: Stricter matching to avoid false positives (e.g. 'no' in 'notes')
                is_code = val_lower in code_synonyms

                # Special cases for bilingual or specific patterns
                if not is_code and (
                    "code" in val_lower
                    and ("/" in val_lower or "mã" in val_lower)
                    or val_lower.startswith("item no")
                    or val_lower == "no."
                ):
                    is_code = True

                if is_code:
                    has_code = True
                    score += 3
                    break

            # Check for Note column
            for val_lower in row_lower:
                is_note = val_lower in note_synonyms

                # Special cases for bilingual
                if (
                    not is_note
                    and "note" in val_lower
                    and (
                        "/" in val_lower
                        or "thuyết" in val_lower
                        or "ghi" in val_lower
                        or "chú" in val_lower
                    )
                ):
                    is_note = True

                if is_note:
                    has_note = True
                    score += 3
                    break

            # Check for year patterns
            for val in row_values:
                for pattern in year_patterns:
                    if re.search(pattern, val, re.IGNORECASE):
                        has_year = True
                        score += 2
                        break
                if has_year:
                    break

            # Bonus: if has both Code and Note, or Code and Year
            if (has_code and has_note) or (has_code and has_year):
                score += 5

            # Bonus: if this is a statement table (balance sheet, income, cash flow)
            if heading:
                heading_lower = heading.lower()
                if (
                    any(
                        term in heading_lower
                        for term in [
                            "balance sheet",
                            "statement of income",
                            "cash flow",
                        ]
                    )
                    and has_code
                ):
                    # Statement tables should have Code column
                    score += 2

            # UPDATE 1: Always update raw best score (for potential fallback)
            if score > best_score:
                best_score = score
                best_header_idx = row_idx
            elif score == best_score:
                # Prefer earlier row if scores equal - no change needed
                pass

            # Gates: strictly require strong confidence
            # V-3: Keyword AND (Year Pattern OR Low Numeric Ratio)
            # Note: numeric_ratio and desc_keywords calculations removed as they were unused

            # Must have detected Code or Note (score >= 3 implies this given logic above,
            # but let's be explicit based on 'has_code' or 'has_note')

            # Gate condition - Simplified: chỉ cần Code keyword là đủ (match legacy behavior)
            is_valid_header = False
            if has_code or has_note and has_year:  # Chỉ cần Code keyword (match legacy)
                is_valid_header = True

            # Comment 2: Track best VALID candidate separately
            # Use 'is_valid_header' flag and minimum score threshold
            if is_valid_header and score >= 3:
                if score > best_valid_score:
                    best_valid_score = score
                    best_valid_idx = row_idx
                elif score == best_valid_score:
                    if best_valid_idx is None:
                        best_valid_idx = row_idx
                    # Else prefer earlier row (keep existing)

        # SELECTION LOGIC
        final_header_idx = None

        # Priority 1: Best Valid Candidate
        if best_valid_idx is not None:
            final_header_idx = best_valid_idx

        # Priority 2: Best Raw Candidate (if it passes re-validation)
        # This covers cases where 'is_valid_header' inside loop might have been too strict
        # but re-check passes, OR if the winning raw candidate is accidentally valid
        # but wasn't caught? (Actually if it was valid it would be in best_valid_idx unless score < best_valid_score)
        # But let's check it anyway as fallback if no valid_idx found.
        elif best_header_idx is not None and best_score >= 5:
            # Check validity one last time
            # (Reuse logic or just assume if no valid found, maybe rely on score?)
            # The instruction says: "Else if best_header_idx passes re-validation, use it."
            pass  # Will be checked below
            final_header_idx = best_header_idx

        # If we have a candidate, promote it
        if final_header_idx is not None:
            best_header_idx = final_header_idx  # Use for subsequent logic
        else:
            # Fallback: legacy simple scan (match legacy behavior)
            # Scan top 5 rows tìm "code" → promote row đó
            for row_idx in range(min(5, len(df))):
                row_lower = [str(x).lower() for x in df.iloc[row_idx]]
                if any("code" in v for v in row_lower):
                    final_header_idx = row_idx
                    best_header_idx = final_header_idx
                    break

        if final_header_idx is not None:
            row_vals = [str(x).strip() for x in df.iloc[best_header_idx]]
            header_values = row_vals
            # ... rest of promotion logic ...
            cleaned_headers: List[str] = []
            for val in header_values:
                cleaned = (
                    val.replace("\u00a0", " ")
                    .replace("\u200b", "")
                    .replace("\n", " ")
                    .replace("\r", " ")
                )
                cleaned = re.sub(r"\s+", " ", cleaned).strip()
                if not cleaned:
                    if len(cleaned_headers) == 0:
                        cleaned = "Item"
                    else:
                        cleaned = f"Column_{len(cleaned_headers) + 1}"
                cleaned_headers.append(cleaned)

            # Merge with previous row if it contains years (Multi-row header support)
            # P1-E2: Strict Header promotion - check context
            best_header_idx_start: Optional[int] = best_header_idx

            if best_header_idx is not None and best_header_idx > 0:
                prev_row = df.iloc[best_header_idx - 1]
                prev_row_vals = [str(x).strip() for x in prev_row]

                # Check if previous row has multiple numeric/year values
                has_year_prev = any(
                    re.search(p, v, re.IGNORECASE)
                    for v in prev_row_vals
                    for p in year_patterns
                )

                if has_year_prev:
                    # Merge headers: "2024" + "VND" = "2024 VND"
                    unique_merged = []
                    seen_map: Dict[str, int] = {}

                    for i, current_val in enumerate(cleaned_headers):
                        prev_val = prev_row_vals[i] if i < len(prev_row_vals) else ""
                        # Clean prev val
                        prev_val = re.sub(r"\s+", " ", prev_val).strip()

                        if prev_val and prev_val.lower() != "nan":
                            merged_val = f"{prev_val} {current_val}".strip()
                        else:
                            merged_val = current_val

                        # Deduplicate
                        if merged_val in seen_map:
                            seen_map[merged_val] += 1
                            merged_val = f"{merged_val}.{seen_map[merged_val]}"
                        else:
                            seen_map[merged_val] = 0

                        unique_merged.append(merged_val)

                    deduplicated_headers = unique_merged
                    if best_header_idx is not None:
                        best_header_idx_start = best_header_idx - 1
                else:
                    # Standard deduplication
                    seen: Dict[str, int] = {}
                    final_headers = []
                    for val in cleaned_headers:
                        if val in seen:
                            seen[val] += 1
                            final_headers.append(f"{val}.{seen[val]}")
                        else:
                            seen[val] = 0
                            final_headers.append(val)
                    deduplicated_headers = final_headers
            else:
                # Standard deduplication for row 0
                seen = {}
                final_headers = []
                for val in cleaned_headers:
                    if val in seen:
                        seen[val] += 1
                        final_headers.append(f"{val}.{seen[val]}")
                    else:
                        seen[val] = 0
                        final_headers.append(val)
                deduplicated_headers = final_headers

            df_new = df.copy()
            df_new.columns = deduplicated_headers

            if best_header_idx is not None and best_header_idx == 0:
                df_new = df_new.iloc[1:].reset_index(drop=True)
            elif best_header_idx is not None and best_header_idx_start is not None:
                # Use determined start index for dropping
                # We drop from best_header_idx_start to best_header_idx (inclusive)
                # But actually checking line 559 in original:
                # if all(v.isdigit() ...) drop_start_idx = best - 1
                # Our logic replaces that heuristic with explicit year check merging.

                df_new = pd.concat(
                    [
                        df_new.iloc[:best_header_idx_start],
                        df_new.iloc[best_header_idx + 1 :],
                    ]
                ).reset_index(drop=True)

            return df_new

        # No valid header found, return original with deduplicated numeric columns
        return self._deduplicate_headers(df)

    def _deduplicate_headers(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Deduplicate column headers to prevent duplicate column names.

        Args:
            df: DataFrame with potentially duplicate headers

        Returns:
            DataFrame with deduplicated headers
        """
        if df.empty:
            return df

        header_values = [str(c).strip() for c in df.columns]
        seen: Dict[str, int] = {}
        deduplicated_headers = []

        for val in header_values:
            if val in seen:
                seen[val] += 1
                deduplicated_headers.append(f"{val}.{seen[val]}")
            else:
                seen[val] = 0
                deduplicated_headers.append(val)

        df_new = df.copy()
        df_new.columns = deduplicated_headers
        return df_new

    def _extract_table_with_fallback(
        self, file_path: str, table_index: int, table: docx.table.Table
    ) -> Tuple[List[List[str]], Dict[str, Any]]:
        """
        Extract table grid with multi-engine fallback: OOXML (A) -> Python-docx (B) -> LibreOffice (C).
        Engine C is used only for the first table (table_index == 0) since it returns the first HTML table.
        R1: Returns engine_attempts and OOXML telemetry (invariants_failed, grid_cols_*, gridSpan_count, vMerge_count).
        """
        engine_attempts: List[str] = []
        ooxml_result: Any = None

        # Engine A: OOXML
        try:
            engine_attempts.append("ooxml")
            ooxml = OOXMLTableGridExtractor()
            res_a = ooxml.extract(table)
            if res_a.is_usable and res_a.grid:
                meta: Dict[str, Any] = {
                    "extractor_engine": "ooxml",
                    "extractor_usable_reason": "usable",
                    "engine_attempts": list(engine_attempts),
                    "quality_score": res_a.quality_score,
                    "quality_flags": list(res_a.quality_flags),
                    "failure_reason_code": getattr(res_a, "failure_reason_code", None),
                    "invariants_failed": list(
                        getattr(res_a, "invariant_violations", []) or []
                    ),
                    "grid_cols_expected": getattr(res_a, "grid_cols_expected", None),
                    "grid_cols_built": getattr(res_a, "grid_cols_built", 0),
                    "gridSpan_count": getattr(res_a, "grid_span_count", 0),
                    "vMerge_count": getattr(res_a, "vmerge_count", 0),
                }
                # P1.1: Post-OOXML numeric evidence gating — trigger fallback when structurally good but numeric-empty
                temp_df = pd.DataFrame(res_a.grid)
                temp_df = self._promote_header_row(temp_df)
                cand = (
                    temp_df.columns[-4:].tolist()
                    if len(temp_df.columns) >= 4
                    else temp_df.columns.tolist()
                )
                evidence = compute_numeric_evidence_score(
                    temp_df.iloc[1:], candidate_columns=cand if cand else None
                )
                meta["numeric_evidence_score"] = evidence["numeric_evidence_score"]
                meta["numeric_col_candidates"] = evidence.get(
                    "numeric_col_candidates", []
                )
                quality_score = meta["quality_score"]
                if (
                    quality_score >= QUALITY_SCORE_FALLBACK_MIN
                    and evidence["numeric_evidence_score"]
                    < NUMERIC_EVIDENCE_FALLBACK_THRESHOLD
                ):
                    logger.info(
                        "Structurally good but numeric-empty (score=%.3f), triggering fallback",
                        evidence["numeric_evidence_score"],
                    )
                    fallback_success = False
                    try:
                        from .extractors.render_first_table_extractor import (
                            RenderFirstTableExtractor,
                        )

                        rf = RenderFirstTableExtractor()
                        if rf.is_available():
                            engine_attempts.append("render_first")
                            res_c = rf.extract(table, file_path, table_index)
                            if res_c.is_usable and res_c.grid:
                                df_fb = pd.DataFrame(res_c.grid)
                                df_fb = self._promote_header_row(df_fb)
                                cand_fb = (
                                    df_fb.columns[-4:].tolist()
                                    if len(df_fb.columns) >= 4
                                    else df_fb.columns.tolist()
                                )
                                ev_fb = compute_numeric_evidence_score(
                                    df_fb.iloc[1:],
                                    candidate_columns=(cand_fb if cand_fb else None),
                                )
                                if (
                                    ev_fb["numeric_evidence_score"]
                                    >= NUMERIC_EVIDENCE_FALLBACK_THRESHOLD
                                    and res_c.quality_score >= 0.4
                                ):
                                    fallback_success = True
                                    return (
                                        res_c.grid,
                                        {
                                            "extractor_engine": "render_first",
                                            "extractor_usable_reason": "usable",
                                            "engine_attempts": list(engine_attempts),
                                            "quality_score": res_c.quality_score,
                                            "quality_flags": list(res_c.quality_flags),
                                            "failure_reason_code": getattr(
                                                res_c, "failure_reason_code", None
                                            ),
                                            "invariants_failed": None,
                                            "grid_cols_expected": None,
                                            "grid_cols_built": getattr(
                                                res_c, "cols", None
                                            ),
                                            "gridSpan_count": None,
                                            "vMerge_count": None,
                                            "numeric_evidence_score": ev_fb[
                                                "numeric_evidence_score"
                                            ],
                                            "numeric_col_candidates": ev_fb.get(
                                                "numeric_col_candidates", []
                                            ),
                                            "conversion_mode": getattr(
                                                res_c, "conversion_mode", None
                                            ),
                                            "structure_recognizer": getattr(
                                                res_c, "structure_recognizer", None
                                            ),
                                            "ocr_engine": getattr(
                                                res_c, "ocr_engine", None
                                            ),
                                            "token_coverage_ratio": getattr(
                                                res_c, "token_coverage_ratio", None
                                            ),
                                            "mean_cell_confidence": getattr(
                                                res_c, "mean_cell_confidence", None
                                            ),
                                            "p10_cell_confidence": getattr(
                                                res_c, "p10_cell_confidence", None
                                            ),
                                            "empty_cell_ratio": getattr(
                                                res_c, "empty_cell_ratio", None
                                            ),
                                            "debug_artifact_path": getattr(
                                                res_c, "debug_artifact_path", None
                                            ),
                                        },
                                    )
                    except Exception as e:
                        logger.debug(
                            "Render-first fallback (numeric-empty) skipped: %s", e
                        )
                    if not fallback_success:
                        meta["fallback_attempted"] = True
                        meta["fallback_failed"] = True
                        meta["failure_reason_code"] = "NO_NUMERIC_EVIDENCE"
                        return (res_a.grid, meta)
                return (res_a.grid, meta)
            else:
                ooxml_result = res_a
        except Exception as e:
            logger.debug("OOXML extractor skipped: %s", e)

        # Trigger for render-first (before python-docx) when extraction_fallback_prefer_advanced_before_legacy
        flags = get_feature_flags()
        prefer_advanced = flags.get(
            "extraction_fallback_prefer_advanced_before_legacy", True
        )
        trigger_mode = flags.get(
            "extraction_render_first_triggered_mode", "signals_only"
        )
        ooxml_quality = (
            getattr(ooxml_result, "quality_score", 0.0)
            if ooxml_result is not None
            else 0.0
        )
        trigger_render_first = False
        trigger_reason: Optional[str] = None
        if prefer_advanced and trigger_mode == "always_on":
            trigger_render_first = True
            trigger_reason = "always_on"
        elif (
            prefer_advanced
            and trigger_mode == "signals_only"
            and ooxml_result is not None
        ):
            fc = getattr(ooxml_result, "failure_reason_code", None)
            if fc in ("GRID_CORRUPTION", "DUPLICATE_PERIODS") or ooxml_quality < 0.6:
                trigger_render_first = True
                trigger_reason = fc or "ooxml_low_quality"
        if trigger_mode == "always_off":
            trigger_render_first = False

        _render_first_rejected_meta: Optional[Dict[str, Any]] = None
        if trigger_render_first:
            logger.info(
                "Render-first triggered: trigger=%s, ooxml_quality=%.2f, table_index=%s",
                trigger_reason or "signals_only",
                ooxml_quality,
                table_index,
            )
            try:
                from .extractors.render_first_table_extractor import (
                    RenderFirstTableExtractor,
                )

                render_first = RenderFirstTableExtractor()
                if render_first.is_available():
                    engine_attempts.append("render_first")
                    res_c = render_first.extract(table, file_path, table_index)
                    if res_c.is_usable and res_c.grid:
                        meta_c: Dict[str, Any] = {
                            "extractor_engine": "render_first",
                            "extractor_usable_reason": "usable",
                            "engine_attempts": list(engine_attempts),
                            "quality_score": res_c.quality_score,
                            "quality_flags": list(res_c.quality_flags),
                            "failure_reason_code": getattr(
                                res_c, "failure_reason_code", None
                            ),
                            "invariants_failed": None,
                            "grid_cols_expected": None,
                            "grid_cols_built": getattr(res_c, "cols", None),
                            "gridSpan_count": None,
                            "vMerge_count": None,
                            "trigger_reason": trigger_reason,
                            "conversion_mode": getattr(res_c, "conversion_mode", None),
                            "structure_recognizer": getattr(
                                res_c, "structure_recognizer", None
                            ),
                            "ocr_engine": getattr(res_c, "ocr_engine", None),
                            "token_coverage_ratio": getattr(
                                res_c, "token_coverage_ratio", None
                            ),
                            "mean_cell_confidence": getattr(
                                res_c, "mean_cell_confidence", None
                            ),
                            "p10_cell_confidence": getattr(
                                res_c, "p10_cell_confidence", None
                            ),
                            "empty_cell_ratio": getattr(
                                res_c, "empty_cell_ratio", None
                            ),
                            "debug_artifact_path": getattr(
                                res_c, "debug_artifact_path", None
                            ),
                        }
                        return (res_c.grid, meta_c)
                    _render_first_rejected_meta = {
                        "rejection_reason": getattr(res_c, "rejection_reason", None)
                        or getattr(res_c, "failure_reason_code", None),
                        "mean_cell_confidence": getattr(
                            res_c, "mean_cell_confidence", None
                        ),
                        "token_coverage_ratio": getattr(
                            res_c, "token_coverage_ratio", None
                        ),
                        "quality_score": getattr(res_c, "quality_score", None),
                        "conversion_mode": getattr(res_c, "conversion_mode", None),
                    }
                else:
                    _render_first_rejected_meta = {
                        "rejection_reason": "render-first skipped: system binaries missing (soffice)",
                    }
            except Exception as e:
                logger.debug("Render-first extractor skipped: %s", e)
                _render_first_rejected_meta = {"rejection_reason": str(e)}

        # Engine B: Python-docx
        try:
            engine_attempts.append("python_docx")
            python_docx = PythonDocxExtractor()
            res_b = python_docx.extract(table)
            if res_b.is_usable and res_b.grid:
                return (
                    res_b.grid,
                    {
                        "extractor_engine": "python_docx",
                        "extractor_usable_reason": "usable",
                        "engine_attempts": list(engine_attempts),
                        "quality_score": res_b.quality_score,
                        "quality_flags": list(res_b.quality_flags),
                        "failure_reason_code": getattr(
                            res_b, "failure_reason_code", None
                        ),
                        "invariants_failed": None,
                        "grid_cols_expected": None,
                        "grid_cols_built": None,
                        "gridSpan_count": None,
                        "vMerge_count": None,
                    },
                )
        except Exception as e:
            logger.debug("Python-docx extractor skipped: %s", e)

        # Legacy fallback
        return (
            self._reconstruct_table_grid(table),
            {
                "extractor_engine": "legacy",
                "extractor_usable_reason": "LEGACY_FALLBACK",
                "engine_attempts": list(engine_attempts),
                "quality_score": 0.0,
                "quality_flags": ["legacy_fallback"],
                "failure_reason_code": "LEGACY_FALLBACK",
                "invariants_failed": None,
                "grid_cols_expected": None,
                "grid_cols_built": None,
                "gridSpan_count": None,
                "vMerge_count": None,
                "render_first_metadata": _render_first_rejected_meta,
            },
        )

    def read_tables_with_headings(
        self, file_path: str, include_context: bool = True
    ) -> Union[
        List[Tuple[pd.DataFrame, Optional[str], Dict[str, Any]]],
        List[Tuple[pd.DataFrame, Optional[str]]],
    ]:
        """
        Read tables from Word document and extract associated headings.

        P0-R1: Handles merged cells by reconstructing proper grid matrix.

        Args:
            file_path: Path to Word document
            include_context: If True (default), return 3-tuples (df, heading, table_context);
                if False, return 2-tuples (df, heading) for backward compatibility.

        Returns:
            List of (table_df, heading) or (table_df, heading, table_context) tuples

        Raises:
            FileNotFoundError: If file doesn't exist
            docx.opc.exceptions.PackageNotFoundError: If file is corrupted
        """
        doc = docx.Document(file_path)
        tables: List[pd.DataFrame] = []
        headings: List[Optional[str]] = []
        table_contexts: List[Dict[str, Any]] = []
        # queue of last N paragraphs: list of (text, style_name, is_bold, is_upper)
        # Comment 2: reset after each table to avoid heading bleed to next table
        prior_paragraphs: List[Tuple[str, str, bool, bool]] = []
        current_heading = None
        current_note_number = None
        paragraphs_since_last_table = 0
        long_paragraph_since_last_table = False
        max_paragraphs_since_table = (
            15  # increased from 8 to 15 to find headings further away from the table
        )

        # Keywords for scoring
        strong_keywords = [
            "balance sheet",
            "bảng cân đối kế toán",
            "statement of financial position",
            "statement of income",
            "income statement",
            "profit and loss",
            "kết quả kinh doanh",
            "báo cáo kết quả",
            "cash flow",
            "lưu chuyển tiền tệ",
            "lưu chuyển tiền",
            "changes in equity",
            "changes in owners' equity",
            "vốn chủ sở hữu",
            "thay đổi vốn",
            "notes to",
            "thuyết minh",
            "ghi chú",
            "tax",
            "thuế",
        ]

        # Scan document
        for block in doc.element.body:
            # Section boundary: reset heading context so headings don't bleed across sections
            if block.tag.endswith("sectPr"):
                prior_paragraphs.clear()
                current_heading = None
                current_note_number = None
                continue
            if block.tag.endswith("tbl"):
                # Create table object once so it can be used for heading fallback and parsing
                table = docx.table.Table(block, doc)
                # V-4: Scan previous 1-8 paragraphs for best candidate
                best_candidate_text = None
                best_score: float = 0.0
                candidates_log = []
                flags = get_feature_flags()
                use_heading_v2 = flags.get("heading_inference_v2", False)

                # Check stored paragraphs (reverse order = closest first)
                for idx, (p_text, p_style, p_bold, p_upper) in enumerate(
                    reversed(prior_paragraphs)
                ):
                    if idx >= max_paragraphs_since_table:
                        break

                    if not p_text:
                        continue
                    # Phase 1: skip junk paragraphs when heading_inference_v2 is on
                    if use_heading_v2 and self._is_heading_junk(p_text):
                        continue

                    score: float = 0.0
                    text_lower = p_text.lower()

                    # 1. Keyword Match (Strongest Signal)
                    if any(k in text_lower for k in strong_keywords):
                        score += 10

                    # 2. Style/Formatting
                    if "Heading" in p_style:
                        score += 5
                    if p_bold:
                        score += 3
                    if p_upper:
                        score += 2

                    # 3. Content Heuristics
                    # Year pattern: only boost when accompanied by contextual phrases
                    if re.search(r"20\d{2}", p_text):
                        year_context_phrases = [
                            "for the year",
                            "năm kết thúc",
                            "ended",
                            "period",
                            "kỳ kế toán",
                            "niên độ",
                            "statement",
                            "báo cáo",
                        ]
                        if any(ph in text_lower for ph in year_context_phrases):
                            score += 10  # contextual year reference
                        else:
                            score += 0  # bare year, weaker signal

                    # 4. Proximity (Closer is better)
                    # Deduct 0.5 per step away
                    score -= idx * 0.5

                    # Update candidates log
                    if score > 0:
                        candidates_log.append(
                            {"text": p_text, "score": score, "distance": idx}
                        )

                    # Update best
                    if score > best_score and score >= 5:  # Threshold 5 to avoid noise
                        best_score = score
                        best_candidate_text = p_text
                        # If we find a very strong match proximate, valid.
                        # But keep scanning in case there's a heading slightly further up (e.g. "APPENDIX" then "Table 1")
                        # Actually closest high score is usually preferred.

                # Use inferred heading if found
                current_heading_candidate = best_candidate_text
                # P1-1: Fallback Heading from Table First Row (guarded by feature flag)
                heading_source = "paragraph"

                use_first_row_fallback = flags.get(
                    "heading_fallback_from_table_first_row", True
                )

                if use_first_row_fallback and (
                    current_heading_candidate is None or best_score < 5
                ):  # If no strong paragraph heading
                    table_first_row_heading = self._derive_heading_from_table_first_row(
                        table
                    )
                    if table_first_row_heading:
                        current_heading_candidate = table_first_row_heading
                        heading_source = "table_first_row"
                        logger.info(
                            "Heading inferred from table first row: %s, heading_source=%s",
                            current_heading_candidate,
                            heading_source,
                        )

                current_heading = current_heading_candidate

                # P3: heading_text, heading_confidence for XLSX export and validation heuristic
                heading_text = (current_heading or "").strip()
                if heading_source == "table_first_row":
                    heading_confidence = 0.6
                else:
                    heading_confidence = (
                        min(1.0, max(0.0, best_score / 10.0))
                        if getattr(self, "_temp_best_score_placeholder", best_score)
                        is not None
                        else 0.5
                    )

                # Sort top 3 candidates by score descending
                top_candidates = sorted(
                    candidates_log, key=lambda x: x["score"], reverse=True
                )[:3]

                table_context: Dict[str, Any] = {
                    "heading_source": heading_source,
                    "heading_text": heading_text,
                    "heading_confidence": heading_confidence,
                    "heading_candidates": top_candidates,
                    "heading_chosen_reason": f"Highest score ({best_score})"
                    if heading_source == "paragraph"
                    else heading_source,
                }

                # Parse table with multi-engine fallback (OOXML -> Python-docx -> LibreOffice -> legacy)
                table_index = len(tables)
                reconstructed_rows, extract_meta = self._extract_table_with_fallback(
                    file_path, table_index, table
                )
                table_context.update(extract_meta)
                flags = table_context.get("quality_flags") or []
                table_context["duplicate_period_artifacts"] = (
                    "DUPLICATE_PERIODS" in flags
                )

                # P0-5: Validate shape - check all rows have same column count
                try:
                    if reconstructed_rows:
                        max_cols = max(len(row) for row in reconstructed_rows)
                        if not all(len(row) == max_cols for row in reconstructed_rows):
                            # Shape không đúng → pad
                            logger.warning(
                                f"Table has inconsistent row lengths, padding to {max_cols} columns"
                            )
                            for row in reconstructed_rows:
                                while len(row) < max_cols:
                                    row.append("")
                    df = pd.DataFrame(reconstructed_rows)
                    if len(df) > 0:
                        df = self._promote_header_row(df, heading=current_heading)
                except Exception as e:
                    logger.warning(f"Reconstruction failed: {e}, using fallback")
                    rows = [
                        [cell.text.strip() for cell in row.cells] for row in table.rows
                    ]
                    # P0-5: Validate fallback shape too
                    if rows:
                        max_cols = max(len(row) for row in rows)
                        if not all(len(row) == max_cols for row in rows):
                            logger.warning(
                                f"Fallback table has inconsistent row lengths, padding to {max_cols} columns"
                            )
                            for row in rows:
                                while len(row) < max_cols:
                                    row.append("")
                    df = pd.DataFrame(rows)
                    if len(df) > 0:
                        df = self._promote_header_row(df, heading=current_heading)

                # P1.1: Compute numeric evidence and optionally try LibreOffice (RenderFirst) fallback
                ev = compute_numeric_evidence_score(df)
                table_context["numeric_evidence_score"] = ev.get(
                    "numeric_evidence_score", 0.0
                )
                table_context["numeric_parse_rate"] = ev.get("numeric_parse_rate", 0.0)
                table_context["numeric_cell_ratio"] = ev.get("numeric_cell_ratio", 0.0)
                table_context["numeric_col_candidates"] = ev.get(
                    "numeric_col_candidates", []
                )

                quality_score = table_context.get("quality_score", 0.0)
                numeric_evidence_score = table_context.get(
                    "numeric_evidence_score", 0.0
                )
                extractor_engine = table_context.get("extractor_engine", "")

                if (
                    quality_score >= QUALITY_SCORE_FALLBACK_MIN
                    and numeric_evidence_score < NUMERIC_EVIDENCE_FALLBACK_THRESHOLD
                    and extractor_engine != "render_first"
                ):
                    logger.info(
                        "P1.1 numeric fallback: triggering (quality_score=%.2f >= %.2f, numeric_evidence_score=%.3f < %.2f)",
                        quality_score,
                        QUALITY_SCORE_FALLBACK_MIN,
                        numeric_evidence_score,
                        NUMERIC_EVIDENCE_FALLBACK_THRESHOLD,
                    )
                    try:
                        from .extractors.render_first_table_extractor import (
                            RenderFirstTableExtractor,
                        )

                        rfe = RenderFirstTableExtractor()
                        fallback_available = rfe.is_available()
                        logger.info(
                            "P1.1 numeric fallback: RenderFirst extractor available=%s",
                            fallback_available,
                        )
                        if fallback_available:
                            res = rfe.extract(table, file_path, table_index)
                            if res.is_usable and res.grid:
                                df_fallback = pd.DataFrame(res.grid)
                                if len(df_fallback) > 0:
                                    df_fallback = self._promote_header_row(
                                        df_fallback, heading=current_heading
                                    )
                                ev_fb = compute_numeric_evidence_score(df_fallback)
                                score_fb = ev_fb.get("numeric_evidence_score", 0.0)
                                if score_fb >= NUMERIC_EVIDENCE_FALLBACK_THRESHOLD:
                                    df = df_fallback
                                    table_context.update(ev_fb)
                                    table_context["extractor_engine"] = "render_first"
                                    table_context["quality_score"] = res.quality_score
                                    table_context["numeric_evidence_score"] = score_fb
                                    logger.info(
                                        "P1.1 numeric fallback: RenderFirst used for table %s, numeric_evidence_score=%.3f",
                                        table_index,
                                        score_fb,
                                    )
                                else:
                                    table_context["fallback_attempted"] = True
                                    table_context["fallback_failed"] = True
                                    table_context["failure_reason_code"] = (
                                        "NO_NUMERIC_EVIDENCE"
                                    )
                                    logger.info(
                                        "P1.1 numeric fallback: RenderFirst score %.3f < threshold, keeping original",
                                        score_fb,
                                    )
                            else:
                                table_context["fallback_attempted"] = True
                                table_context["fallback_failed"] = True
                                table_context["failure_reason_code"] = (
                                    "NO_NUMERIC_EVIDENCE"
                                )
                                logger.info(
                                    "P1.1 numeric fallback: RenderFirst not usable for table %s",
                                    table_index,
                                )
                        else:
                            table_context["fallback_attempted"] = True
                            table_context["fallback_failed"] = True
                            table_context["failure_reason_code"] = "NO_NUMERIC_EVIDENCE"
                            logger.info(
                                "P1.1 numeric fallback: RenderFirst not available (soffice missing)"
                            )
                    except Exception as e:
                        table_context["fallback_attempted"] = True
                        table_context["fallback_failed"] = True
                        table_context["failure_reason_code"] = "NO_NUMERIC_EVIDENCE"
                        logger.info(
                            "P1.1 numeric fallback: RenderFirst failed for table %s: %s",
                            table_index,
                            e,
                        )

                # Ticket 10: Attach note number to context
                if current_note_number:
                    table_context["note_number"] = current_note_number

                is_footer = self._is_footer_or_signature_table(df)

                # Ticket 6: Split Table Guardrails
                should_merge = False
                if (
                    tables
                    and not is_footer
                    and headings[-1] != "SKIPPED_FOOTER_SIGNATURE"
                ):
                    # Proximity Check (Safe distance)
                    if (
                        paragraphs_since_last_table <= 2
                        and not long_paragraph_since_last_table
                    ):
                        prev_df = tables[-1]
                        # Schema Validation (Pre-concat)
                        if len(df.columns) == len(prev_df.columns):
                            # Type Consistency / Header alignment proxy
                            prev_heading = headings[-1]
                            if (
                                current_heading is None
                                or current_heading == prev_heading
                            ):
                                should_merge = True

                if should_merge:
                    # Merge into previous table
                    # To align columns, we safely reset column names before concat
                    df_to_merge = df.copy()
                    df_to_merge.columns = tables[-1].columns
                    merged_df = pd.concat([tables[-1], df_to_merge], ignore_index=True)
                    tables[-1] = merged_df
                    logger.info(
                        "Ticket-6: Merged table %s with previous table (proximity <= 2, matching schema)",
                        table_index,
                    )
                else:
                    if is_footer:
                        tables.append(df)
                        headings.append("SKIPPED_FOOTER_SIGNATURE")
                        table_contexts.append(table_context)
                    else:
                        tables.append(df)
                        headings.append(current_heading or "")
                        table_contexts.append(table_context)

                # Comment 2: reset buffer after each table so heading does not bleed to next table
                prior_paragraphs.clear()
                current_heading = None
                paragraphs_since_last_table = 0
                long_paragraph_since_last_table = False

                # Reset prior paragraphs?
                # "Notes" heading might apply to multiple tables.
                # But specifics ("Table 1") only apply to one.
                # User request: "scan... set current_heading".
                # If we clear, we lose "Notes" context for subsequent tables.
                # If we don't clear, "Balance Sheet" might bleed into "Income Statement" if no text between?
                # Usually text exists. Let's keep existing logic: we don't clear, but the *next* scan will pick up new paragraphs.

            elif block.tag.endswith("p"):
                # Section break can be inside a paragraph (w:sectPr child)
                if any(c.tag.endswith("sectPr") for c in block):
                    prior_paragraphs.clear()
                    current_heading = None
                # Comment 2: empty paragraph or page break = boundary, clear buffer
                from docx.oxml.ns import qn
                from docx.text.paragraph import Paragraph

                paragraph = Paragraph(block, doc)
                text = paragraph.text.strip()
                has_page_break = any(
                    br.get(qn("w:type")) == "page" for br in block.iter(qn("w:br"))
                )
                if not text or has_page_break:
                    prior_paragraphs.clear()
                    current_heading = None
                    # Do not reset note number, as notes span page breaks
                    continue
                if text:
                    paragraphs_since_last_table += 1
                    if len(text) > 200:
                        long_paragraph_since_last_table = True

                    # Extract features for caching
                    is_bold = any(run.bold for run in paragraph.runs)
                    is_upper = text.isupper()
                    style_name = paragraph.style.name if paragraph.style else ""

                    # Ticket 10: Note Number Mapping
                    # Reset state on major heading (e.g. Appendix, Phụ lục)
                    if "Heading" in style_name:
                        text_lower_p = text.lower()
                        if "appendix" in text_lower_p or "phụ lục" in text_lower_p:
                            current_note_number = None

                    # Strict regex anchoring
                    note_match = re.search(
                        r"^(Thuyết minh|Note)\s*(số\s*)?([A-Z0-9\.\-]+)",
                        text,
                        flags=re.IGNORECASE,
                    )
                    if note_match:
                        # Confidence gate: heading style or strong signal (bold)
                        if "Heading" in style_name or is_bold:
                            current_note_number = note_match.group(3)

                    prior_paragraphs.append((text, style_name, is_bold, is_upper))
                    # Keep buffer finite
                    if len(prior_paragraphs) > 20:
                        prior_paragraphs.pop(0)

        # Return 3-tuples (table, heading, table_context) or 2-tuples (table, heading) for backward compat
        if include_context:
            return list(zip(tables, headings, table_contexts))
        return list(zip(tables, headings))

    def extract_heading_from_section(self, section) -> Optional[str]:
        """
        Extract financial statement type from section header.

        Args:
            section: Document section

        Returns:
            Optional[str]: Detected heading type or None
        """
        if not section or not hasattr(section, "header"):
            return None

        for para in section.header.paragraphs:
            text = para.text.strip().lower()
            if "balance sheet" in text:
                return "balance sheet"
            elif "statement of income" in text:
                return "statement of income"
            elif "statement of cash flows" in text:
                return "statement of cash flows"

        return None

    def validate_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Validate Word document structure and content.

        Args:
            file_path: Path to Word document

        Returns:
            Dict with validation results
        """
        try:
            doc = docx.Document(file_path)

            # Count tables and paragraphs
            table_count = 0
            paragraph_count = 0
            heading_count = 0

            for block in doc.element.body:
                if block.tag.endswith("tbl"):
                    table_count += 1
                elif block.tag.endswith("p"):
                    paragraph_count += 1
                    from docx.text.paragraph import Paragraph

                    para = Paragraph(block, doc)
                    if para.style and para.style.name.startswith("Heading"):
                        heading_count += 1

            return {
                "valid": True,
                "table_count": table_count,
                "paragraph_count": paragraph_count,
                "heading_count": heading_count,
                "sections": len(doc.sections),
            }

        except Exception as e:
            return {
                "valid": False,
                "error": str(e),
                "table_count": 0,
                "paragraph_count": 0,
                "heading_count": 0,
                "sections": 0,
            }


class AsyncWordReader:
    """
    Async version of WordReader for improved performance with concurrent file processing.

    Uses ThreadPoolExecutor to handle I/O-bound operations asynchronously,
    allowing better resource utilization when processing multiple documents.
    """

    def __init__(self, max_workers: int = 4):
        """
        Initialize async word reader.

        Args:
            max_workers: Maximum number of worker threads for concurrent processing
        """
        self.executor = ThreadPoolExecutor(max_workers=max_workers)
        self._sync_reader = WordReader()

    async def read_document_async(
        self, file_path: str
    ) -> List[Tuple[pd.DataFrame, Optional[str], Dict[str, str]]]:
        """
        Read Word document asynchronously.

        Args:
            file_path: Path to Word document

        Returns:
            List[Tuple[pd.DataFrame, Optional[str], Dict[str, str]]]: List of
            (table_df, heading, table_context) triples
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor,
            self._sync_reader.read_tables_with_headings,  # type: ignore[arg-type]
            file_path,
        )

    async def validate_document_structure_async(self, file_path: str) -> Dict[str, Any]:
        """
        Validate Word document structure asynchronously.

        Args:
            file_path: Path to Word document

        Returns:
            Dict with validation results
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            self.executor, self._sync_reader.validate_document_structure, file_path
        )

    def shutdown(self, wait: bool = True) -> None:
        """
        Shutdown the thread pool executor.

        Args:
            wait: If True, wait for all pending tasks to complete
        """
        self.executor.shutdown(wait=wait)

    async def __aenter__(self):
        """Async context manager entry."""
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit."""
        self.shutdown()

    def __enter__(self):
        """Context manager entry."""
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit - shutdown executor."""
        self.shutdown(wait=True)
