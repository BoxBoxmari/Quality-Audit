import pandas as pd
from docx import Document

from quality_audit.io.word_reader import WordReader
from quality_audit.services.audit_service import AuditService


def _make_sample_docx(tmp_path):
    doc_path = tmp_path / "sample.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(str(doc_path))
    return str(doc_path)


def test_render_first_triggered_before_python_docx_on_low_quality(tmp_path):
    from unittest.mock import patch

    sample_word_path = _make_sample_docx(tmp_path)
    reader = WordReader()
    doc = __import__("docx").Document(sample_word_path)
    table = doc.tables[0]

    class _LowQualityOOXML:
        grid = []
        quality_score = 0.4
        quality_flags = []
        failure_reason_code = None
        invariant_violations = []

        @property
        def is_usable(self):
            return False

    class _RenderFirstOK:
        is_usable = True
        grid = [["H1", "H2"], ["1", "2"]]
        quality_score = 0.8
        quality_flags = []

    with patch("quality_audit.io.word_reader.OOXMLTableGridExtractor") as mock_ooxml:
        mock_ooxml.return_value.extract.return_value = _LowQualityOOXML()
        with patch(
            "quality_audit.io.extractors.render_first_table_extractor.RenderFirstTableExtractor"
        ) as mock_render_first:
            mock_render_first.return_value.is_available.return_value = True
            mock_render_first.return_value.extract.return_value = _RenderFirstOK()
            with patch(
                "quality_audit.io.word_reader.get_feature_flags",
                return_value={
                    "extraction_fallback_prefer_advanced_before_legacy": True,
                    "extraction_render_first_triggered_mode": "signals_only",
                    "heading_inference_v2": True,
                    "heading_fallback_from_table_first_row": True,
                },
            ):
                mock_ocr = type("MockOCR", (), {"is_available": lambda self: True})()
                with patch(
                    "quality_audit.io.extractors.ocr.get_best_ocr_engine",
                    return_value=mock_ocr,
                ):
                    _, meta = reader._extract_table_with_fallback(
                        sample_word_path, 0, table
                    )

    assert "engine_attempts" in meta
    assert "ooxml" in meta["engine_attempts"]
    assert "render_first" in meta["engine_attempts"]
    assert meta["extractor_engine"] == "render_first"


def test_render_first_not_triggered_when_ooxml_quality_is_good(tmp_path):
    from unittest.mock import patch

    sample_word_path = _make_sample_docx(tmp_path)
    reader = WordReader()
    doc = __import__("docx").Document(sample_word_path)
    table = doc.tables[0]

    class _GoodOOXML:
        grid = [["A", "B"], ["1", "2"]]
        quality_score = 0.9
        quality_flags = []
        failure_reason_code = None
        invariant_violations = []

        @property
        def is_usable(self):
            return True

    with patch("quality_audit.io.word_reader.OOXMLTableGridExtractor") as mock_ooxml:
        mock_ooxml.return_value.extract.return_value = _GoodOOXML()
        with patch(
            "quality_audit.config.feature_flags.get_feature_flags",
            return_value={
                "extraction_fallback_prefer_advanced_before_legacy": True,
                "extraction_render_first_triggered_mode": "signals_only",
            },
        ):
            _, meta = reader._extract_table_with_fallback(sample_word_path, 0, table)

    assert meta.get("engine_attempts") == ["ooxml"]
    assert meta["extractor_engine"] == "ooxml"


def test_parity_no_numeric_dispatch_returns_info_no_numeric_evidence(monkeypatch):
    monkeypatch.setattr(
        "quality_audit.services.audit_service.get_feature_flags",
        lambda: {"legacy_parity_mode": True},
    )
    monkeypatch.setattr(
        "quality_audit.services.audit_service.ValidatorFactory.get_validator",
        lambda *_args, **_kwargs: (None, "SKIPPED_NO_NUMERIC_EVIDENCE"),
    )

    service = AuditService()
    service.context.set_last_classification_context(
        {
            "classifier_primary_type": "FS_BALANCE_SHEET",
            "classifier_confidence": 0.95,
        }
    )

    df = pd.DataFrame(
        {"Description": ["Assets", "Liabilities"], "Notes": ["n/a", "n/a"]}
    )
    result = service._validate_single_table(
        df,
        heading="Balance Sheet",
        table_context={"heading_source": "paragraph"},
    )

    assert result.rule_id == "NO_NUMERIC_EVIDENCE"
    assert result.status_enum == "INFO"
