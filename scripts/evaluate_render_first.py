"""
Render-first evaluation harness.

Computes CER (Character Error Rate) and structure fidelity metrics
on a gold set of tables for CI quality tracking.

Usage:
    python scripts/evaluate_render_first.py [--gold-dir tests/fixtures/gold_set]

Output:
    reports/render_first_eval.json
"""

import argparse
import json
import logging
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


@dataclass
class CellEvaluation:
    """Per-cell evaluation result."""

    row: int
    col: int
    expected: str
    actual: str
    cer: float  # Character Error Rate


@dataclass
class TableEvaluation:
    """Per-table evaluation result."""

    table_id: str
    source_file: str
    extraction_success: bool
    failure_reason: Optional[str] = None
    cells: List[CellEvaluation] = field(default_factory=list)
    mean_cer: float = 0.0
    structure_match: bool = True
    expected_rows: int = 0
    expected_cols: int = 0
    actual_rows: int = 0
    actual_cols: int = 0
    quality_score: float = 0.0
    quality_flags: List[str] = field(default_factory=list)


@dataclass
class EvaluationReport:
    """Full evaluation report."""

    total_tables: int = 0
    successful_extractions: int = 0
    failed_extractions: int = 0
    mean_cer: float = 0.0
    structure_match_rate: float = 0.0
    tables: List[TableEvaluation] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "total_tables": self.total_tables,
            "successful_extractions": self.successful_extractions,
            "failed_extractions": self.failed_extractions,
            "mean_cer": self.mean_cer,
            "structure_match_rate": self.structure_match_rate,
            "tables": [
                {
                    "table_id": t.table_id,
                    "source_file": t.source_file,
                    "extraction_success": t.extraction_success,
                    "failure_reason": t.failure_reason,
                    "mean_cer": t.mean_cer,
                    "structure_match": t.structure_match,
                    "expected_rows": t.expected_rows,
                    "expected_cols": t.expected_cols,
                    "actual_rows": t.actual_rows,
                    "actual_cols": t.actual_cols,
                    "quality_score": t.quality_score,
                    "quality_flags": t.quality_flags,
                }
                for t in self.tables
            ],
        }


def compute_cer(expected: str, actual: str) -> float:
    """
    Compute Character Error Rate between expected and actual strings.

    CER = (insertions + deletions + substitutions) / len(expected)
    Uses Levenshtein distance.
    """
    if not expected:
        return 0.0 if not actual else 1.0

    try:
        from Levenshtein import distance
    except ImportError:
        # Fallback to simple implementation if python-Levenshtein not available
        logger.warning("python-Levenshtein not installed, using fallback CER")
        return _fallback_cer(expected, actual)

    return distance(expected, actual) / len(expected)


def _fallback_cer(expected: str, actual: str) -> float:
    """Simple CER fallback without external deps."""
    if expected == actual:
        return 0.0
    if not expected:
        return 1.0 if actual else 0.0

    # Simple comparison - not Levenshtein but gives rough estimate
    matches = sum(1 for a, b in zip(expected, actual) if a == b)
    return 1.0 - (matches / max(len(expected), len(actual)))


def load_ground_truth(gt_path: Path) -> List[List[str]]:
    """Load ground truth CSV file."""
    import csv

    with open(gt_path, encoding="utf-8") as f:
        reader = csv.reader(f)
        return list(reader)


def evaluate_table(
    docx_path: Path,
    gt_path: Path,
    table_id: str,
) -> TableEvaluation:
    """
    Evaluate render-first extraction against ground truth.

    Args:
        docx_path: Path to DOCX file
        gt_path: Path to ground truth CSV
        table_id: Identifier for this table

    Returns:
        TableEvaluation with metrics
    """
    from docx import Document

    from quality_audit.io.extractors.render_first_table_extractor import (
        RenderFirstTableExtractor,
    )

    evaluation = TableEvaluation(
        table_id=table_id,
        source_file=str(docx_path),
    )

    # Load ground truth
    try:
        gt_grid = load_ground_truth(gt_path)
        evaluation.expected_rows = len(gt_grid)
        evaluation.expected_cols = len(gt_grid[0]) if gt_grid else 0
    except Exception as e:
        evaluation.extraction_success = False
        evaluation.failure_reason = f"Failed to load ground truth: {e}"
        return evaluation

    # Run extraction
    try:
        doc = Document(docx_path)
        if not doc.tables:
            evaluation.extraction_success = False
            evaluation.failure_reason = "No tables found in document"
            return evaluation

        extractor = RenderFirstTableExtractor(save_debug_artifacts=True)
        result = extractor.extract(doc.tables[0], str(docx_path), 0)

        if not result.is_usable:
            evaluation.extraction_success = False
            evaluation.failure_reason = result.failure_reason_code
            evaluation.quality_flags = list(result.quality_flags)
            return evaluation

        evaluation.extraction_success = True
        evaluation.actual_rows = result.rows
        evaluation.actual_cols = result.cols
        evaluation.quality_score = result.quality_score
        evaluation.quality_flags = list(result.quality_flags)

        # Check structure match
        evaluation.structure_match = (
            evaluation.expected_rows == evaluation.actual_rows
            and evaluation.expected_cols == evaluation.actual_cols
        )

        # Compute per-cell CER
        actual_grid = result.grid
        cers = []

        for row_idx in range(min(len(gt_grid), len(actual_grid))):
            gt_row = gt_grid[row_idx]
            actual_row = actual_grid[row_idx] if row_idx < len(actual_grid) else []

            for col_idx in range(
                min(len(gt_row), len(actual_row) if actual_row else 0)
            ):
                expected_text = gt_row[col_idx].strip()
                actual_text = (
                    actual_row[col_idx].strip() if col_idx < len(actual_row) else ""
                )
                cer = compute_cer(expected_text, actual_text)
                cers.append(cer)

                evaluation.cells.append(
                    CellEvaluation(
                        row=row_idx,
                        col=col_idx,
                        expected=expected_text,
                        actual=actual_text,
                        cer=cer,
                    )
                )

        evaluation.mean_cer = sum(cers) / len(cers) if cers else 0.0

    except Exception as e:
        evaluation.extraction_success = False
        evaluation.failure_reason = str(e)

    return evaluation


def run_evaluation(gold_dir: Path, output_path: Path) -> EvaluationReport:
    """
    Run full evaluation on gold set.

    Expects:
        gold_dir/table_001.docx + gold_dir/table_001_gt.csv
        gold_dir/table_002.docx + gold_dir/table_002_gt.csv
        ...

    Args:
        gold_dir: Directory containing gold set files
        output_path: Path to save JSON report

    Returns:
        EvaluationReport
    """
    report = EvaluationReport()

    # Find all DOCX files
    docx_files = sorted(gold_dir.glob("*.docx"))
    logger.info("Found %d DOCX files in gold set", len(docx_files))

    for docx_path in docx_files:
        table_id = docx_path.stem
        gt_path = docx_path.with_suffix("").parent / f"{table_id}_gt.csv"

        if not gt_path.exists():
            logger.warning("No ground truth for %s, skipping", table_id)
            continue

        logger.info("Evaluating %s...", table_id)
        evaluation = evaluate_table(docx_path, gt_path, table_id)
        report.tables.append(evaluation)
        report.total_tables += 1

        if evaluation.extraction_success:
            report.successful_extractions += 1
        else:
            report.failed_extractions += 1
            logger.warning("  FAIL: %s", evaluation.failure_reason)

    # Compute aggregate metrics
    if report.tables:
        successful = [t for t in report.tables if t.extraction_success]
        if successful:
            report.mean_cer = sum(t.mean_cer for t in successful) / len(successful)
            report.structure_match_rate = sum(
                1 for t in successful if t.structure_match
            ) / len(successful)

    # Save report
    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(report.to_dict(), f, indent=2)

    logger.info("Evaluation complete. Report saved to %s", output_path)
    logger.info(
        "Results: %d/%d successful (%.1f%%), Mean CER: %.4f, Structure match: %.1f%%",
        report.successful_extractions,
        report.total_tables,
        100 * report.successful_extractions / max(1, report.total_tables),
        report.mean_cer,
        100 * report.structure_match_rate,
    )

    return report


def main():
    parser = argparse.ArgumentParser(
        description="Evaluate render-first table extraction on gold set"
    )
    parser.add_argument(
        "--gold-dir",
        type=Path,
        default=Path("tests/fixtures/gold_set"),
        help="Directory containing gold set DOCX files and ground truth CSVs",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("reports/render_first_eval.json"),
        help="Output JSON report path",
    )
    parser.add_argument(
        "--cer-threshold",
        type=float,
        default=0.05,
        help="CER threshold for CI pass (default: 0.05)",
    )
    parser.add_argument(
        "--structure-threshold",
        type=float,
        default=0.90,
        help="Structure match rate threshold for CI pass (default: 0.90)",
    )

    args = parser.parse_args()

    if not args.gold_dir.exists():
        logger.error("Gold set directory not found: %s", args.gold_dir)
        sys.exit(1)

    report = run_evaluation(args.gold_dir, args.output)

    # Check CI thresholds
    ci_pass = True
    if report.mean_cer > args.cer_threshold:
        logger.error(
            "CI FAIL: Mean CER %.4f exceeds threshold %.4f",
            report.mean_cer,
            args.cer_threshold,
        )
        ci_pass = False

    if report.structure_match_rate < args.structure_threshold:
        logger.error(
            "CI FAIL: Structure match rate %.2f below threshold %.2f",
            report.structure_match_rate,
            args.structure_threshold,
        )
        ci_pass = False

    sys.exit(0 if ci_pass else 1)


if __name__ == "__main__":
    main()
