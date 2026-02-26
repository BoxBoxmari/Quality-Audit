"""
Tests for multi-engine fallback: OOXML -> render-first (conditional) -> Python-docx -> legacy.
"""

import pytest

from quality_audit.io.word_reader import WordReader


@pytest.fixture
def sample_word_path(tmp_path):
    """Create a sample .docx with one table for fallback tests."""
    from docx import Document

    doc_path = tmp_path / "sample.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"
    doc.save(str(doc_path))
    return str(doc_path)


def test_extract_table_with_fallback_returns_grid_and_metadata(sample_word_path):
    """_extract_table_with_fallback returns (rows, metadata) with extractor_engine and quality_score."""
    reader = WordReader()
    doc = __import__("docx").Document(sample_word_path)
    table = doc.tables[0]
    grid, meta = reader._extract_table_with_fallback(sample_word_path, 0, table)
    assert isinstance(grid, list)
    assert isinstance(meta, dict)
    assert "extractor_engine" in meta
    assert meta["extractor_engine"] in (
        "ooxml",
        "python_docx",
        "libreoffice",
        "legacy",
        "render_first",
    )
    assert "quality_score" in meta
    assert "quality_flags" in meta


def test_read_tables_with_headings_includes_table_context(sample_word_path):
    """read_tables_with_headings returns table_context with extractor_engine for traceability."""
    reader = WordReader()
    result = reader.read_tables_with_headings(sample_word_path)
    assert isinstance(result, list)
    for item in result:
        assert isinstance(item, tuple)
        assert len(item) >= 3
        _, _, table_context = item[0], item[1], item[2]
        assert isinstance(table_context, dict)
        assert "extractor_engine" in table_context


def test_render_first_triggered_before_python_docx_on_low_quality(sample_word_path):
    """When OOXML returns low quality_score, render-first is attempted before python-docx."""
    from unittest.mock import patch

    reader = WordReader()
    doc = __import__("docx").Document(sample_word_path)
    table = doc.tables[0]

    class MockOOXMLResult:
        grid = []
        quality_score = 0.4
        quality_flags = []
        failure_reason_code = None
        invariant_violations = []

        @property
        def is_usable(self):
            return False

    class MockRenderFirstResult:
        is_usable = True
        grid = [["H1", "H2"], ["1", "2"]]
        quality_score = 0.8
        quality_flags = []

    with patch("quality_audit.io.word_reader.OOXMLTableGridExtractor") as mock_ooxml:
        mock_ooxml.return_value.extract.return_value = MockOOXMLResult()
        with patch(
            "quality_audit.io.extractors.render_first_table_extractor.RenderFirstTableExtractor"
        ) as mock_render_first:
            mock_render_first.return_value.is_available.return_value = True
            mock_render_first.return_value.extract.return_value = (
                MockRenderFirstResult()
            )
            with patch(
                "quality_audit.config.feature_flags.get_feature_flags",
                return_value={
                    "extraction_fallback_prefer_advanced_before_legacy": True,
                    "extraction_render_first_triggered_mode": "signals_only",
                },
            ):
                grid, meta = reader._extract_table_with_fallback(
                    sample_word_path, 0, table
                )
    assert "engine_attempts" in meta
    assert "ooxml" in meta["engine_attempts"]
    assert "render_first" in meta["engine_attempts"]
    assert meta["extractor_engine"] == "render_first"


def test_render_first_not_triggered_when_signals_only_and_ooxml_good(sample_word_path):
    """When OOXML returns high quality, render-first is not called (engine_attempts == ['ooxml'])."""
    from unittest.mock import patch

    reader = WordReader()
    doc = __import__("docx").Document(sample_word_path)
    table = doc.tables[0]

    class MockOOXMLResultGood:
        grid = [["A", "B"], ["1", "2"]]
        quality_score = 0.9
        quality_flags = []
        failure_reason_code = None
        invariant_violations = []

        @property
        def is_usable(self):
            return True

    with patch("quality_audit.io.word_reader.OOXMLTableGridExtractor") as mock_ooxml:
        mock_ooxml.return_value.extract.return_value = MockOOXMLResultGood()
        with patch(
            "quality_audit.config.feature_flags.get_feature_flags",
            return_value={
                "extraction_fallback_prefer_advanced_before_legacy": True,
                "extraction_render_first_triggered_mode": "signals_only",
            },
        ):
            grid, meta = reader._extract_table_with_fallback(sample_word_path, 0, table)
    assert meta.get("engine_attempts") == ["ooxml"]
    assert meta["extractor_engine"] == "ooxml"
