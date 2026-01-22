import re
import sys

import docx
import pandas as pd
from openpyxl import Workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter, quote_sheetname
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo

# =========================
# Constants & Regex helpers
# =========================
_SHEET_NAME_CLEAN_RE = re.compile(r"[:\\/*?\[\]]")
# cross_check_cache = {}
BSPL_cross_check_cache = {}
BSPL_cross_check_mark = []
_CODE_COL_NAME = "code"
_NOTE_COL_NAME = "note"
_CODE_VALID_RE = re.compile(r"^[0-9]+[A-Z]?$")
_HEADER_DATE_RE = re.compile(r"\d{4}|\d{1,2}/\d{1,2}/\d{2,4}")
TABLES_NEED_COLUMN_CHECK = {
    "long-term prepaid expenses",
    "tangible fixed assets",
    "intangible fixed assets",
    "chi phí trả trước dài hạn",
    "tài sản cố định hữu hình",
    "tài sản cố định vô hình",
    "taxes payable to state treasury",
    "thuế và các khoản phải nộp nhà nước",
    "borrowings",
    "borrowings, bonds and finance lease liabilities",
    "short-term borrowings",
    "long-term borrowings",
}
# Dạng bảng 1: bảng có thể có subtotal và dòng cross ref là ở grand total
CROSS_CHECK_TABLES_FORM_1 = {
    "accounts receivable from customers",
    "accounts receivable from customers detailed by significant customer",
    "accounts receivable from customers detailed by significant customers",
    "receivables on construction contracts according to stages of completion",
    "payables on construction contracts according to stages of completion"
    "deferred tax assets and liabilities",
    "deferred tax assets",
    "deferred tax liabilities" "accrued expenses",
    "accrued expenses – short-term",
    "accrued expenses - short-term",
    "accrued expenses – long-term",
    "accrued expenses - long-term",
    "unearned revenue",
    "unearned revenue – short-term",
    "unearned revenue – long-term",
    "other payables",
    "other payables – short-term",
    "other payables – long-term",
    "long-term borrowings",
    "long-term borrowings, bonds and financial lease liabilities",
    "long-term bonds and financial lease liabilities",
    "long-term financial lease liabilities",
    "long-term bonds",
}

# Dạng bảng 2: bảng có thể có subtotal, cross ref ở cả subtotal & grand total
CROSS_CHECK_TABLES_FORM_2 = {
    "revenue from sales of goods and provision of services",
    "revenue from sales of goods",
    "revenue from provision of services",
}

# Dạng bảng 2: bảng không có subtotal nhưng không phải standard table
# (standard table là bảng có 3 cột, côt 2 là CY balance, cột 3 là PY
# balance)
CROSS_CHECK_TABLES_FORM_3 = {
    "investments",
    "trading securities",
    "held-to-maturity investments",
    "equity investments in other entities",
    "equity investments in other entity",
    "bad and doubtful debts",
    "shortage of assets awaiting resolution",
    "inventories",
    "long-term work in progress",
    "construction in progress",
    "long-term prepaid expenses",
    "accounts payable to suppliers",
    "accounts payable to suppliers detailed by significant suppliers",
    "accounts payable to suppliers detailed by significant supplier"
    "taxes and others payable to state treasury",
    "taxes and others receivable from state treasury",
    "taxes and others receivable from and payable to state treasury",
    "taxes receivable from state treasury",
    "taxes payable to state treasury",
    "short-term borrowings",
    "short-term borrowings, bonds and finance lease liabilities",
    "short-term bonds and finance lease liabilities",
    "short-term bonds",
    "preference shares",
    "provisions",
    "short-term provisions",
    "long-term provisions",
    "share capital",
    "contributed capital",
}

valid_codes = {"222", "223", "225", "226", "228", "229", "231", "232"}
TABLES_NEED_CHECK_SEPARATELY = {
    "tangible fixed assets",
    "intangible fixed assets",
    "tài sản cố định hữu hình",
    "tài sản cố định vô hình",
}
TABLES_WITHOUT_TOTAL = {
    "business costs by element",
    "Production and business costs by elements",
    "non-cash investing activity",
    "non-cash investing activities",
    "significant transactions with related parties",
    "significant transactions with related companies",
    "corresponding figures",
}
RE_PARTY_TABLE = {
    "related parties",
    "related party",
    "related companies",
    "related company",
}

# Màu tô
GREEN_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
BLUE_FILL = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
RED_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
INFO_FILL = PatternFill(start_color="DAE8FC", end_color="DAE8FC", fill_type="solid")

GREEN_FONT = Font(color="32CD32")  # Xanh
RED_FONT = Font(color="FF0000")  # Đỏ
RIGHT_ALIGN = Alignment(horizontal="right")  # Canh phải


# =========
# Utilities
# =========
def shorten_sheet_name(name, max_length=20):
    """Rút gọn tên sheet để không vượt quá 20 ký tự và loại bỏ ký tự đặc biệt."""
    name = _SHEET_NAME_CLEAN_RE.sub("_", str(name))
    return name[:max_length]


def normalize_numeric_column(value):
    """Chuyển chuỗi số thành số thực, xử lý dấu phẩy và dấu ngoặc."""
    if isinstance(value, str):
        value = value.replace(",", "").replace("(", "-").replace(")", "")
    return pd.to_numeric(value, errors="coerce")


def _dfpos_to_excel(row_idx_0based: int, col_idx_0based: int):
    """
    dataframe_to_rows(..., header=True) -> Excel hàng 1 là header cột.
    Nên: Excel_row = df_row + 2 ; Excel_col = df_col + 1
    """
    return row_idx_0based + 2, col_idx_0based + 1


def is_red_fill(cell):
    fill = cell.fill
    return (
        fill.start_color.rgb == RED_FILL.start_color.rgb
        and fill.end_color.rgb == RED_FILL.end_color.rgb
        and fill.fill_type == RED_FILL.fill_type
    )


def apply_cell_marks(ws, marks: list, start_row, start_col):
    """
    marks: list các dict {'row': int(df_row), 'col': int(df_col), 'ok': bool|None, 'comment': str|None}
    """
    for m in marks:
        r, c = _dfpos_to_excel(m["row"], m["col"])
        cell = ws.cell(row=r + start_row, column=c + start_col)
        if m.get("ok") is True:
            if not is_red_fill(cell):  # chỉ tô xanh nếu chưa bị tô đỏ
                cell.fill = GREEN_FILL
        elif m.get("ok") is False:
            cell.fill = RED_FILL
        # comment cho ô lỗi
        if m.get("comment"):
            try:
                # Tránh đè comment cũ: nối thêm
                if cell.comment:
                    new_text = cell.comment.text + "\n" + str(m["comment"])
                else:
                    new_text = str(m["comment"])
                cell.comment = Comment(text=new_text, author="AutoCheck")
            except Exception:
                # comment lỗi thì bỏ qua, không ảnh hưởng phần còn lại
                pass


def apply_crossref_marks(ws, marks: list, start_row, start_col):
    """
    marks: list các dict {'row': int(df_row), 'col': int(df_col), 'ok': bool|None, 'comment': str|None}
    """
    for m in marks:
        r, c = _dfpos_to_excel(m["row"], m["col"])
        cell = ws.cell(row=r + start_row, column=c + start_col)
        cell.alignment = RIGHT_ALIGN  # Áp dụng canh phải cho mọi ô

        # Đổi màu chữ theo trạng thái ok
        if m.get("ok") is True:
            cell.font = GREEN_FONT
            cell.value = "PASS"

        elif m.get("ok") is False:
            cell.font = RED_FONT
            cell.value = "FAIL"
            try:
                # Tránh đè comment cũ: nối thêm
                if cell.comment:
                    new_text = cell.comment.text + "\n" + str(m["comment"])
                else:
                    new_text = str(m["comment"])
                cell.comment = Comment(text=new_text, author="AutoCheck")
            except Exception:
                # comment lỗi thì bỏ qua, không ảnh hưởng phần còn lại
                pass


# ------------------------------
# Reading Word tables with header
# ------------------------------
def read_word_tables_with_headings(file_path):
    """Đọc bảng từ file Word và lấy heading gần nhất trước mỗi bảng."""
    doc = docx.Document(file_path)
    tables = []
    headings = []
    current_heading = None
    sec = 0
    current_section = doc.sections[sec]

    for block in doc.element.body:
        if block.tag.endswith("tbl"):
            #            Nếu chưa có heading, kiểm tra header của section
            if (
                current_heading is None
                or current_heading == "balance sheet"
                or current_heading == "statement of income"
            ):
                for para in current_section.header.paragraphs:
                    text = para.text.strip().lower()
                    if "balance sheet" in text:
                        current_heading = "balance sheet"
                        break
                    elif "statement of income" in text:
                        current_heading = "statement of income"
                        break
                    elif "statement of cash flows" in text:
                        current_heading = "statement of cash flows"
                        break
            # Đọc bảng
            table = docx.table.Table(block, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(pd.DataFrame(rows))
            headings.append(current_heading)
        elif block.tag.endswith("p"):
            paragraph = docx.text.paragraph.Paragraph(block, doc)
            # Kiểm tra nếu có sectPr (section break)
            sectPr = block.xpath(".//w:sectPr")
            if sectPr:
                # Cập nhật section hiện tại
                sec = sec + 1
                current_section = doc.sections[sec]
            if paragraph.style.name.startswith("Heading"):
                current_heading = paragraph.text.strip()

    return list(zip(tables, headings))


# --------------------------------------
# Find total row index
# --------------------------------------
def find_total_row_index(df, _strict=True):
    """
    Xác định dòng tổng theo heuristics (không dùng keywords):
    - Ưu tiên dòng số "cuối cùng" mà dòng ngay trước đó TRỐNG HOÀN TOÀN (không chữ, không số).
    - Nếu không tìm thấy mẫu này:
        + _strict=True  => coi như bảng KHÔNG có dòng tổng -> trả về None
        + _strict=False => fallback: trả về dòng số cuối cùng của bảng
    """

    # ---- Helpers nội bộ ----
    def _as_numeric_series(row):
        return row.map(normalize_numeric_column)

    def _is_numeric_row(row) -> bool:
        ser = _as_numeric_series(row)
        return ser.notna().any()

    def _is_empty_row(row) -> bool:
        # “Trống” = không có số và không có text hữu ích (loại '-' '—' ngoặc, dấu phẩy, khoảng trắng)
        def _strip_text(x):
            s = str(x).strip()
            s = s.replace("-", "").replace("–", "").replace("—", "")
            s = s.replace("(", "").replace(")", "").replace(",", "")
            return s.strip()

        has_text = any(_strip_text(c) != "" for c in row)
        has_num = _is_numeric_row(row)
        return (not has_text) and (not has_num)

    # ---- Tìm tất cả các dòng có số ----
    numeric_rows = [i for i in range(len(df)) if _is_numeric_row(df.iloc[i])]
    if not numeric_rows:
        return None  # Bảng hoàn toàn không có số

    # ---- Từ dưới lên: chọn dòng số mà dòng trước đó "trống hoàn toàn" ----
    for i in reversed(numeric_rows):
        prev_empty = True
        if i - 1 >= 0:
            prev_empty = _is_empty_row(df.iloc[i - 1])
        if prev_empty:
            return i

    # ---- Không tìm thấy mẫu "trước trống" ----
    if _strict:
        return None  # KHÔNG có dòng tổng theo heuristics
    else:
        return numeric_rows[-1]  # Fallback: dùng dòng số cuối (nếu anh muốn nới tay)


# ======================
# Checking logic (router)
# ======================
def _parse_num(v):
    val = normalize_numeric_column(v)
    return 0.0 if pd.isna(val) else float(val)


def _norm_code(x: str) -> str:
    s = str(x).strip()
    s = s.replace("_", "").replace("**", "").replace("\u2212", "-").replace("–", "-")
    s = re.sub(r"[^0-9A-Za-z]", "", s)
    return s.upper()


# --------------------------
# Balance Sheet casting
# --------------------------
_BALANCE_RULES = {
    "100": ["110", "120", "130", "140", "150"],
    "110": ["111", "112"],
    "120": ["121", "122", "123"],
    "130": ["131", "132", "133", "134", "135", "136", "137", "139"],
    "140": ["141", "149"],
    "150": ["151", "152", "153", "154", "155"],
    "200": ["210", "220", "230", "240", "250", "260"],
    "210": ["211", "212", "213", "214", "215", "216", "219"],
    "220": ["221", "224", "227"],
    "221": ["222", "223"],
    "224": ["225", "226"],
    "227": ["228", "229"],
    "230": ["231", "232"],
    "240": ["241", "242"],
    "250": ["251", "252", "253", "254", "255"],
    "260": ["261", "262", "263", "268", "269"],
    "270": ["100", "200"],
    "300": ["310", "330"],
    "310": [
        "311",
        "312",
        "313",
        "314",
        "315",
        "316",
        "317",
        "318",
        "319",
        "320",
        "321",
        "322",
        "323",
        "324",
    ],
    "330": [
        "331",
        "332",
        "333",
        "334",
        "335",
        "336",
        "337",
        "338",
        "339",
        "340",
        "341",
        "342",
        "343",
    ],
    "400": ["410", "430"],
    "410": [
        "411",
        "412",
        "413",
        "414",
        "415",
        "416",
        "417",
        "418",
        "419",
        "420",
        "421",
        "422",
        "429",
    ],
    "411": ["411A", "411B"],
    "421": ["421A", "421B"],
    "430": ["431", "432"],
    "440": ["300", "400"],
}
_BALANCE_NEW_RULES = {
    "100": ["110", "120", "130", "140", "150", "160"],
    "110": ["111", "112"],
    "120": ["121", "122", "123", "124", "125", "126"],
    "130": ["131", "132", "133", "134", "135", "136", "137"],
    "140": ["141", "142"],
    "150": ["151", "152", "153"],
    "160": ["161", "162", "163", "164", "165"],
    "200": ["210", "220", "230", "240", "250", "260", "270"],
    "210": ["211", "212", "213", "214", "215", "216"],
    "220": ["221", "224", "227"],
    "221": ["222", "223"],
    "224": ["225", "226"],
    "227": ["228", "229"],
    "230": ["231", "236", "237", "238"],
    "231": ["232", "233"],
    "233": ["234", "235"],
    "240": ["241", "242"],
    "250": ["251", "252"],
    "260": ["261", "262", "263", "264", "265", "266"],
    "270": ["261", "262", "263", "264"],
    "280": ["100", "200"],
    "300": ["310", "330"],
    "310": [
        "311",
        "312",
        "313",
        "314",
        "315",
        "316",
        "317",
        "318",
        "319",
        "320",
        "321",
        "322",
        "323",
        "324",
        "325",
    ],
    "330": [
        "331",
        "332",
        "333",
        "334",
        "335",
        "336",
        "337",
        "338",
        "339",
        "340",
        "341",
        "342",
        "343",
        "344",
    ],
    "400": ["410", "430"],
    "410": [
        "411",
        "412",
        "413",
        "414",
        "415",
        "416",
        "417",
        "418",
        "419",
        "420",
        "421",
        "422",
        "429",
    ],
    "411": ["411A", "411B"],
    "421": ["421A", "421B"],
    "430": ["431", "432"],
    "440": ["300", "400"],
}


def _find_header_idx_has_code(df: pd.DataFrame):
    for i in range(len(df)):
        row_strs = df.iloc[i].astype(str).str.lower()
        if row_strs.str.contains(_CODE_COL_NAME).any():
            return i
    return None


def _check_balance_sheet(df: pd.DataFrame) -> dict:
    # 1) Tìm hàng tiêu đề có chữ 'Code'
    header_idx = _find_header_idx_has_code(df)
    if header_idx is None:
        return {
            "status": "WARN: Balance sheet - không tìm thấy cột 'Code' để kiểm tra",
            "marks": [],
        }

    # 2) Dựng bảng dữ liệu với header từ hàng 'Code'
    header = [str(c).strip() for c in df.iloc[header_idx].tolist()]
    tmp = df.iloc[header_idx + 1:].copy()
    tmp.columns = header

    # 3) Xác định cột 'Code' và 2 cột số (cột hiện tại & so sánh)
    code_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _CODE_COL_NAME), None
    )
    if code_col is None:
        return {
            "status": "WARN: Balance sheet - không xác định được cột 'Code'",
            "marks": [],
        }

    note_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _NOTE_COL_NAME), None
    )

    candidate_num_cols = [c for c in tmp.columns if _HEADER_DATE_RE.search(str(c))]
    if len(candidate_num_cols) >= 2:
        cur_col, prior_col = candidate_num_cols[:2]  # giữ thứ tự xuất hiện
    else:
        # fallback: dùng 2 cột cuối cùng
        cur_col, prior_col = tmp.columns[-2], tmp.columns[-1]

    # 4) Map code -> (current, prior) chỉ trong bảng hiện tại
    data = {}
    code_rowpos = {}  # map code -> tmp_row_index (0-based)

    for ridx, row in tmp.iterrows():
        code = _norm_code(row.get(code_col, ""))
        if not code:
            continue
        # Chấp nhận pattern 3 chữ số +/- hậu tố chữ (vd. 421B)
        if not re.match(r"^[0-9]+[A-Z]?$", code):
            continue
        cur_val = _parse_num(row.get(cur_col, ""))
        prior_val = _parse_num(row.get(prior_col, ""))
        # Nếu code lặp lại, ưu tiên bản có số (tránh ghi đè = 0)
        if code in data:
            old_cur, old_pr = data[code]
            if abs(cur_val) + abs(prior_val) == 0 and abs(old_cur) + abs(old_pr) != 0:
                continue
        data[code] = (cur_val, prior_val)
        if row.get(note_col, "") != "" "":
            acc_name = row.get(tmp.columns[0]).strip().lower()
            BSPL_cross_check_cache[acc_name] = (cur_val, prior_val)
        else:
            if code in valid_codes or code in [
                "141",
                "149",
                "251",
                "252",
                "253",
                "254",
            ]:
                if code in ["251", "252", "253"]:
                    try:
                        old_cur, old_pr = BSPL_cross_check_cache[
                            "investments in other entities"
                        ]
                    except BaseException:
                        old_cur = 0
                        old_pr = 0
                    BSPL_cross_check_cache["investments in other entities"] = (
                        cur_val + old_cur,
                        prior_val + old_pr,
                    )

                BSPL_cross_check_cache[code] = (cur_val, prior_val)
        # chỉ set hàng đầu tiên có số liệu
        code_rowpos.setdefault(code, ridx - tmp.index[0])  # quy về 0-based liên tục
    # vị trí cột trong df gốc
    try:
        cur_col_pos = header.index(cur_col)
        prior_col_pos = header.index(prior_col)
    except ValueError:
        cur_col_pos = len(header) - 2
        prior_col_pos = len(header) - 1

    # 5) Kiểm tra: skip nếu không có mã cha; skip nếu không có TẤT CẢ mã con;
    #    nếu thiếu MỘT VÀI mã con -> coi các mã thiếu = 0 để tính.
    issues = []
    marks = []

    for parent, children in _BALANCE_RULES.items():
        parent_norm = _norm_code(parent)
        if parent_norm not in data:
            continue

        have_any = False
        cur_sum = prior_sum = 0.0
        missing = []
        for ch in children:
            ch_norm = _norm_code(ch)
            if ch_norm in data:
                ccur, cprior = data[ch_norm]
                cur_sum += ccur
                prior_sum += cprior
                have_any = True
            else:
                missing.append(ch_norm)

        if not have_any:
            # Không có TẤT CẢ mã con trong BẢNG NÀY -> skip
            continue

        parent_cur, parent_prior = data[parent_norm]
        diff_cur = cur_sum - parent_cur
        diff_prio = prior_sum - parent_prior
        is_ok_cy = abs(diff_cur) == 0
        is_ok_py = abs(diff_prio) == 0

        # Nếu biết vị trí dòng mã cha -> đánh dấu 2 ô số (cur/prior)
        if parent_norm in code_rowpos:
            df_row = (
                header_idx + 1 + code_rowpos[parent_norm]
            )  # vị trí trong df (0-based)
            comment = (
                f"{parent_norm} = sum({','.join(children)}); "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f}; "
                f"Thực tế={parent_cur:,.0f}/{parent_prior:,.0f}; "
                f"Δ={diff_cur:,.0f}/{diff_prio:,.0f}"
                + (f"; Thiếu={','.join(missing)}" if missing else "")
            )
            marks.append(
                {
                    "row": df_row,
                    "col": cur_col_pos,
                    "ok": is_ok_cy,
                    "comment": (None if is_ok_cy else comment),
                }
            )
            marks.append(
                {
                    "row": df_row,
                    "col": prior_col_pos,
                    "ok": is_ok_py,
                    "comment": (None if is_ok_py else comment),
                }
            )

        if not is_ok_cy or not is_ok_py:
            issues.append(comment)

    # 6) Kết luận trả về (ghi trực tiếp vào cột trạng thái của sheet bảng này)
    if not issues:
        status = "PASS: Balance sheet - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues) - 10} dòng)" if len(issues) > 10 else ""
        status = f"FAIL: Balance sheet - kiểm tra công thức: {
            len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": []}


# ---------------------------------
# Statement of Income casting
# ---------------------------------
def _sum_weighted(data, children):
    have_any = False
    cur_sum = prior_sum = 0.0
    missing = []
    for token in children:
        token = str(token).strip()
        sign = -1 if token.startswith("-") else 1
        code = token[1:] if sign == -1 else token
        cn = _norm_code(code)
        if cn in data:
            ccur, cprior = data[cn]
            cur_sum += sign * ccur
            prior_sum += sign * cprior
            have_any = True
        else:
            missing.append(cn if sign == 1 else f"-{cn}")
    return have_any, cur_sum, prior_sum, missing


def _check_income_statement(df: pd.DataFrame) -> dict:
    # === Kiểm tra công thức Statement of Income ngay trong nhánh này ===
    # 1) Tìm hàng tiêu đề có chữ 'Code'
    header_idx = _find_header_idx_has_code(df)
    if header_idx is None:
        return {
            "status": "WARN: Statement of income - không tìm thấy cột 'Code' để kiểm tra",
            "marks": [],
            "cross_ref_marks": [],
        }

    # 2) Dựng bảng dữ liệu với header từ hàng 'Code'
    header = [str(c).strip() for c in df.iloc[header_idx].tolist()]
    tmp = df.iloc[header_idx + 1:].copy()
    tmp.columns = header
    tmp = tmp.reset_index(drop=True)

    # 3) Xác định cột 'Code' và 2 cột số (cột hiện tại & so sánh)
    code_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _CODE_COL_NAME), None
    )
    if code_col is None:
        return {
            "status": "WARN: Statement of income - không xác định được cột 'Code'",
            "marks": [],
            "cross_ref_marks": [],
        }

    note_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _NOTE_COL_NAME), None
    )
    candidate_num_cols = [c for c in tmp.columns if _HEADER_DATE_RE.search(str(c))]
    if len(candidate_num_cols) >= 2:
        cur_col, prior_col = candidate_num_cols[:2]  # giữ thứ tự xuất hiện
    else:
        # fallback: dùng 2 cột cuối cùng
        cur_col, prior_col = tmp.columns[-2], tmp.columns[-1]

    # 4) Map code -> (current, prior) từ bảng hiện tại
    data = {}
    code_rowpos = {}
    for ridx, row in tmp.iterrows():
        code = _norm_code(row.get(code_col, ""))
        # Chấp nhận các code dạng số (giữ leading zero), có thể có hậu tố chữ (hiếm)
        if code == "" or not re.match(r"^[0-9]+[A-Z]?$", code):
            continue
        cur_val = _parse_num(row.get(cur_col, ""))
        prior_val = _parse_num(row.get(prior_col, ""))
        if code in data:
            old_cur, old_pr = data[code]
            if abs(cur_val) + abs(prior_val) == 0 and abs(old_cur) + abs(old_pr) != 0:
                continue
        data[code] = (cur_val, prior_val)
        if row.get(note_col, "") != "" "":
            acc_name = row.get(tmp.columns[0]).strip().lower()
            BSPL_cross_check_cache[acc_name] = (cur_val, prior_val)
            if code in ["51", "52"]:
                try:
                    old_cur, old_pr = BSPL_cross_check_cache["income tax"]
                except BaseException:
                    old_cur = 0
                    old_pr = 0
                BSPL_cross_check_cache["income tax"] = (
                    cur_val + old_cur,
                    prior_val + old_pr,
                )

        else:
            if code in ["50"]:
                BSPL_cross_check_cache[code] = (cur_val, prior_val)

        code_rowpos.setdefault(code, ridx)

    try:
        cur_col_pos = header.index(cur_col)
        prior_col_pos = header.index(prior_col)
    except ValueError:
        cur_col_pos = len(header) - 2
        prior_col_pos = len(header) - 1

    issues = []
    marks = []

    def check(parent, children, label=None):
        parent_norm = _norm_code(parent)
        if parent_norm not in data:
            return  # skip nếu không có mã cha
        have_any, cur_sum, prior_sum, missing = _sum_weighted(data, children)
        if not have_any:
            return  # skip nếu không có TẤT CẢ mã con
        ac_cur, ac_pr = data[parent_norm]
        dc = cur_sum - ac_cur
        dp = prior_sum - ac_pr
        is_ok_cy = abs(dc) == 0
        is_ok_py = abs(dp) == 0
        if parent_norm in code_rowpos:
            df_row = header_idx + 1 + code_rowpos[parent_norm]
            comment = (
                f"{parent_norm} = {' + '.join(children).replace('+ -', ' - ')}; "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f}; Thực tế={ac_cur:,.0f}/{ac_pr:,.0f}; Δ={dc:,.0f}/{dp:,.0f}"
                + (f"; Thiếu={','.join(missing)}" if missing else "")
            )
            marks.append(
                {
                    "row": df_row,
                    "col": cur_col_pos,
                    "ok": is_ok_cy,
                    "comment": (None if is_ok_cy else comment),
                }
            )
            marks.append(
                {
                    "row": df_row,
                    "col": prior_col_pos,
                    "ok": is_ok_py,
                    "comment": (None if is_ok_py else comment),
                }
            )
        if not is_ok_cy or not is_ok_py:
            issues.append(comment)

    # 7) Áp công thức
    # 10 = 01 - 02
    check("10", ["01", "-02"])
    # 20 = 10 - 11
    check("20", ["10", "-11"])
    # 20 = 01 - 11
    check("20", ["01", "-02", "-11"])
    # 30 = 20 + (21 - 22) + 24 - (25 + 26)
    check("30", ["20", "21", "-22", "24", "-25", "-26"])
    # 40 = 31 - 32
    check("40", ["31", "-32"])
    # 50 = 30 + 40
    check("50", ["30", "40"])
    # 60 = 50 - 51 - 52
    check("60", ["50", "-51", "-52"])
    # 60 = 61 + 62, chỉ kiểm khi có đủ 61 & 62
    check("60", ["61", "62"])

    # 8) Kết luận trả về

    if not issues:
        status = "PASS: Statement of income - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues) - 10} dòng)" if len(issues) > 10 else ""
        status = f"FAIL: Statement of income - kiểm tra công thức: {
            len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": []}


# ----------------------------
# Statement of Cash Flows casting
# ----------------------------
def _check_cash_flows(df: pd.DataFrame) -> dict:
    # 1) Tìm hàng tiêu đề có chữ 'Code'
    header_idx = _find_header_idx_has_code(df)
    if header_idx is None:
        return {
            "status": "WARN: Cash flows - không tìm thấy cột 'Code' để kiểm tra",
            "marks": [],
            "cross_ref_marks": [],
        }

    # 2) Dựng bảng dữ liệu với header từ hàng 'Code'
    header = [str(c).strip() for c in df.iloc[header_idx].tolist()]
    tmp = df.iloc[header_idx + 1:].copy()
    tmp.columns = header
    tmp = tmp.reset_index(drop=True)

    # 3) Xác định cột 'Code' và 2 cột số (cột hiện tại & so sánh)
    code_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _CODE_COL_NAME), None
    )
    if code_col is None:
        return {
            "status": "WARN: Cash flows - không xác định được cột 'Code'",
            "marks": [],
            "cross_ref_marks": [],
        }

    candidate_num_cols = [c for c in tmp.columns if _HEADER_DATE_RE.search(str(c))]
    if len(candidate_num_cols) >= 2:
        cur_col, prior_col = candidate_num_cols[:2]  # giữ thứ tự xuất hiện
    else:
        cur_col, prior_col = tmp.columns[-2], tmp.columns[-1]

    # 4) Map code -> (current, prior) và lưu vị trí hàng
    data = {}
    code_rowpos = {}
    rows_cache = []  # [(ridx, code_norm, cur, prior)]

    for ridx, row in tmp.iterrows():
        code = _norm_code(row.get(code_col, ""))
        cur_val = _parse_num(row.get(cur_col, ""))
        prior_val = _parse_num(row.get(prior_col, ""))

        # Giữ cache thô để còn bắt "dòng không có code nhưng có tổng" (phục vụ Code 18)
        rows_cache.append((ridx, code, cur_val, prior_val))

        # Chỉ cộng dồn cho các mã code hợp lệ dạng số (có thể kèm suffix chữ)
        # Ví dụ: '05', '14', '21', hoặc '411A' (nếu có)
        if code and _CODE_VALID_RE.match(code):
            # PASS: Cộng dồn nếu cùng mã xuất hiện nhiều dòng (vd. nhiều dòng Code '05')
            agg_cur, agg_pr = data.get(code, (0.0, 0.0))
            data[code] = (agg_cur + cur_val, agg_pr + prior_val)
            code_rowpos.setdefault(code, ridx)

    # 5) Bắt trường hợp dòng "không có code nhưng có tổng" ngay trước Code 14~20 => đặt tạm là Code 18
    #    Code 18 = 08 + 09 + 10 + 11 + 12 + 13
    target_set = {str(i) for i in range(14, 21)}  # '14'..'20'
    first_block_idx = None
    for ridx, code, _, _ in rows_cache:
        if code in target_set:
            first_block_idx = ridx
            break
    if first_block_idx is not None and first_block_idx > 0:
        idx18 = first_block_idx - 1
        prev_code = rows_cache[idx18][1]
        prev_cur = rows_cache[idx18][2]
        prev_pr = rows_cache[idx18][3]
        # Nếu ngay trước là dòng không có code nhưng có số ⇒ coi đó là Code '18'
        if (prev_code == "" or prev_code is None) and (
            abs(prev_cur) + abs(prev_pr) != 0
        ):
            data["18"] = (prev_cur, prev_pr)
            code_rowpos.setdefault("18", idx18)
    try:
        cur_col_pos = header.index(cur_col)
        prior_col_pos = header.index(prior_col)
    except ValueError:
        cur_col_pos = len(header) - 2
        prior_col_pos = len(header) - 1

    issues = []
    marks = []

    def check(parent, children, label=None):
        parent_norm = _norm_code(parent)
        if parent_norm not in data:
            return
        have_any, cur_sum, prior_sum, missing = _sum_weighted(data, children)
        if not have_any:
            return
        ac_cur, ac_pr = data[parent_norm]
        dc = cur_sum - ac_cur
        dp = prior_sum - ac_pr

        is_ok_cy = abs(dc) == 0
        is_ok_py = abs(dp) == 0

        if parent_norm in code_rowpos:
            df_row = header_idx + 1 + code_rowpos[parent_norm]
            comment = (
                f"{parent_norm} = {' + '.join(children).replace('+ -', ' - ')}; "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f}; Thực tế={ac_cur:,.0f}/{ac_pr:,.0f}; Δ={dc:,.0f}/{dp:,.0f}"
                + (f"; Thiếu={','.join(missing)}" if missing else "")
            )
            marks.append(
                {
                    "row": df_row,
                    "col": cur_col_pos,
                    "ok": is_ok_cy,
                    "comment": (None if is_ok_cy else comment),
                }
            )
            marks.append(
                {
                    "row": df_row,
                    "col": prior_col_pos,
                    "ok": is_ok_py,
                    "comment": (None if is_ok_py else comment),
                }
            )
        if not is_ok_cy or not is_ok_py:
            issues.append(comment)
        """if abs(dc) != 0 or abs(dp) != 0:
            formula_txt = ' + '.join(children).replace('+ -', ' - ')
            issues.append(
                f"{parent_norm} = {formula_txt} | "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f} | "
                f"Thực tế={ac_cur:,.0f}/{ac_pr:,.0f} | "
                f"Δ={dc:,.0f}/{dp:,.0f}"
                + (f" | Thiếu={','.join(missing)}" if missing else "")
            )"""

    # 8) Áp công thức cho Cash flows
    # 08 = 01 + 02 + 03 + 04 + 05 + 06 + 07
    check("08", ["01", "02", "03", "04", "05", "06", "07"])
    # 18 = 08 + 09 + 10 + 11 + 12 + 13 (nếu đã bắt được dòng 18, thì so khớp)
    if "18" in data:
        check("18", ["08", "09", "10", "11", "12", "13"])
    # 20 = 08 + 09 + 10 + 11 + 12 + 13 + 14 + 15 + 16 + 17
    check("20", ["08", "09", "10", "11", "12", "13", "14", "15", "16", "17"])
    # 30 = 21 + 22 + 23 + 24 + 25 + 26 + 27
    check("30", ["21", "22", "23", "24", "25", "26", "27"])
    # 40 = 31 + 32 + 33 + 34 + 35 + 36
    check("40", ["31", "32", "33", "34", "35", "36"])
    # 50 = 20 + 30 + 40
    check("50", ["20", "30", "40"])
    # 70 = 50 + 60 + 61
    check("70", ["50", "60", "61"])

    # 9) Kết luận trả về

    if not issues:
        status = "PASS: Statement of cash flows - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues) - 10} dòng)" if len(issues) > 10 else ""
        status = f"FAIL: Statement of cash flows - kiểm tra công thức: {
            len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": []}


# ---------------------------------------------
# Changes in owners’ equity casting
# ---------------------------------------------
def _norm_text(s: str) -> str:
    s = str(s).strip().lower()
    s = s.replace("’", "'").replace("‘", "'")  # chuẩn hoá apostrophe
    s = re.sub(r"\s+", " ", s)
    return s


def _row_starts_with_balance_at(row) -> bool:
    # Tìm ô text đầu tiên của dòng; nếu bắt đầu bằng "balance at" => True
    for c in row:
        t = str(c).strip()
        if t != "":
            return _norm_text(t).startswith("balance at")
    return False


def _check_changes_in_equity(df: pd.DataFrame) -> dict:
    # --- Chuẩn hoá dữ liệu số ---
    df_numeric = df.map(normalize_numeric_column)
    issues = []
    marks = []
    # ======================================================
    # A) KIỂM TRA HÀNG: tìm các dòng "Balance at ..."
    # ======================================================
    balance_rows_idx = [
        i for i in range(len(df)) if _row_starts_with_balance_at(df.iloc[i])
    ]
    balance_rows_idx = sorted(set(balance_rows_idx))
    # Hàng header giả định ở index 0
    header_idx = 0
    first_data_idx = 1 if len(df) > 1 else 0
    # 1) Dòng "Balance at" thứ 2 (nếu có)
    if len(balance_rows_idx) >= 2:
        idx2 = balance_rows_idx[1]
        # Tổng của tất cả các dòng trước đó (trừ header) => từ first_data_idx .. idx2-1
        expected_series_2 = df_numeric.iloc[first_data_idx:idx2].sum(skipna=True)
        actual_series_2 = df_numeric.iloc[idx2]
        for col in range(df.shape[1]):
            exp_val = expected_series_2.iloc[col]
            act_val = actual_series_2.iloc[col]
            if not (pd.isna(exp_val) and pd.isna(act_val)):
                exp_val = 0.0 if pd.isna(exp_val) else float(exp_val)
                act_val = 0.0 if pd.isna(act_val) else float(act_val)
                diff = exp_val - act_val
                is_ok = abs(diff) == 0
                comment = f"HÀNG: Balance at (thứ 2) - Cột {
                    col +
                    1}: Tính={
                    exp_val:,.0f}, Trên bảng={
                    act_val:,.0f}, Δ={
                    diff:,.0f}"
                marks.append(
                    {
                        "row": idx2,
                        "col": col,
                        "ok": is_ok,
                        "comment": (None if is_ok else comment),
                    }
                )
                if not is_ok:
                    issues.append(comment)
    else:
        issues.append("INFO: HÀNG: Không tìm đủ 'Balance at' (thứ 2) để kiểm tra.")
    # 2) Dòng "Balance at" thứ 3 (nếu có)
    if len(balance_rows_idx) >= 3:
        idx3 = balance_rows_idx[2]
        # Tổng của tất cả các dòng trước đó tới và bao gồm cả dòng "Balance at" thứ 2
        # => từ first_data_idx .. idx3-1 (bao gồm idx2 nếu idx2 < idx3)
        expected_series_3 = df_numeric.iloc[idx2:idx3].sum(skipna=True)
        actual_series_3 = df_numeric.iloc[idx3]
        for col in range(df.shape[1]):
            exp_val = expected_series_3.iloc[col]
            act_val = actual_series_3.iloc[col]
            if not (pd.isna(exp_val) and pd.isna(act_val)):
                exp_val = 0.0 if pd.isna(exp_val) else float(exp_val)
                act_val = 0.0 if pd.isna(act_val) else float(act_val)
                diff = exp_val - act_val
                is_ok = abs(diff) == 0
                comment = f"HÀNG: Balance at (thứ 3) - Cột {
                    col +
                    1}: Tính={
                    exp_val:,.0f}, Trên bảng={
                    act_val:,.0f}, Δ={
                    diff:,.0f}"
                marks.append(
                    {
                        "row": idx3,
                        "col": col,
                        "ok": is_ok,
                        "comment": (None if is_ok else comment),
                    }
                )
                if not is_ok:
                    issues.append(comment)
    else:
        issues.append("INFO: HÀNG: Không tìm đủ 'Balance at' (thứ 3) để kiểm tra.")
    # ======================================================
    # B) KIỂM TRA CỘT: Total owners’ equity / Total
    # ======================================================
    # Xác định tiêu đề cột (dòng header): dùng row 0
    header_row = df.iloc[header_idx].astype(str).tolist() if len(df) > 0 else []
    header_norm = [_norm_text(x) for x in header_row]

    # Chuẩn hoá các biến thể của tiêu đề
    def is_total_owners_equity(text_norm: str) -> bool:
        # chấp nhận total owners' equity / total owners’ equity (curly / straight
        # apostrophe)
        return (
            "total owners' equity" in text_norm or "total owners’ equity" in text_norm
        )

    def is_total_col(text_norm: str) -> bool:
        # cột "Total" (tránh nhầm với "Total owners' equity")
        return (len(text_norm) < 15) and "total" in text_norm

    # Tìm chỉ số cột
    toe_idx = None
    total_idx = None
    for j, t in enumerate(header_norm):
        if toe_idx is None and is_total_owners_equity(t):
            toe_idx = j
        elif total_idx is None and is_total_col(t):
            total_idx = j

    # Hàm kiểm tra một cột tổng theo quy tắc: giá trị ô = SUM tất cả các cột
    # TRƯỚC đó (có thể nhiều cột text -> NaN)
    def check_column_sum(
        target_col_idx: int,
        start_col_idx=int,
        include_upto: int = None,
        label: str = "CỘT",
    ):
        """
        target_col_idx: cột đích để so sánh.
        include_upto: cột tới hạn để cộng (mặc định = target_col_idx-1).
                    Nếu muốn "Total" bao gồm đến cột "Total owners’ equity" (khi toe trước total) thì
                    truyền include_upto = max(include_upto, toe_idx).
        """
        nonlocal issues
        if target_col_idx is None or target_col_idx <= 0:
            return
        end_idx = (
            target_col_idx - 1
            if include_upto is None
            else min(include_upto, target_col_idx - 1)
        )
        if end_idx < 0:
            return

        for r in range(first_data_idx, len(df)):
            row_series = df_numeric.iloc[r]
            left_part = row_series.iloc[start_col_idx: end_idx + 1]
            expected = left_part.sum(skipna=True)
            actual = row_series.iloc[target_col_idx]
            if pd.isna(expected) and pd.isna(actual):
                continue
            expected = 0.0 if pd.isna(expected) else float(expected)
            actual = 0.0 if pd.isna(actual) else float(actual)
            diff = expected - actual
            is_ok = abs(diff) == 0
            comment = f"{label}: Dòng {
                r + 1} - Tính={
                expected:,.0f}, Trên bảng={
                actual:,.0f}, Δ={
                diff:,.0f}"
            marks.append(
                {
                    "row": r,
                    "col": target_col_idx,
                    "ok": is_ok,
                    "comment": (None if is_ok else comment),
                }
            )
            if not is_ok:
                issues.append(comment)

    # Trường hợp có cả 2 cột
    if toe_idx is not None and total_idx is not None:
        # Nếu thứ tự bất thường (Total đứng trước TOEquity), vẫn kiểm từng cột độc lập bằng quy tắc “sum các cột trước đó”.
        # 1) Total owners’ equity = SUM(các cột trước đó)
        check_column_sum(
            target_col_idx=toe_idx,
            start_col_idx=1,
            include_upto=None,
            label="CỘT: Total owners’ equity",
        )
        # 2) Total = SUM(các cột trước đó tới và bao gồm TOEquity)
        #    => Nếu TOEquity đứng trước Total, include_upto = total_idx-1 (đã bao gồm TOEquity theo vị trí)
        #    Logic "sum các cột trước đó" đã tự bao hàm, nên chỉ cần:
        check_column_sum(
            target_col_idx=total_idx,
            start_col_idx=toe_idx,
            include_upto=None,
            label="CỘT: Total",
        )
    else:
        # Chỉ có 1 trong 2 cột
        if toe_idx is not None:
            check_column_sum(
                target_col_idx=toe_idx,
                start_col_idx=1,
                include_upto=None,
                label="CỘT: Total owners’ equity",
            )
        if total_idx is not None:
            check_column_sum(
                target_col_idx=total_idx,
                start_col_idx=1,
                include_upto=None,
                label="CỘT: Total",
            )

    # ======================================================
    # KẾT LUẬN
    # ======================================================
    # Lọc bỏ các thông tin "INFO" nếu anh muốn chỉ báo lỗi cứng; ở đây em để
    # nguyên để biết tình trạng dữ liệu.
    hard_issues = [m for m in issues if not m.startswith("INFO")]
    if not hard_issues:
        status = (
            "PASS: Changes in owners’ equity - kiểm tra công thức: KHỚP (0 sai lệch)"
        )
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues) - 10} dòng)" if len(issues) > 10 else ""
        status = f"FAIL: Changes in owners’ equity - kiểm tra công thức: {
            len(hard_issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": []}


# -----------------------------------------------------------------------
# Income tax disclosure including DTAs and DTLs
# -----------------------------------------------------------------------
def _check_income_tax_rec_table(df: pd.DataFrame) -> dict:
    df_numeric = df.map(normalize_numeric_column)
    # Tìm dòng chứa "Accounting profit before tax" hoặc "Accounting loss before tax"
    profit_row_idx = None
    for i, row in df.iterrows():
        row_text = " ".join(str(cell).lower() for cell in row)
        if (
            "accounting profit before tax" in row_text
            or "accounting loss before tax" in row_text
            or "accounting profit/(loss) before tax" in row_text
            or "accounting (loss)/profit before tax" in row_text
        ):
            profit_row_idx = i
            account_name = "50"
            break

    if profit_row_idx is None:
        return {
            "status": "INFO: Không có dòng Accounting profit before tax hay Accounting loss before tax",
            "marks": [],
            "cross_ref_marks": [],
        }

    issues = []
    marks = []
    cross_ref_marks = []
    # cross check accounting profit/(loss) before tax
    CY_bal = (
        0
        if pd.isna(df_numeric.iloc[i, len(df.columns) - 2])
        else df_numeric.iloc[i, len(df.columns) - 2]
    )
    PY_bal = (
        0
        if pd.isna(df_numeric.iloc[i, len(df.columns) - 1])
        else df_numeric.iloc[i, len(df.columns) - 1]
    )
    if account_name not in BSPL_cross_check_mark:
        cross_check_with_BSPL(
            df,
            cross_ref_marks,
            issues,
            account_name,
            CY_bal,
            PY_bal,
            i,
            len(df.columns) - 2,
            0,
            -1,
        )
        BSPL_cross_check_mark.append(account_name)

    try:
        # Secure input validation delegated to new modular system
        from quality_audit.io.file_handler import get_validated_tax_rate

        tax_rate = get_validated_tax_rate()
        if tax_rate is None:
            return {
                "status": "WARN: Không thể đọc thuế suất từ người dùng",
                "marks": [],
                "cross_ref_marks": [],
            }
    except Exception:
        return {
            "status": "WARN: Không thể đọc thuế suất từ người dùng",
            "marks": [],
            "cross_ref_marks": [],
        }

    # Bước 1: So sánh %tax rate của dòng profit với dòng tax rate cho từng cột
    profit_row = df_numeric.iloc[profit_row_idx]
    tax_rate_row_idx = None
    for i in range(profit_row_idx + 1, len(df)):
        row_text = " ".join(str(cell).lower() for cell in df.iloc[i])
        if (
            "tax at the company’s tax rate" in row_text
            or "tax at the group’s tax rate" in row_text
        ):
            tax_rate_row_idx = i
            break

    if tax_rate_row_idx is not None:
        tax_row = df_numeric.iloc[tax_rate_row_idx]
        for col in range(len(df.columns)):
            profit_val = profit_row[col]
            tax_val = tax_row[col]
            if not pd.isna(profit_val) and not pd.isna(tax_val):
                expected_tax = profit_val * tax_rate
                diff = expected_tax - tax_val
                is_ok = abs(round(diff)) == 0
                comment = f"Bước 1 - Cột {
                    col +
                    1}: {
                    tax_rate *
                    100}% lợi nhuận = {
                    expected_tax:,.2f}, Thuế trên bảng = {
                    tax_val:,.2f}, Sai lệch = {
                    diff:,.2f}"
                marks.append(
                    {
                        "row": tax_rate_row_idx,
                        "col": col,
                        "ok": is_ok,
                        "comment": (None if is_ok else comment),
                    }
                )
                if not is_ok:
                    issues.append(comment)
        # Bước 2: Cộng dồn từ dòng sau dòng thuế đến dòng trống
        sum1 = [0.0] * len(df.columns)
        i = tax_rate_row_idx
        while i < len(df):
            row = df.iloc[i]
            if all(str(cell).strip() == "" for cell in row):
                break
            for col in range(len(df.columns)):
                val = df_numeric.iloc[i, col]
                if not pd.isna(val):
                    sum1[col] += val
            i += 1

        if i < len(df) - 1:
            total1_row = df_numeric.iloc[i + 1]
            for col in range(len(df.columns)):
                total_val = total1_row[col]
                if not pd.isna(total_val):
                    diff = sum1[col] - total_val
                    is_ok = abs(round(diff)) == 0
                    comment = f"Bước 2 - Cột {
                        col +
                        1}: Tổng chi tiết = {
                        sum1[col]:,.2f}, Tổng 1 = {
                        total_val:,.2f}, Sai lệch = {
                        diff:,.2f}"
                    marks.append(
                        {
                            "row": i + 1,
                            "col": col,
                            "ok": is_ok,
                            "comment": (None if is_ok else comment),
                        }
                    )
                    if not is_ok:
                        issues.append(comment)

            # Bước 3: Cộng tiếp các dòng sau nếu có số liệu
            sum2 = sum1.copy()
            j = i + 2
            while j < len(df) - 1:
                row = df.iloc[j]
                if all(str(cell).strip() == "" for cell in row):
                    break
                for col in range(len(df.columns)):
                    val = df_numeric.iloc[j, col]
                    if not pd.isna(val):
                        sum2[col] += val
                j += 1

            if j < len(df):
                total2_row = df_numeric.iloc[len(df) - 1]
                for col in range(len(df.columns)):
                    total_val = total2_row[col]
                    if not pd.isna(total_val):
                        diff = sum2[col] - total_val
                        is_ok = abs(round(diff)) == 0
                        comment = f"Bước 3 - Cột {
                            col +
                            1}: Tổng cộng dồn = {
                            sum2[col]:,.2f}, Tổng 2 = {
                            total_val:,.2f}, Sai lệch = {
                            diff:,.2f}"
                        marks.append(
                            {
                                "row": len(df) - 1,
                                "col": col,
                                "ok": is_ok,
                                "comment": (None if is_ok else comment),
                            }
                        )
                        if not is_ok:
                            issues.append(comment)

        account_name = "income tax"
        CY_bal = (
            0
            if pd.isna(df_numeric.iloc[len(df) - 1, len(df.columns) - 2])
            else df_numeric.iloc[len(df) - 1, len(df.columns) - 2]
        )
        PY_bal = (
            0
            if pd.isna(df_numeric.iloc[len(df) - 1, len(df.columns) - 1])
            else df_numeric.iloc[len(df) - 1, len(df.columns) - 1]
        )
        if account_name not in BSPL_cross_check_mark:
            cross_check_with_BSPL(
                df,
                cross_ref_marks,
                issues,
                account_name,
                CY_bal,
                PY_bal,
                len(df) - 1,
                len(df.columns) - 2,
                0,
                -1,
            )
            BSPL_cross_check_mark.append(account_name)

    if not issues:
        status = "PASS: Reconciliation of effective tax rate - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues) - 10} dòng)" if len(issues) > 10 else ""
        status = f"FAIL: Reconciliation of effective tax rate - kiểm tra công thức: {
            len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": cross_ref_marks}


def _check_income_tax_remaining_tables(df: pd.DataFrame) -> dict:
    df_numeric = df.map(normalize_numeric_column)
    issues = []
    marks = []

    def find_block_sum(start_idx):
        sum_vals = [0.0] * len(df.columns)
        count = 0
        i = start_idx + 1
        while i < len(df):
            row = df.iloc[i]
            if all(str(cell).strip() == "" for cell in row):
                break
            for col in range(len(df.columns)):
                val = df_numeric.iloc[i, col]
                if not pd.isna(val):
                    sum_vals[col] += val
            count += 1
            i += 1
        return sum_vals, count, i

    def compare_sum_with_total(sum_vals, total_row, end_row):
        for col in range(len(df.columns)):
            total_val = total_row[col]
            if not pd.isna(total_val):
                diff = sum_vals[col] - total_val
                is_ok = abs(round(diff)) == 0
                comment = f"Cột {
                    col +
                    1}: Tổng chi tiết = {
                    sum_vals[col]:,.2f}, Tổng trên bảng = {
                    total_val:,.2f}, Sai lệch = {
                    diff:,.0f}"
                marks.append(
                    {
                        "row": end_row + 1,
                        "col": col,
                        "ok": is_ok,
                        "comment": (None if is_ok else comment),
                    }
                )
                if not is_ok:
                    issues.append(comment)

    # Bước 1: Current tax expense
    current_idx = None
    for i, row in df.iterrows():
        row_text = " ".join(str(cell).lower() for cell in row)
        if "current tax expense" in row_text or "deferred tax assets" in row_text:
            current_idx = i
            break

    total1 = [0.0] * len(df.columns)
    if current_idx is not None:
        sum1, count1, end1 = find_block_sum(current_idx)
        if count1 > 1 and end1 < len(df) - 1:
            total1_row = df_numeric.iloc[end1 + 1]
            compare_sum_with_total(sum1, total1_row, end1)
        total1 = sum1

    # Bước 2: Deferred tax expense/(income)
    deferred_keywords = [
        "deferred tax expense/(income)",
        "deferred tax expense",
        "deferred tax income",
        "deferred tax (income)/expense",
        "deferred tax liabilities",
    ]
    deferred_idx = None
    for i, row in df.iterrows():
        row_text = " ".join(str(cell).lower() for cell in row)
        if any(k in row_text for k in deferred_keywords):
            deferred_idx = i
            break

    total2 = [0.0] * len(df.columns)
    if deferred_idx is not None:
        sum2, count2, end2 = find_block_sum(deferred_idx)
        if count2 > 1 and end2 < len(df) - 1:
            total2_row = df_numeric.iloc[end2 + 1]
            compare_sum_with_total(sum2, total2_row, end2)
        total2 = sum2

    # Bước 3: Tổng 1 + Tổng 2 so với dòng cuối bảng
    final_row = df_numeric.iloc[-1]
    for col in range(len(df.columns)):
        combined = total1[col] + total2[col]
        final_val = final_row[col]
        if not pd.isna(final_val):
            diff = combined - final_val
            is_ok = abs(round(diff)) == 0
            comment = f"Bước 3 - Cột {
                col +
                1}: Tổng cộng = {
                combined:,.2f}, Dòng cuối = {
                final_val:,.2f}, Sai lệch = {
                diff:,.0f}"
            marks.append(
                {
                    "row": len(df) - 1,
                    "col": col,
                    "ok": is_ok,
                    "comment": (None if is_ok else comment),
                }
            )
            if not is_ok:
                issues.append(comment)

    if not issues:
        status = "PASS Kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues) - 10} dòng)" if len(issues) > 10 else ""
        status = f"FAIL Kiểm tra công thức: {len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks}


# -----------------------------------------------------------------------
# FS casting
# -----------------------------------------------------------------------


def cross_check_with_BSPL(
    df,
    cross_ref_marks,
    issues,
    account_name,
    CY_bal,
    PY_bal,
    CY_row,
    CY_col,
    gap_row,
    gap_col,
):
    if account_name in BSPL_cross_check_cache:
        BSPL_CY_bal, BSPL_PY_bal = BSPL_cross_check_cache[account_name]
        diffCB = BSPL_CY_bal - CY_bal
        diffOB = BSPL_PY_bal - PY_bal
        is_okCB = abs(round(diffCB)) == 0
        is_okOB = abs(round(diffOB)) == 0

        if (
            (account_name in TABLES_NEED_CHECK_SEPARATELY)
            or (account_name in valid_codes)
            or (
                account_name
                in ["50", "construction in progress", "long-term prepaid expenses"]
            )
            or ("revenue" in account_name)
        ):
            CY_row = CY_row - 1
            CY_col = len(df.columns)

        commentCB = (
            f"BSPL = {BSPL_CY_bal:,.2f}, Note = {CY_bal:,.2f}, Sai lệch = {diffCB:,.0f}"
        )
        cross_ref_marks.append(
            {
                "row": CY_row,
                "col": CY_col,
                "ok": is_okCB,
                "comment": (None if is_okCB else commentCB),
            }
        )
        if not is_okCB:
            issues.append(commentCB)

        commentOB = (
            f"BSPL = {BSPL_PY_bal:,.2f}, Note = {PY_bal:,.2f}, Sai lệch = {diffOB:,.0f}"
        )
        cross_ref_marks.append(
            {
                "row": CY_row - gap_row,
                "col": CY_col - gap_col,
                "ok": is_okOB,
                "comment": (None if is_okOB else commentOB),
            }
        )
        if not is_okOB:
            issues.append(commentOB)


def check_table_total(df, heading=None):
    """Kiểm tra dòng tổng của bảng có khớp với chi tiết không và trả về trạng thái chi tiết."""

    df_numeric = df.map(normalize_numeric_column)
    total_row_idx = find_total_row_index(df)
    last_col_idx = len(df.columns) - 1
    heading_lower = heading.lower().strip() if heading else ""
    check_column_total = heading_lower in [
        name.lower() for name in TABLES_NEED_COLUMN_CHECK
    ]
    check_separately_total = heading_lower in [
        name.lower() for name in TABLES_NEED_CHECK_SEPARATELY
    ]
    check_tables_without_total = heading_lower in [
        name.lower() for name in TABLES_WITHOUT_TOTAL
    ]

    if total_row_idx is None and not check_column_total:
        return {
            "status": "INFO: Bảng không có dòng/cột tổng",
            "marks": [],
            "cross_ref_marks": [],
        }

    check_tables_without_figures = False
    subset = df_numeric.iloc[2:total_row_idx]
    if subset.isna().all().all():
        check_tables_without_figures = True

    try:
        if (
            check_tables_without_total or check_tables_without_figures
        ) and not check_column_total:
            cross_ref_marks = []
            issues = []
            if heading_lower not in CROSS_CHECK_TABLES_FORM_3 and (
                check_tables_without_figures and not check_tables_without_total
            ):
                if (
                    not any(term in heading_lower for term in RE_PARTY_TABLE)
                    and heading_lower not in BSPL_cross_check_mark
                ):
                    account_name = heading_lower
                    CY_bal = df_numeric.iloc[total_row_idx, len(df.columns) - 2] or 0
                    PY_bal = df_numeric.iloc[total_row_idx, len(df.columns) - 1] or 0
                    cross_check_with_BSPL(
                        df,
                        cross_ref_marks,
                        issues,
                        account_name,
                        CY_bal,
                        PY_bal,
                        total_row_idx,
                        len(df.columns) - 2,
                        0,
                        -1,
                    )
                    BSPL_cross_check_mark.append(account_name)

            return {
                "status": "INFO: Bảng không bao gồm số/số tổng",
                "marks": [],
                "cross_ref_marks": cross_ref_marks,
            }

        else:
            if heading_lower == "balance sheet":
                return _check_balance_sheet(df)
            elif heading_lower == "statement of income":
                return _check_income_statement(df)
            elif heading_lower == "statement of cash flows":
                return _check_cash_flows(df)
            elif heading_lower == "changes in owners’ equity":
                return _check_changes_in_equity(df)
            elif "reconciliation of effective tax rate" in heading_lower:
                return _check_income_tax_rec_table(df)
            elif (
                "recognised in the statement of income" in heading_lower
                or "recognised in the balance sheet" in heading_lower
                or "recognised in consolidated statement of income" in heading_lower
                or "deferred tax assets and liabilities" in heading_lower
                or "deferred tax assets" in heading_lower
                or "deferred tax liabilities" in heading_lower
                or "recognised in consolidated balance sheet" in heading_lower
            ):
                return _check_income_tax_remaining_tables(df)
            else:
                # -----------------------------------------------------------------------
                # Cast dòng tổng trong bảng
                # -----------------------------------------------------------------------
                if check_separately_total:
                    # -----------------------------------------------------------------------
                    # #TABLES_NEED_CHECK_SEPARATELY included fixed asets disclosure
                    # -----------------------------------------------------------------------
                    # Tìm dòng chứa từ khóa tổng
                    cost_keywords = ["cost", "giá vốn"]
                    AD_keywords = [
                        "accumulated depreciation",
                        "accumulated amortisation",
                        "khấu hao lũy kế",
                        "hao mòn lũy kế",
                    ]
                    NBV_keywords = ["net book value", "giá trị còn lại"]

                    for i, row in df.iterrows():
                        if any(
                            keyword.lower() in str(cell).lower()
                            for cell in row
                            for keyword in cost_keywords
                        ):
                            cost_start_row_idx = i
                        if any(
                            keyword.lower() in str(cell).lower()
                            for cell in row
                            for keyword in AD_keywords
                        ):
                            AD_start_row_idx = i
                        if any(
                            keyword.lower() in str(cell).lower()
                            for cell in row
                            for keyword in NBV_keywords
                        ):
                            NBV_start_row_idx = i

                    cost_detail_sum = df_numeric.iloc[
                        cost_start_row_idx: AD_start_row_idx - 2
                    ].sum(skipna=True)
                    cost_total_row = df_numeric.iloc[AD_start_row_idx - 1]
                    AD_detail_sum = df_numeric.iloc[
                        AD_start_row_idx: NBV_start_row_idx - 2
                    ].sum(skipna=True)
                    AD_total_row = df_numeric.iloc[NBV_start_row_idx - 1]
                    OB_detail_cal = (
                        df_numeric.iloc[cost_start_row_idx + 1]
                        - df_numeric.iloc[AD_start_row_idx + 1]
                    )
                    CB_detail_cal = cost_total_row - AD_total_row
                    OB_NBV_total_row = df_numeric.iloc[NBV_start_row_idx + 1]
                    CB_NBV_total_row = df_numeric.iloc[NBV_start_row_idx + 2]

                    issues = []
                    marks = []
                    cross_ref_marks = []

                    for col in range(len(df.columns)):
                        if not pd.isna(cost_total_row[col]) and not pd.isna(
                            cost_detail_sum[col]
                        ):
                            diff = cost_detail_sum[col] - cost_total_row[col]
                            is_ok = abs(round(diff)) == 0
                            comment = f"DÒNG TỔNG (GV) - Cột {
                                col +
                                1}: Tính lại={
                                cost_detail_sum[col]:,.2f}, Trên bảng={
                                cost_total_row[col]:,.2f}, Sai lệch={
                                diff:,.2f}"
                            marks.append(
                                {
                                    "row": AD_start_row_idx - 1,
                                    "col": col,
                                    "ok": is_ok,
                                    "comment": (None if is_ok else comment),
                                }
                            )
                            if not is_ok:
                                issues.append(comment)

                    for col in range(len(df.columns)):
                        if not pd.isna(AD_total_row[col]) and not pd.isna(
                            AD_detail_sum[col]
                        ):
                            diff = AD_detail_sum[col] - AD_total_row[col]
                            is_ok = abs(round(diff)) == 0
                            comment = f"DÒNG TỔNG (AD) - Cột {
                                col +
                                1}: Tính lại={
                                AD_detail_sum[col]:,.2f}, Trên bảng={
                                AD_total_row[col]:,.2f}, Sai lệch={
                                diff:,.2f}"
                            marks.append(
                                {
                                    "row": NBV_start_row_idx - 1,
                                    "col": col,
                                    "ok": is_ok,
                                    "comment": (None if is_ok else comment),
                                }
                            )
                            if not is_ok:
                                issues.append(comment)

                    for col in range(len(df.columns)):
                        if (
                            not pd.isna(OB_NBV_total_row[col])
                            and not pd.isna(CB_NBV_total_row[col])
                            and not pd.isna(OB_detail_cal[col])
                            and not pd.isna(CB_detail_cal[col])
                        ):
                            diffOB = OB_detail_cal[col] - OB_NBV_total_row[col]
                            diffCB = CB_detail_cal[col] - CB_NBV_total_row[col]
                            is_okOB = abs(round(diffOB)) == 0
                            is_okCB = abs(round(diffCB)) == 0
                            commentOB = f"DÒNG TỔNG (OB NBV) - Cột {
                                col +
                                1}: Tính lại={
                                OB_detail_cal[col]:,.2f}, Trên bảng={
                                OB_NBV_total_row[col]:,.2f}, Sai lệch={
                                diff:,.2f}"
                            marks.append(
                                {
                                    "row": NBV_start_row_idx + 1,
                                    "col": col,
                                    "ok": is_okOB,
                                    "comment": (None if is_okOB else commentOB),
                                }
                            )
                            if not is_okOB:
                                issues.append(commentOB)

                            commentCB = f"DÒNG TỔNG (CB NBV) - Cột {
                                col +
                                1}: Tính lại={
                                CB_detail_cal[col]:,.2f}, Trên bảng={
                                CB_NBV_total_row[col]:,.2f}, Sai lệch={
                                diff:,.2f}"
                            marks.append(
                                {
                                    "row": NBV_start_row_idx + 2,
                                    "col": col,
                                    "ok": is_okCB,
                                    "comment": (None if is_okCB else commentCB),
                                }
                            )
                            if not is_okCB:
                                issues.append(commentCB)

                    # NBV
                    account_name = heading_lower
                    CY_bal = (
                        df_numeric.iloc[NBV_start_row_idx + 2, len(df.columns) - 1] or 0
                    )
                    PY_bal = (
                        df_numeric.iloc[NBV_start_row_idx + 1, len(df.columns) - 1] or 0
                    )
                    if account_name not in BSPL_cross_check_mark:
                        cross_check_with_BSPL(
                            df,
                            cross_ref_marks,
                            issues,
                            account_name,
                            CY_bal,
                            PY_bal,
                            NBV_start_row_idx + 2,
                            len(df.columns) - 1,
                            1,
                            0,
                        )
                        BSPL_cross_check_mark.append(account_name)

                    # Costs
                    CY_bal = (
                        df_numeric.iloc[AD_start_row_idx - 1, len(df.columns) - 1] or 0
                    )
                    PY_bal = (
                        df_numeric.iloc[cost_start_row_idx + 1, len(df.columns) - 1]
                        or 0
                    )

                    if heading_lower == "tangible fixed assets":
                        account_name = "222"
                    elif heading_lower == "finance lease tangible fixed assets":
                        account_name = "225"
                    elif heading_lower == "intangible fixed assets":
                        account_name = "228"
                    elif heading_lower == "investment property":
                        account_name = "231"
                    if account_name not in BSPL_cross_check_mark:
                        cross_check_with_BSPL(
                            df,
                            cross_ref_marks,
                            issues,
                            account_name,
                            CY_bal,
                            PY_bal,
                            AD_start_row_idx - 1,
                            len(df.columns) - 1,
                            AD_start_row_idx - 1 - (cost_start_row_idx + 1),
                            0,
                        )
                        BSPL_cross_check_mark.append(account_name)

                    # Accumulated depreciation
                    CY_bal = (
                        df_numeric.iloc[NBV_start_row_idx - 1, len(df.columns) - 1]
                    ) * -1 or 0
                    PY_bal = (
                        df_numeric.iloc[AD_start_row_idx + 1, len(df.columns) - 1]
                    ) * -1 or 0

                    if heading_lower == "tangible fixed assets":
                        account_name = "223"
                    elif heading_lower == "finance lease tangible fixed assets":
                        account_name = "226"
                    elif heading_lower == "intangible fixed assets":
                        account_name = "229"
                    elif heading_lower == "investment property":
                        account_name = "232"
                    if account_name not in BSPL_cross_check_mark:
                        cross_check_with_BSPL(
                            df,
                            cross_ref_marks,
                            issues,
                            account_name,
                            CY_bal,
                            PY_bal,
                            NBV_start_row_idx - 1,
                            len(df.columns) - 1,
                            NBV_start_row_idx - 1 - (AD_start_row_idx + 1),
                            0,
                        )
                        BSPL_cross_check_mark.append(account_name)

                else:
                    issues = []
                    marks = []
                    cross_ref_marks = []
                    if find_total_row_index is not None:

                        def find_block_sum(start_idx):
                            sum_vals = [0.0] * len(df.columns)
                            count = 0
                            i = start_idx + 1
                            while i < len(df):
                                row = df.iloc[i]
                                if all(str(cell).strip() == "" for cell in row):
                                    break
                                for col in range(len(df.columns)):
                                    val = df_numeric.iloc[i, col]
                                    if not pd.isna(val):
                                        sum_vals[col] += val
                                count += 1
                                i += 1
                            return sum_vals, count, i

                        def compare_sum_with_total(sum_vals, total_row, end_row):
                            for col in range(len(df.columns)):
                                total_val = total_row[col]
                                if not pd.isna(total_val):
                                    diff = sum_vals[col] - total_val
                                    is_ok = abs(round(diff)) == 0
                                    comment = f"Cột {
                                        col +
                                        1}: Tổng chi tiết = {
                                        sum_vals[col]:,.2f}, Tổng trên bảng = {
                                        total_val:,.2f}, Sai lệch = {
                                        diff:,.0f}"
                                    marks.append(
                                        {
                                            "row": end_row + 1,
                                            "col": col,
                                            "ok": is_ok,
                                            "comment": (None if is_ok else comment),
                                        }
                                    )
                                    if not is_ok:
                                        issues.append(comment)

                        start_idx = 0
                        while start_idx < len(df):
                            row = df.iloc[start_idx]
                            row_text = " ".join(str(x).lower() for x in row)
                            if (
                                all(str(cell).strip() == "" for cell in row)
                                or "equity investments" in row_text
                            ):
                                break
                            start_idx += 1

                        total1 = [0.0] * len(df.columns)
                        sum1, count1, end1 = find_block_sum(start_idx)
                        if count1 > 1 and end1 < len(df) - 1:
                            total1_row = df_numeric.iloc[end1 + 1]
                            compare_sum_with_total(sum1, total1_row, end1)
                        total1 = sum1

                        if count1 == 1:
                            start_idx = end1 - 1
                        else:
                            start_idx = end1 + 1

                        if total_row_idx > start_idx + 1:
                            total2 = [0.0] * len(df.columns)
                            sum2, count2, end2 = find_block_sum(start_idx)
                            if count2 > 1 and end2 < len(df) - 1:
                                total2_row = df_numeric.iloc[end2 + 1]
                                compare_sum_with_total(sum2, total2_row, end2)
                            if "revenue from" in heading_lower:
                                if total2_row.dropna().gt(0).all():
                                    for col in range(len(df.columns)):
                                        total2[col] = -sum2[col]
                            else:
                                total2 = sum2
                            final_row = df_numeric.iloc[len(df) - 1]

                            for col in range(len(df.columns)):
                                combined = total1[col] + total2[col]
                                final_val = final_row[col]
                                if not pd.isna(final_val):
                                    diff = combined - final_val
                                    is_ok = abs(round(diff)) == 0
                                    comment = f"Dòng Grand total - Cột {
                                        col +
                                        1}: Tổng cộng = {
                                        combined:,.2f}, Dòng cuối = {
                                        final_val:,.2f}, Sai lệch = {
                                        diff:,.0f}"
                                    marks.append(
                                        {
                                            "row": len(df) - 1,
                                            "col": col,
                                            "ok": is_ok,
                                            "comment": (None if is_ok else comment),
                                        }
                                    )
                                    if not is_ok:
                                        issues.append(comment)

                            if heading_lower in CROSS_CHECK_TABLES_FORM_1:
                                account_name = heading_lower
                                CY_bal = (
                                    df_numeric.iloc[len(df) - 1, len(df.columns) - 2]
                                    or 0
                                )
                                PY_bal = (
                                    df_numeric.iloc[len(df) - 1, len(df.columns) - 1]
                                    or 0
                                )
                                if (
                                    "accounts receivable from customers"
                                    in heading_lower
                                ):
                                    account_name = "accounts receivable from customers"
                                if account_name not in BSPL_cross_check_mark:
                                    cross_check_with_BSPL(
                                        df,
                                        cross_ref_marks,
                                        issues,
                                        account_name,
                                        CY_bal,
                                        PY_bal,
                                        len(df) - 1,
                                        len(df.columns) - 2,
                                        0,
                                        -1,
                                    )
                                    BSPL_cross_check_mark.append(account_name)

                            elif heading_lower in CROSS_CHECK_TABLES_FORM_2:
                                account_name = heading_lower
                                CY_bal = (
                                    df_numeric.iloc[end1 + 1, len(df.columns) - 2] or 0
                                )
                                PY_bal = (
                                    df_numeric.iloc[end1 + 1, len(df.columns) - 1] or 0
                                )
                                if account_name not in BSPL_cross_check_mark:
                                    cross_check_with_BSPL(
                                        df,
                                        cross_ref_marks,
                                        issues,
                                        account_name,
                                        CY_bal,
                                        PY_bal,
                                        end1 + 1,
                                        len(df.columns) - 2,
                                        0,
                                        -1,
                                    )
                                    BSPL_cross_check_mark.append(account_name)

                                account_name = "revenue deductions"
                                CY_bal = (
                                    df_numeric.iloc[end2 + 1, len(df.columns) - 2] or 0
                                )
                                PY_bal = (
                                    df_numeric.iloc[end2 + 1, len(df.columns) - 1] or 0
                                )
                                if account_name not in BSPL_cross_check_mark:
                                    cross_check_with_BSPL(
                                        df,
                                        cross_ref_marks,
                                        issues,
                                        account_name,
                                        CY_bal,
                                        PY_bal,
                                        end2 + 1,
                                        len(df.columns) - 2,
                                        0,
                                        -1,
                                    )
                                    BSPL_cross_check_mark.append(account_name)

                                account_name = "net revenue (10 = 01 - 02)"
                                CY_bal = (
                                    df_numeric.iloc[len(df) - 1, len(df.columns) - 2]
                                    or 0
                                )
                                PY_bal = (
                                    df_numeric.iloc[len(df) - 1, len(df.columns) - 1]
                                    or 0
                                )
                                if account_name not in BSPL_cross_check_mark:
                                    cross_check_with_BSPL(
                                        df,
                                        cross_ref_marks,
                                        issues,
                                        account_name,
                                        CY_bal,
                                        PY_bal,
                                        len(df) - 1,
                                        len(df.columns) - 2,
                                        0,
                                        -1,
                                    )
                                    BSPL_cross_check_mark.append(account_name)

                        else:
                            if heading_lower not in CROSS_CHECK_TABLES_FORM_3:
                                account_name = heading_lower
                                if (
                                    "accounts receivable from customers"
                                    in heading_lower
                                ):
                                    account_name = "accounts receivable from customers"
                                if account_name not in BSPL_cross_check_mark:
                                    CY_bal = (
                                        df_numeric.iloc[end1 + 1, len(df.columns) - 2]
                                        or 0
                                    )
                                    PY_bal = (
                                        df_numeric.iloc[end1 + 1, len(df.columns) - 1]
                                        or 0
                                    )
                                    cross_check_with_BSPL(
                                        df,
                                        cross_ref_marks,
                                        issues,
                                        account_name,
                                        CY_bal,
                                        PY_bal,
                                        end1 + 1,
                                        len(df.columns) - 2,
                                        0,
                                        -1,
                                    )
                                    BSPL_cross_check_mark.append(account_name)

                            else:

                                def search_col_and_cross_ref(
                                    key_word, account_name, total_row_xref
                                ):
                                    CY_col = 0
                                    PY_col = 0
                                    for j, row in df.iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if key_word in row_text:
                                            for col in range(len(df.columns)):
                                                if key_word in row.get(col, "").lower():
                                                    if CY_col == 0:
                                                        CY_col = col
                                                        PY_col = CY_col
                                                    else:
                                                        PY_col = col
                                                        break
                                            break

                                    if (
                                        "short-term borrowings" in account_name
                                        or "short-term bonds" in account_name
                                        or "short-term borrowing" in account_name
                                        or "short-term bond" in account_name
                                    ):
                                        temp_col = PY_col
                                        PY_col = CY_col
                                        CY_col = temp_col

                                    CY_bal = (
                                        0
                                        if pd.isna(
                                            df_numeric.iloc[total_row_xref, CY_col]
                                        )
                                        else df_numeric.iloc[total_row_xref, CY_col]
                                    )
                                    PY_bal = (
                                        0
                                        if pd.isna(
                                            df_numeric.iloc[total_row_xref, PY_col]
                                        )
                                        else df_numeric.iloc[total_row_xref, PY_col]
                                    )
                                    if CY_col != 0 or PY_col != 0:
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            total_row_xref,
                                            CY_col,
                                            0,
                                            CY_col - PY_col,
                                        )
                                    BSPL_cross_check_mark.append(account_name)

                                def search_row_and_cross_ref(account_name, col_xref):
                                    CY_row = end1 + 1
                                    PY_row = 0
                                    for j, row in df.iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if "opening balance" in row_text:
                                            PY_row = j
                                            break

                                    CY_bal = (
                                        0
                                        if pd.isna(df_numeric.iloc[CY_row, col_xref])
                                        else df_numeric.iloc[CY_row, col_xref]
                                    )
                                    PY_bal = (
                                        0
                                        if pd.isna(df_numeric.iloc[PY_row, col_xref])
                                        else df_numeric.iloc[PY_row, col_xref]
                                    )
                                    if CY_row != 0 or PY_row != 0:
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            CY_row,
                                            col_xref,
                                            CY_row - PY_row,
                                            0,
                                        )
                                    BSPL_cross_check_mark.append(account_name)

                                if heading_lower == "bad and doubtful debts":
                                    account_name = "allowance for doubtful debts"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, end1 + 1
                                        )

                                elif heading_lower == "inventories":
                                    # search and cross ref costs
                                    account_name = "141"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, end1 + 1
                                        )
                                    # search and cross ref allowance
                                    account_name = "149"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, end1 + 1
                                        )

                                elif (
                                    heading_lower
                                    == "equity investments in other entity"
                                    or heading_lower
                                    == "equity investments in other entities"
                                ):
                                    # search and cross ref costs (separate) or carrying
                                    # amounts (consol)
                                    account_name = "investments in other entities"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, end1 + 1
                                        )
                                        search_col_and_cross_ref(
                                            "carrying amounts", account_name, end1 + 1
                                        )
                                    # search and cross ref allowance
                                    account_name = "254"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance for diminution in value",
                                            account_name,
                                            end1 + 1,
                                        )

                                elif heading_lower == "construction in progress":
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_row_and_cross_ref(account_name, 1)

                                elif heading_lower == "long-term prepaid expenses":
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_row_and_cross_ref(
                                            account_name, len(df.columns) - 1
                                        )

                                elif "accounts payable to suppliers" in heading_lower:
                                    # search and cross ref costs
                                    account_name = "accounts payable to suppliers"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, end1 + 1
                                        )

                                elif "taxes" in heading_lower:
                                    # search and cross ref costs
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    end1 + 1, len(df.columns) - 1
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                end1 + 1, len(df.columns) - 1
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(df_numeric.iloc[end1 + 1, 1])
                                            else df_numeric.iloc[end1 + 1, 1]
                                        )
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            end1 + 1,
                                            len(df.columns) - 1,
                                            0,
                                            (len(df.columns) - 1) - 1,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                                elif (
                                    "short-term borrowings" in heading_lower
                                    or "short-term bonds" in heading_lower
                                    or "short-term borrowing" in heading_lower
                                    or "short-term bond" in heading_lower
                                ):
                                    # search and cross ref costs
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "carrying", account_name, start_idx
                                        )

                # -----------------------------------------------------------------------
                # Cast cột tổng trong bảng, Chỉ kiểm tra cột tổng nếu heading nằm trong danh sách
                # -----------------------------------------------------------------------
                if check_column_total and last_col_idx > 1:

                    def is_borrowing_movement(df) -> bool:
                        borrowing_movement = False
                        for i, row in df.iterrows():
                            row_text = " ".join(str(cell).lower() for cell in row)
                            if "amount within repayment capacity" in row_text:
                                borrowing_movement = True
                                break
                        return borrowing_movement

                    def compare_sum_with_total(start_col, end_col):
                        for i in range(total_row_idx + 1):  # Duyệt từng dòng chi tiết
                            row = df_numeric.iloc[i]
                            # row_sum = row.drop(labels=df.columns[last_col_idx]).sum(skipna=True)
                            row_sum = row[df.columns[start_col:end_col]].sum(
                                skipna=True
                            )
                            col_total_val = row.iloc[end_col]
                            if not pd.isna(col_total_val) and not pd.isna(row_sum):
                                diff = row_sum - col_total_val
                                is_ok = abs(round(diff)) == 0
                                comment = f"CỘT TỔNG - Dòng {
                                    i +
                                    1}: Tính lại={
                                    row_sum:,.2f}, Trên bảng={
                                    col_total_val:,.2f}, Sai lệch={
                                    diff:,.2f}"
                                marks.append(
                                    {
                                        "row": i,
                                        "col": end_col,
                                        "ok": is_ok,
                                        "comment": (None if is_ok else comment),
                                    }
                                )
                                if not is_ok:
                                    issues.append(comment)

                    start_cidx = 0
                    end_cidx = last_col_idx
                    if "borrowings" not in heading_lower:
                        compare_sum_with_total(start_cidx, end_cidx)
                    else:
                        if "borrowings" in heading_lower and is_borrowing_movement(df):
                            start_cidx = 2
                            end_cidx = last_col_idx - 1
                            compare_sum_with_total(start_cidx, end_cidx)
                # -----------------------------------------------------------------------
                # Trả kết quả
                # -----------------------------------------------------------------------
                if not issues:
                    status = "PASS: Kiểm tra công thức: KHỚP (0 sai lệch)"
                else:
                    preview = "; ".join(issues[:10])
                    more = f" ... (+{len(issues) -
                                     10} dòng)" if len(issues) > 10 else ""
                    status = f"FAIL: Kiểm tra công thức: {
                        len(issues)} sai lệch. {preview}{more}"

                return {
                    "status": status,
                    "marks": marks,
                    "cross_ref_marks": cross_ref_marks,
                }

    except Exception:
        return {
            "status": "WARN Không xác định được dòng tổng hoặc lỗi khi kiểm tra",
            "marks": [],
            "cross_ref_marks": [],
        }


# =====================
# Excel writing routines
# =====================
def apply_status_colors(ws):
    """Tô màu trạng thái kiểm tra."""
    for row in ws.iter_rows(min_row=2, min_col=2, max_col=2):
        for cell in row:
            if "PASS" in cell.value:
                cell.fill = PatternFill(start_color="C6EFCE", fill_type="solid")
            elif "FAIL" in cell.value:
                cell.fill = PatternFill(start_color="FFC7CE", fill_type="solid")
            elif "INFO" in cell.value:
                cell.fill = PatternFill(start_color="87CEEB", fill_type="solid")
            else:
                cell.fill = PatternFill(start_color="FFEB9C", fill_type="solid")


def write_summary_sheet_option_1(ws, results, sheet_names, wb):
    """Ghi sheet tổng hợp trạng thái kiểm tra và tạo hyperlink đến từng bảng."""
    ws.title = "Tổng hợp kiểm tra"
    ws.append(["Tên bảng", "Trạng thái kiểm tra"])
    for i, (result, sheet_name) in enumerate(zip(results, sheet_names)):
        cell = ws.cell(row=i + 2, column=1, value=sheet_name)
        if sheet_name in wb.sheetnames:
            quoted_name = quote_sheetname(sheet_name)
            cell.hyperlink = f"#{quoted_name}!A1"
            cell.style = "Hyperlink"
        #    ws.cell(row=i+2, column=2, value=status)
        ws.cell(row=i + 2, column=2, value=result.get("status"))
    apply_status_colors(ws)


def write_summary_sheet_option_2(ws, results, sheet_positions, wb):
    """Ghi sheet tổng hợp trạng thái kiểm tra và tạo hyperlink đến từng bảng trong sheet 'FS casting'."""
    ws.title = "Tổng hợp kiểm tra"
    ws.append(["Tên bảng", "Trạng thái kiểm tra"])

    for i, (result, (heading, start_row)) in enumerate(zip(results, sheet_positions)):
        cell = ws.cell(row=i + 2, column=1, value=heading)
        cell.hyperlink = f"#'FS casting'!A{start_row}"
        cell.style = "Hyperlink"
        ws.cell(row=i + 2, column=2, value=result.get("status"))

    apply_status_colors(ws)


def write_table_sheets_option_1(wb, table_heading_pairs, results):
    """Ghi từng bảng vào sheet riêng biệt, dùng heading nếu có, thêm trạng thái và nút quay lại."""
    sheet_names = []
    for i, ((table, heading), result) in enumerate(zip(table_heading_pairs, results)):
        table.columns = table.columns.map(str)
        raw_name = heading if heading else f"Bảng {i + 1}"
        sheet_name = shorten_sheet_name(raw_name)
        if sheet_name in wb.sheetnames:
            sheet_name += f"_{i + 1}"
        sheet_names.append(sheet_name)

        ws = wb.create_sheet(title=sheet_name)
        for row in dataframe_to_rows(table, index=False, header=True):
            ws.append(row)

        # Áp dụng highlight & comment theo marks
        apply_cell_marks(ws, result.get("marks", []))
        # Tạo Table (để stripe dễ nhìn) sau khi đã tô highlight (không ảnh hưởng
        # fill ô đã set)
        max_col = get_column_letter(ws.max_column)
        max_row = ws.max_row
        tab = Table(displayName=f"Table{i + 1}", ref=f"A1:{max_col}{max_row}")
        tab.tableStyleInfo = TableStyleInfo(
            name="TableStyleMedium9", showRowStripes=True
        )
        ws.add_table(tab)

        # Ghi trạng thái kiểm tra sau bảng
        ws.cell(row=max_row + 2, column=1, value="Trạng thái kiểm tra:")
        ws.cell(row=max_row + 2, column=2, value=result.get("status"))

        # Thêm hyperlink quay lại sheet tổng hợp
        back_cell = ws.cell(row=max_row + 3, column=1, value="⬅ Quay lại Tổng hợp")
        back_cell.hyperlink = "#'Tổng hợp kiểm tra'!A1"
        back_cell.style = "Hyperlink"
    return sheet_names


def write_table_sheets_option_2(wb, table_heading_pairs, results):
    """Ghi tất cả các bảng vào một sheet duy nhất 'FS casting', mỗi bảng cách nhau vài dòng."""
    sheet_name = "FS casting"
    ws = wb.create_sheet(title=sheet_name)
    sheet_positions = []  # lưu vị trí bắt đầu của từng bảng
    current_row = 1

    for i, ((table, heading), result) in enumerate(zip(table_heading_pairs, results)):
        table.columns = table.columns.map(str)
        raw_name = heading if heading else f"Bảng {i + 1}"
        sheet_positions.append((raw_name, current_row))  # lưu heading và vị trí dòng

        # Ghi heading
        ws.cell(row=current_row, column=1, value=raw_name)
        current_row += 1

        # Ghi bảng
        start_row = current_row
        start_col = 1

        for row_idx, row in enumerate(
            dataframe_to_rows(table, index=False, header=True)
        ):
            for j, value in enumerate(row):
                # Sanitize value to prevent Excel formula injection
                safe_value = sanitize_excel_value(value)

                # Kiểm tra nếu value là số (hoặc chuỗi có thể chuyển thành số)
                if row_idx == 0:
                    ws.cell(row=current_row, column=start_col + j, value=safe_value)
                else:
                    if isinstance(safe_value, (int, float)):
                        cell = ws.cell(
                            row=current_row, column=start_col + j, value=safe_value
                        )
                        cell.number_format = '_(* #,##0_);_(* (#,##0);_(* "-"??_);_(@_)'
                    else:
                        # Thử chuyển đổi chuỗi thành số
                        try:
                            num_val = float(
                                str(safe_value)
                                .replace(",", "")
                                .replace("(", "-")
                                .replace(")", "")
                            )
                            cell = ws.cell(
                                row=current_row, column=start_col + j, value=num_val
                            )
                            cell.number_format = (
                                '_(* #,##0_);_(* (#,##0);_(* "-"??_);_(@_)'
                            )
                        except ValueError:
                            ws.cell(
                                row=current_row, column=start_col + j, value=safe_value
                            )
            current_row += 1

        end_row = current_row - 1
        end_col = start_col + len(table.columns) - 1

        # Tạo định dạng bảng Excel
        table_range = f"{
            ws.cell(
                row=start_row,
                column=start_col).coordinate}:{
            ws.cell(
                row=end_row,
                column=end_col).coordinate}"
        excel_table = Table(displayName=f"Table_{i + 1}", ref=table_range)
        style = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        excel_table.tableStyleInfo = style
        ws.add_table(excel_table)

        # Áp dụng highlight & comment
        apply_cell_marks(ws, result.get("marks", []), start_row - 1, start_col - 1)
        apply_crossref_marks(
            ws, result.get("cross_ref_marks", []), start_row, start_col - 1
        )

        # Ghi trạng thái kiểm tra
        ws.cell(row=current_row + 1, column=1, value="Trạng thái kiểm tra:")
        ws.cell(row=current_row + 1, column=2, value=result.get("status"))
        back_cell = ws.cell(row=current_row + 2, column=1, value="⬅ Quay lại Tổng hợp")
        back_cell.hyperlink = "#'Tổng hợp kiểm tra'!A1"
        back_cell.style = "Hyperlink"

        current_row += 4  # cách vài dòng trước khi ghi bảng tiếp theo

    return sheet_positions


# =========================
# Thêm hàm hỗ trợ cross-check
# =========================
"""def apply_cross_check(ws, start_row, start_col, table, heading, cache):

    max_col = len(table.columns)
    for i, row in enumerate(table.itertuples(index=False), start=start_row + 1):
        row_text = " ".join(str(x).lower() for x in row)
        for tbl_name, (cur_val, prior_val) in cache.items():
            if tbl_name.lower() in row_text and not any(x in row_text for x in ['150', '260']):
                # Ghi Current Year và Prior Year
                ws.cell(row=i, column=start_col + max_col + 1, value=cur_val)
                ws.cell(row=i, column=start_col + max_col + 2, value=prior_val)

                # Lấy giá trị gốc để tính diff
                orig_cur = normalize_numeric_column(ws.cell(row=i, column=start_col + max_col - 2).value)
                orig_prior = normalize_numeric_column(ws.cell(row=i, column=start_col + max_col-1).value)

                # Diff cho Current Year
                diff_cur = (orig_cur or 0) - cur_val
                cell_diff_cur = ws.cell(row=i, column=start_col + max_col + 3, value=diff_cur)
                cell_diff_cur.fill = GREEN_FILL if diff_cur == 0 else RED_FILL

                # Diff cho Prior Year
                diff_prior = (orig_prior or 0) - prior_val
                cell_diff_prior = ws.cell(row=i, column=start_col + max_col + 4, value=diff_prior)
                cell_diff_prior.fill = GREEN_FILL if diff_prior == 0 else RED_FILL"""


def export_check_result(word_path, excel_path):
    """Quy trình chính: đọc Word, kiểm tra, ghi Excel."""

    table_heading_pairs = read_word_tables_with_headings(word_path)
    results = [
        check_table_total(table, heading) for table, heading in table_heading_pairs
    ]

    wb = Workbook()
    summary_ws = wb.active
    # Option 1: mỗi bảng 1 sheet
    # sheet_names = write_table_sheets_option_1(wb, table_heading_pairs, results)
    # write_summary_sheet_option_1(summary_ws, results, sheet_names, wb)

    # Option 2: tất cả bảng trong 1 sheet
    sheet_positions = write_table_sheets_option_2(wb, table_heading_pairs, results)
    write_summary_sheet_option_2(summary_ws, results, sheet_positions, wb)

    wb.save(excel_path)
    # Safe file opening delegated to new modular system


# Main execution - now delegates to the new modular system
if __name__ == "__main__":
    # #region agent log
    import json

    debug_log_path = r"c:\Users\Admin\Downloads\Quality Audit\.cursor\debug.log"

    def debug_log(hypothesis_id, message, data=None):
        log_entry = {
            "id": f"log_{int(__import__('time').time() * 1000)}",
            "timestamp": int(__import__("time").time() * 1000),
            "location": "Quality Audit.py:1753",
            "message": message,
            "data": data or {},
            "sessionId": "debug-session",
            "runId": "run1",
            "hypothesisId": hypothesis_id,
        }
        try:
            with open(debug_log_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(log_entry) + "\n")
        except BaseException:
            pass  # Fail silently if logging fails

    # #endregion

    debug_log("E", "__main__ block entered", {"name": __name__})

    try:
        # Import the new main function
        from main import main

        debug_log("A", "main import successful")
    except Exception as e:
        debug_log(
            "A", "main import failed", {"error": str(e), "error_type": type(e).__name__}
        )
        raise

    # Execute with new secure, modular system
    try:
        debug_log("B", "about to call main()", {"sys_argv": sys.argv})
        main()
        debug_log("B", "main() call completed successfully")
    except Exception as e:
        debug_log(
            "B", "main() call failed", {"error": str(e), "error_type": type(e).__name__}
        )
        raise
