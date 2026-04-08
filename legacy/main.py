import os
import re
import tkinter as tk
from tkinter import filedialog

import docx
import numpy as np
import pandas as pd
from openpyxl import Workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter, quote_sheetname
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.worksheet.table import Table, TableStyleInfo

# =========================
# Constants & Regex helpers
# =========================
_SHEET_NAME_CLEAN_RE = re.compile(r"[:\\/*?\[\]]")
# cross_check_cache = {}
BSPL_cross_check_cache = {}
BSPL_cross_check_mark = []
_CODE_COL_NAME = "code"
_NOTE_COL_NAME = "note"
_CODE_VALID_RE = re.compile(r"^[0-9]+[A-Z]?$")
_HEADER_DATE_RE = re.compile(r"\d{4}|\d{1,2}/\d{1,2}/\d{2,4}")

TABLES_WITHOUT_TOTAL = {
    "geographical segments",
    "non-cash investing activity",
    "non-cash investing activities",
    "non-cash investing and financing activities",
    "significant transactions with related parties",
    "significant transactions with related companies",
    "fees paid and payable to the auditors",
    "corresponding figures",
    # tạm add một vài bảng chưa được xử lý ở version này
    "interest rate risk",
    "fair values versus carrying amounts",
}

TABLES_NEED_CHECK_SEPARATELY = {
    "tangible fixed assets",
    "finance lease tangible fixed assets",
    "intangible fixed assets",
    "mature livestock producing periodic products",
    "investment property held to earn rental",
    "investment property held for capital appreciation",
    "goodwill",
}

TABLES_NEED_COLUMN_CHECK = {
    #'business segments',
    "acquisition of subsidiary",
    "tangible fixed assets",
    "finance lease tangible fixed assets",
    "intangible fixed assets",
    "mature livestock producing periodic products",
    "investment property held to earn rental",
    "investment property held for capital appreciation",
    "long-term deferred expenses",
    "movement in temporary differences during the period",
    "movement in temporary differences during the year",
    "taxes and others payable to state treasury",
    "taxes payable to state treasury",
    "taxes and others receivable from state treasury",
    "taxes receivable from state treasury",
    "taxes and others receivable from the state",
    "taxes receivable from the state",
    "taxes and others payable to the state",
    "taxes payable to the state",
    "others payable to the state",
    "taxes and others payable to the state – long-term",
    "taxes payable to the state – long-term",
    "others payable to the state – long-term",
    "borrowings",
    "borrowings, bonds and finance lease liabilities",
    "short-term borrowings, bonds and finance lease liabilities",
    "long-term borrowings, bonds and financial lease liabilities",
    "short-term borrowings",
    "long-term borrowings",
    "finance lease liabilities",
    "short-term provisions",
    "long-term provisions",
    "changes in owners’ equity",
}

# Dạng bảng 1A: bảng có thể có subtotal và dòng cross ref là ở grand total
CROSS_CHECK_TABLES_FORM_1A = {
    "receivables on construction contracts according to stages of completion",
    "payables on construction contracts according to stages of completion",
    "deferred tax assets and liabilities",
    "deferred tax assets",
    "deferred tax liabilities",
    "recognised deferred tax assets and liabilities",
    "accounts payable to suppliers",
    "accounts payable to suppliers detailed by significant suppliers",
    "accounts payable to suppliers detailed by significant supplier",
    "accounts payable to suppliers classified by payment terms",
    "accrued expenses",
    "accrued expenses – short-term",
    "accrued expenses - short-term",
    "accrued expenses – long-term",
    "accrued expenses - long-term",
    "deferred revenue",
    "deferred revenue – short-term",
    "deferred revenue – long-term",
    "long-term deferred revenue",
    "other payables",
    "other payables – short-term",
    "other payables – long-term",
    "long-term borrowings",
    "long-term bonds and financial lease liabilities",
    "long-term financial lease liabilities",
    "long-term bonds",
}

# Dạng bảng 1B: bảng có thể có subtotal và dòng cross ref là ở grand total

CROSS_CHECK_TABLES_FORM_1B = {"cash", "cash and cash equivalents", "cash in banks"}

# Dạng bảng 2: bảng có thể có subtotal, cross ref ở cả subtotal & grand total
CROSS_CHECK_TABLES_FORM_2 = {
    "biological assets – short-term",
    "biological assets – long-term, other than mature livestock producing periodic products",
    "revenue from sales of goods and provision of services",
    "revenue from sales of goods",
    "revenue from provision of services",
}

# Dạng bảng 2: bảng không có subtotal nhưng không phải standard table (standard table là bảng có 3 cột, côt 2 là CY balance, cột 3 là PY balance)
CROSS_CHECK_TABLES_FORM_3 = {
    "acquisition of subsidiary",
    "business segments",
    "investments",
    "trading securities",
    "held-to-maturity investments",
    "equity investments in other entities",
    "equity investments in other entity",
    "accounts receivable from customers",
    "accounts receivable from customers detailed by significant customer",
    "accounts receivable from customers detailed by significant customers",
    "accounts receivable from customers classified by payment term",
    "other short-term receivables comprised:",
    "other long-term receivables comprised:",
    "other receivables",
    "other short-term receivables",
    "other long-term receivables",
    "receivables from BCC contracts under joint control",
    "bad and doubtful debts",
    "shortage of assets awaiting resolution",
    "inventories",
    "long-term work in progress",
    "construction in progress",
    "biological assets – short-term",
    "biological assets – long-term, other than mature livestock producing periodic products",
    "long-term deferred expenses",
    "investment property held for capital appreciation",
    "taxes and others payable to state treasury",
    "taxes and others receivable from state treasury",
    "taxes and others receivable from and payable to state treasury",
    "taxes receivable from state treasury",
    "taxes payable to state treasury",
    "taxes and others receivable from the state",
    "taxes receivable from the state",
    "taxes and others payable to the state",
    "taxes payable to the state",
    "others payable to the state",
    "taxes and others payable to the state – long-term",
    "taxes payable to the state – long-term",
    "others payable to the state – long-term",
    "short-term borrowings",
    "short-term borrowings, bonds and finance lease liabilities",
    "short-term bonds and finance lease liabilities",
    "short-term bonds",
    "convertible bonds",
    "preference shares",
    "provisions",
    "short-term provisions",
    "long-term provisions",
    "changes in owners’ equity",
    "share capital",
    "contributed capital",
}

valid_codes = {"222", "223", "225", "226", "228", "229", "241", "242", "234", "235"}

RE_PARTY_TABLE = {
    "related parties",
    "related party",
    "related companies",
    "related company",
}

# Màu tô
GREEN_FILL = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
BLUE_FILL = PatternFill(start_color="90EE90", end_color="90EE90", fill_type="solid")
RED_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
INFO_FILL = PatternFill(start_color="DAE8FC", end_color="DAE8FC", fill_type="solid")

GREEN_FONT = Font(color="32CD32")  # Xanh
RED_FONT = Font(color="FF0000")  # Đỏ
RIGHT_ALIGN = Alignment(horizontal="right")  # Canh phải


# =========
# Utilities
# =========
def shorten_sheet_name(name, max_length=20):
    """Rút gọn tên sheet để không vượt quá 20 ký tự và loại bỏ ký tự đặc biệt."""
    name = _SHEET_NAME_CLEAN_RE.sub("_", str(name))
    return name[:max_length]


def normalize_numeric_column(value):
    """Chuyển chuỗi số thành số thực, xử lý dấu phẩy và dấu ngoặc."""
    if isinstance(value, str):
        value = value.replace(",", "").replace("(", "-").replace(")", "")
    return pd.to_numeric(value, errors="coerce")


def _dfpos_to_excel(row_idx_0based: int, col_idx_0based: int):
    """
    dataframe_to_rows(..., header=True) -> Excel hàng 1 là header cột.
    Nên: Excel_row = df_row + 2 ; Excel_col = df_col + 1
    """
    return row_idx_0based + 2, col_idx_0based + 1


def is_red_fill(cell):
    fill = cell.fill
    return (
        fill.start_color.rgb == RED_FILL.start_color.rgb
        and fill.end_color.rgb == RED_FILL.end_color.rgb
        and fill.fill_type == RED_FILL.fill_type
    )


def apply_cell_marks(ws, marks: list, start_row, start_col):
    """
    marks: list các dict {'row': int(df_row), 'col': int(df_col), 'ok': bool|None, 'comment': str|None}
    """
    for m in marks:
        r, c = _dfpos_to_excel(m["row"], m["col"])
        cell = ws.cell(row=r + start_row, column=c + start_col)
        if m.get("ok") is True:
            if not is_red_fill(cell):  # chỉ tô xanh nếu chưa bị tô đỏ
                cell.fill = GREEN_FILL
        elif m.get("ok") is False:
            cell.fill = RED_FILL
        # comment cho ô lỗi
        if m.get("comment"):
            try:
                # Tránh đè comment cũ: nối thêm
                if cell.comment:
                    new_text = cell.comment.text + "\n" + str(m["comment"])
                else:
                    new_text = str(m["comment"])

                comment = Comment(text=new_text, author="AutoCheck")
                comment.visible = True
                cell.comment = comment

                # cell.comment = Comment(text=new_text, author="AutoCheck")

            except Exception:
                # comment lỗi thì bỏ qua, không ảnh hưởng phần còn lại
                pass


def apply_crossref_marks(ws, marks: list, start_row, start_col):
    """
    marks: list các dict {'row': int(df_row), 'col': int(df_col), 'ok': bool|None, 'comment': str|None}
    """
    for m in marks:
        r, c = _dfpos_to_excel(m["row"], m["col"])
        cell = ws.cell(row=r + start_row, column=c + start_col)
        cell.alignment = RIGHT_ALIGN  # Áp dụng canh phải cho mọi ô

        # Đổi màu chữ theo trạng thái ok
        if m.get("ok") is True:
            cell.font = GREEN_FONT
            cell.value = "✔"

        elif m.get("ok") is False:
            cell.font = RED_FONT
            cell.value = "❌"
            try:
                # Tránh đè comment cũ: nối thêm
                if cell.comment:
                    new_text = cell.comment.text + "\n" + str(m["comment"])
                else:
                    new_text = str(m["comment"])

                comment = Comment(text=new_text, author="AutoCheck")
                comment.visible = True
                cell.comment = comment

                # cell.comment = Comment(text=new_text, author="AutoCheck")
            except Exception:
                # comment lỗi thì bỏ qua, không ảnh hưởng phần còn lại
                pass


# ------------------------------
# Reading Word tables with header: Đọc bảng từ file Word và lấy heading gần nhất trước mỗi bảng.
# ------------------------------
def read_word_tables_with_headings(file_path):
    doc = docx.Document(file_path)
    tables = []
    headings = []
    current_heading = None
    sec = 0
    current_section = doc.sections[sec]

    for block in doc.element.body:
        if block.tag.endswith("tbl"):
            #            Nếu chưa có heading, kiểm tra header của section
            if (
                current_heading is None
                or current_heading == "financial position"
                or current_heading == "statement of income"
            ):
                for para in current_section.header.paragraphs:
                    text = para.text.strip().lower()
                    if "financial position" in text:
                        current_heading = "financial position"
                        break
                    elif "statement of income" in text:
                        current_heading = "statement of income"
                        break
                    elif "statement of cash flows" in text:
                        current_heading = "statement of cash flows"
                        break
            # Đọc bảng
            table = docx.table.Table(block, doc)
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(pd.DataFrame(rows))
            headings.append(current_heading)
        elif block.tag.endswith("p"):
            paragraph = docx.text.paragraph.Paragraph(block, doc)
            # Kiểm tra nếu có sectPr (section break)
            sectPr = block.xpath(".//w:sectPr")
            if sectPr:
                # Cập nhật section hiện tại
                sec = sec + 1
                current_section = doc.sections[sec]
            if paragraph.style.name.startswith("Heading"):
                current_heading = paragraph.text.strip()

    return list(zip(tables, headings))


# --------------------------------------
# Find total row index: Ưu tiên dòng số "cuối cùng" mà dòng ngay trước đó TRỐNG HOÀN TOÀN (không chữ, không số).
# --------------------------------------
def find_total_row_index(df, headling_lower, _strict=True):
    """
    Xác định dòng tổng theo heuristics (không dùng keywords):
    - Ưu tiên dòng số "cuối cùng" mà dòng ngay trước đó TRỐNG HOÀN TOÀN (không chữ, không số).
    - Nếu không tìm thấy mẫu này:
        + _strict=True  => coi như bảng KHÔNG có dòng tổng -> trả về None
        + _strict=False => fallback: trả về dòng số cuối cùng của bảng
    """

    # ---- Helpers nội bộ ----
    def _as_numeric_series(row):
        return row.map(normalize_numeric_column)

    def _is_numeric_row(row) -> bool:
        ser = _as_numeric_series(row)
        return ser.notna().any()

    def _is_empty_row(row) -> bool:
        # “Trống” = không có số và không có text hữu ích (loại '-' '—' ngoặc, dấu phẩy, khoảng trắng)
        def _strip_text(x):
            s = str(x).strip()
            s = s.replace("-", "").replace("–", "").replace("—", "")
            s = s.replace("(", "").replace(")", "").replace(",", "")
            return s.strip()

        has_text = any(_strip_text(c) != "" for c in row)
        has_num = _is_numeric_row(row)
        return (not has_text) and (not has_num)

    # ---- Tìm tất cả các dòng có số ----
    numeric_rows = [i for i in range(len(df)) if _is_numeric_row(df.iloc[i])]
    if not numeric_rows:
        return None  # Bảng hoàn toàn không có số

    # ---- Từ dưới lên: chọn dòng số mà dòng trước đó "trống hoàn toàn" ----
    for i in reversed(numeric_rows):
        prev_empty = True
        if i - 1 >= 0:
            prev_empty = _is_empty_row(df.iloc[i - 1])
        if prev_empty:
            if (
                "accrued expenses" in headling_lower
                or "deferred revenue" in headling_lower
                or "other payables" in headling_lower
                or "short-term provisions" in headling_lower
            ) and (numeric_rows[-1] - i >= 2):
                return numeric_rows[-1]
            elif (
                "straight bonds and bonds convertible to a variable number of shares"
                in headling_lower
                or "convertible bonds" in headling_lower
                or "preference shares" in headling_lower
            ):
                while i > 0:
                    row = df.iloc[i]
                    row_text = " ".join(str(x).lower() for x in row)
                    if "within" not in row_text and "after" not in row_text:
                        if _is_numeric_row(df.iloc[i]):
                            return i
                        else:
                            i = i - 1
                    else:
                        i = i - 1
            elif "acquisition of subsidiary" in headling_lower:
                while i > 0:
                    row = df.iloc[i]
                    row_text = " ".join(str(x).lower() for x in row)
                    if "net identifiable" in row_text:
                        return i
                    else:
                        i = i - 1
            elif "business segments" in headling_lower:
                row = df.iloc[i]
                row_text = " ".join(str(x).lower() for x in row)
                if "after tax" in row_text:
                    while i > 0:
                        row = df.iloc[i]
                        row_text = " ".join(str(x).lower() for x in row)
                        if "segment revenue" in row_text:
                            return i
                        else:
                            i = i - 1
                else:
                    return i
            else:
                return i

    # ---- Không tìm thấy mẫu "trước trống" ----
    if _strict:
        return None  # KHÔNG có dòng tổng theo heuristics
    else:
        return numeric_rows[-1]  # Fallback: dùng dòng số cuối


# ======================
# Checking logic (router)
# ======================
def _parse_num(v):
    val = normalize_numeric_column(v)
    return 0.0 if pd.isna(val) else float(val)


def _norm_code(x: str) -> str:
    s = str(x).strip()
    s = s.replace("_", "").replace("**", "").replace("\u2212", "-").replace("–", "-")
    s = re.sub(r"[^0-9A-Za-z]", "", s)
    return s.upper()


# --------------------------
# Financial Position casting
# --------------------------
_FINANCIAL_POSTION_RULES = {
    "100": ["110", "120", "130", "140", "150", "160"],
    "110": ["111", "112"],
    "120": ["121", "122", "123", "124", "125", "126"],
    "130": ["131", "132", "133", "134", "135", "136", "137"],
    "140": ["141", "142"],
    "150": ["151", "152", "153"],
    "160": ["161", "162", "163", "164", "165"],
    "200": ["210", "220", "230", "240", "250", "260", "270"],
    "210": ["211", "212", "213", "214", "215", "216"],
    "220": ["221", "224", "227"],
    "221": ["222", "223"],
    "224": ["225", "226"],
    "227": ["228", "229"],
    "230": ["231", "236", "237", "238"],
    "231": ["232", "233"],
    "233": ["234", "235"],
    "240": ["241", "242"],
    "250": ["251", "252"],
    "260": ["261", "262", "263", "264", "265", "266"],
    "270": ["271", "272", "273", "274", "279"],
    "280": ["100", "200"],
    "300": ["310", "330"],
    "310": [
        "311",
        "312",
        "313",
        "314",
        "315",
        "316",
        "317",
        "318",
        "319",
        "320",
        "321",
        "322",
        "323",
        "324",
        "325",
    ],
    "330": [
        "331",
        "332",
        "333",
        "334",
        "335",
        "336",
        "337",
        "338",
        "339",
        "340",
        "341",
        "342",
        "343",
        "344",
    ],
    "400": [
        "411",
        "412",
        "413",
        "414",
        "415",
        "416",
        "417",
        "418",
        "419",
        "420",
        "429",
    ],
    "411": ["411A", "411B"],
    "420": ["420A", "420B"],
    "440": ["300", "400"],
}


def _find_header_idx_has_code(df: pd.DataFrame):
    for i in range(len(df)):
        row_strs = df.iloc[i].astype(str).str.lower()
        if row_strs.str.contains(_CODE_COL_NAME).any():
            return i
    return None


def _check_financial_position(df: pd.DataFrame) -> dict:
    # 1) Tìm hàng tiêu đề có chữ 'Code'
    header_idx = _find_header_idx_has_code(df)
    if header_idx is None:
        return {
            "status": "⚠️ Financial Position - không tìm thấy cột 'Code' để kiểm tra",
            "marks": [],
        }

    # 2) Dựng bảng dữ liệu với header từ hàng 'Code'
    header = [str(c).strip() for c in df.iloc[header_idx].tolist()]
    tmp = df.iloc[header_idx + 1 :].copy()
    tmp.columns = header

    # 3) Xác định cột 'Code' và 2 cột số (cột hiện tại & so sánh)
    code_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _CODE_COL_NAME), None
    )
    if code_col is None:
        return {
            "status": "⚠️ Financial Position - không xác định được cột 'Code'",
            "marks": [],
        }

    note_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _NOTE_COL_NAME), None
    )

    candidate_num_cols = [c for c in tmp.columns if _HEADER_DATE_RE.search(str(c))]
    if len(candidate_num_cols) >= 2:
        cur_col, prior_col = candidate_num_cols[:2]  # giữ thứ tự xuất hiện
    else:
        # fallback: dùng 2 cột cuối cùng
        cur_col, prior_col = tmp.columns[-2], tmp.columns[-1]

    # 4) Map code -> (current, prior) chỉ trong bảng hiện tại
    data = {}
    code_rowpos = {}  # map code -> tmp_row_index (0-based)
    issues = []
    marks = []
    cross_ref_marks = []

    for ridx, row in tmp.iterrows():
        code = _norm_code(row.get(code_col, ""))
        if not code:
            continue
        # Chấp nhận pattern 3 chữ số +/- hậu tố chữ (vd. 421B)
        if not re.match(r"^[0-9]+[A-Z]?$", code):
            continue
        cur_val = _parse_num(row.get(cur_col, ""))
        prior_val = _parse_num(row.get(prior_col, ""))
        # Nếu code lặp lại, ưu tiên bản có số (tránh ghi đè = 0)
        if code in data:
            old_cur, old_pr = data[code]
            if abs(cur_val) + abs(prior_val) == 0 and abs(old_cur) + abs(old_pr) != 0:
                continue
        data[code] = (cur_val, prior_val)
        if code == "280":
            BSPL_cross_check_cache["440"] = (cur_val, prior_val)
        if code == "440":
            account_name = code
            if account_name not in BSPL_cross_check_mark:
                CY_bal = cur_val
                PY_bal = prior_val
                cross_check_with_BSPL(
                    df,
                    cross_ref_marks,
                    issues,
                    account_name,
                    CY_bal,
                    PY_bal,
                    ridx,
                    len(df.columns) - 2,
                    0,
                    -1,
                )
                BSPL_cross_check_mark.append(account_name)

        if row.get(note_col, "") != "" "":
            acc_name = row.get(tmp.columns[0]).strip().lower()
            BSPL_cross_check_cache[acc_name] = (cur_val, prior_val)
            if code in [
                "123",
                "124",
                "264",
                "265",
                "266",
                "411",
                "412",
                "413",
                "414",
                "415",
                "416",
                "417",
                "418",
                "419",
                "420",
                "429",
            ]:
                BSPL_cross_check_cache[code] = (cur_val, prior_val)
            if code in ["261", "262", "263"]:
                try:
                    old_cur, old_pr = BSPL_cross_check_cache[
                        "investments in other entities"
                    ]
                except:
                    old_cur = 0
                    old_pr = 0
                BSPL_cross_check_cache["investments in other entities"] = (
                    cur_val + old_cur,
                    prior_val + old_pr,
                )

            if code in ["131", "211"]:
                try:
                    old_cur, old_pr = BSPL_cross_check_cache[
                        "accounts receivable from customers-combined"
                    ]
                except:
                    old_cur = 0
                    old_pr = 0
                BSPL_cross_check_cache[
                    "accounts receivable from customers-combined"
                ] = (cur_val + old_cur, prior_val + old_pr)

            if code in ["311", "331"]:
                try:
                    old_cur, old_pr = BSPL_cross_check_cache[
                        "accounts payable to suppliers-combined"
                    ]
                except:
                    old_cur = 0
                    old_pr = 0
                BSPL_cross_check_cache["accounts payable to suppliers-combined"] = (
                    cur_val + old_cur,
                    prior_val + old_pr,
                )

            if code in ["272", "342"]:
                try:
                    old_cur, old_pr = BSPL_cross_check_cache["Net_DTA_DTL"]
                except:
                    old_cur = 0
                    old_pr = 0
                if code == "342":
                    cur_val = -cur_val
                    prior_val = -prior_val
                BSPL_cross_check_cache["Net_DTA_DTL"] = (
                    cur_val + old_cur,
                    prior_val + old_pr,
                )

        else:
            if (
                code in valid_codes
                or code
                in [
                    "123",
                    "124",
                    "141",
                    "142",
                    "151",
                    "152",
                    "153",
                    "261",
                    "262",
                    "263",
                    "264",
                    "265",
                    "266",
                ]
                or code
                in [
                    "411",
                    "412",
                    "413",
                    "414",
                    "415",
                    "416",
                    "417",
                    "418",
                    "419",
                    "420",
                    "420B",
                    "429",
                ]
            ):
                if code in ["261", "262", "263"]:
                    try:
                        old_cur, old_pr = BSPL_cross_check_cache[
                            "investments in other entities"
                        ]
                    except:
                        old_cur = 0
                        old_pr = 0
                    BSPL_cross_check_cache["investments in other entities"] = (
                        cur_val + old_cur,
                        prior_val + old_pr,
                    )
                BSPL_cross_check_cache[code] = (cur_val, prior_val)
        # chỉ set hàng đầu tiên có số liệu
        code_rowpos.setdefault(code, ridx - tmp.index[0])  # quy về 0-based liên tục
    # vị trí cột trong df gốc
    try:
        cur_col_pos = header.index(cur_col)
        prior_col_pos = header.index(prior_col)
    except ValueError:
        cur_col_pos = len(header) - 2
        prior_col_pos = len(header) - 1

    # 5) Kiểm tra: skip nếu không có mã cha; skip nếu không có TẤT CẢ mã con;
    #    nếu thiếu MỘT VÀI mã con -> coi các mã thiếu = 0 để tính.

    for parent, children in _FINANCIAL_POSTION_RULES.items():
        parent_norm = _norm_code(parent)
        if parent_norm not in data:
            continue

        have_any = False
        cur_sum = prior_sum = 0.0
        missing = []
        for ch in children:
            ch_norm = _norm_code(ch)
            if ch_norm in data:
                ccur, cprior = data[ch_norm]
                cur_sum += ccur
                prior_sum += cprior
                have_any = True
            else:
                missing.append(ch_norm)

        if not have_any:
            # Không có TẤT CẢ mã con trong BẢNG NÀY -> skip
            continue

        parent_cur, parent_prior = data[parent_norm]
        diff_cur = cur_sum - parent_cur
        diff_prio = prior_sum - parent_prior
        is_ok_cy = abs(diff_cur) == 0
        is_ok_py = abs(diff_prio) == 0

        # Nếu biết vị trí dòng mã cha -> đánh dấu 2 ô số (cur/prior)
        if parent_norm in code_rowpos:
            df_row = (
                header_idx + 1 + code_rowpos[parent_norm]
            )  # vị trí trong df (0-based)
            comment = (
                f"{parent_norm} = sum({','.join(children)}); "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f}; "
                f"Thực tế={parent_cur:,.0f}/{parent_prior:,.0f}; "
                f"Δ={diff_cur:,.0f}/{diff_prio:,.0f}"
                + (f"; Thiếu={','.join(missing)}" if missing else "")
            )
            marks.append(
                {
                    "row": df_row,
                    "col": cur_col_pos,
                    "ok": is_ok_cy,
                    "comment": (None if is_ok_cy else comment),
                }
            )
            marks.append(
                {
                    "row": df_row,
                    "col": prior_col_pos,
                    "ok": is_ok_py,
                    "comment": (None if is_ok_py else comment),
                }
            )

        if not is_ok_cy or not is_ok_py:
            issues.append(comment)

    # 6) Kết luận trả về (ghi trực tiếp vào cột trạng thái của sheet bảng này)
    if not issues:
        status = "✅ Financial Position - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues)-10} dòng)" if len(issues) > 10 else ""
        status = f"❌ Financial Position - kiểm tra công thức: {len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": cross_ref_marks}


# ---------------------------------
# Statement of Income casting
# ---------------------------------
def _sum_weighted(data, children):
    have_any = False
    cur_sum = prior_sum = 0.0
    missing = []
    for token in children:
        token = str(token).strip()
        sign = -1 if token.startswith("-") else 1
        code = token[1:] if sign == -1 else token
        cn = _norm_code(code)
        if cn in data:
            ccur, cprior = data[cn]
            cur_sum += sign * ccur
            prior_sum += sign * cprior
            have_any = True
        else:
            missing.append(cn if sign == 1 else f"-{cn}")
    return have_any, cur_sum, prior_sum, missing


def _check_income_statement(df: pd.DataFrame) -> dict:
    # === Kiểm tra công thức Statement of Income ngay trong nhánh này ===
    # 1) Tìm hàng tiêu đề có chữ 'Code'
    header_idx = _find_header_idx_has_code(df)
    if header_idx is None:
        return {
            "status": "⚠️ Statement of income - không tìm thấy cột 'Code' để kiểm tra",
            "marks": [],
            "cross_ref_marks": [],
        }

    # 2) Dựng bảng dữ liệu với header từ hàng 'Code'
    header = [str(c).strip() for c in df.iloc[header_idx].tolist()]
    tmp = df.iloc[header_idx + 1 :].copy()
    tmp.columns = header
    tmp = tmp.reset_index(drop=True)

    # 3) Xác định cột 'Code' và 2 cột số (cột hiện tại & so sánh)
    code_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _CODE_COL_NAME), None
    )
    if code_col is None:
        return {
            "status": "⚠️ Statement of income - không xác định được cột 'Code'",
            "marks": [],
            "cross_ref_marks": [],
        }

    note_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _NOTE_COL_NAME), None
    )
    candidate_num_cols = [c for c in tmp.columns if _HEADER_DATE_RE.search(str(c))]
    if len(candidate_num_cols) >= 2:
        cur_col, prior_col = candidate_num_cols[:2]  # giữ thứ tự xuất hiện
    else:
        # fallback: dùng 2 cột cuối cùng
        cur_col, prior_col = tmp.columns[-2], tmp.columns[-1]

    # 4) Map code -> (current, prior) từ bảng hiện tại
    data = {}
    code_rowpos = {}
    issues = []
    marks = []
    cross_ref_marks = []

    for ridx, row in tmp.iterrows():
        code = _norm_code(row.get(code_col, ""))
        # Chấp nhận các code dạng số (giữ leading zero), có thể có hậu tố chữ (hiếm)
        if code == "" or not re.match(r"^[0-9]+[A-Z]?$", code):
            continue
        cur_val = _parse_num(row.get(cur_col, ""))
        prior_val = _parse_num(row.get(prior_col, ""))
        if code in data:
            old_cur, old_pr = data[code]
            if abs(cur_val) + abs(prior_val) == 0 and abs(old_cur) + abs(old_pr) != 0:
                continue
        data[code] = (cur_val, prior_val)
        if row.get(note_col, "") != "" "":
            acc_name = row.get(tmp.columns[0]).strip().lower()
            BSPL_cross_check_cache[acc_name] = (cur_val, prior_val)
            BSPL_cross_check_cache[code] = (cur_val, prior_val)
            if code in ["51", "52"]:
                try:
                    old_cur, old_pr = BSPL_cross_check_cache["income tax"]
                except:
                    old_cur = 0
                    old_pr = 0
                BSPL_cross_check_cache["income tax"] = (
                    cur_val + old_cur,
                    prior_val + old_pr,
                )

        else:
            if code in ["50", "60"]:
                BSPL_cross_check_cache[code] = (cur_val, prior_val)
                if code == "60":
                    account_name = "420B"
                    CY_bal = cur_val
                    PY_bal = prior_val
                    cross_check_with_BSPL(
                        df,
                        cross_ref_marks,
                        issues,
                        account_name,
                        CY_bal,
                        PY_bal,
                        ridx + header_idx + 1,
                        len(df.columns) - 2,
                        0,
                        -1,
                    )

        code_rowpos.setdefault(code, ridx)

    try:
        cur_col_pos = header.index(cur_col)
        prior_col_pos = header.index(prior_col)
    except ValueError:
        cur_col_pos = len(header) - 2
        prior_col_pos = len(header) - 1

    def check(parent, children, label=None):
        parent_norm = _norm_code(parent)
        if parent_norm not in data:
            return  # skip nếu không có mã cha
        have_any, cur_sum, prior_sum, missing = _sum_weighted(data, children)
        if not have_any:
            return  # skip nếu không có TẤT CẢ mã con
        ac_cur, ac_pr = data[parent_norm]
        dc = cur_sum - ac_cur
        dp = prior_sum - ac_pr
        is_ok_cy = abs(dc) == 0
        is_ok_py = abs(dp) == 0
        if parent_norm in code_rowpos:
            df_row = header_idx + 1 + code_rowpos[parent_norm]
            comment = (
                f"{parent_norm} = {' + '.join(children).replace('+ -',' - ')}; "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f}; Thực tế={ac_cur:,.0f}/{ac_pr:,.0f}; Δ={dc:,.0f}/{dp:,.0f}"
                + (f"; Thiếu={','.join(missing)}" if missing else "")
            )
            marks.append(
                {
                    "row": df_row,
                    "col": cur_col_pos,
                    "ok": is_ok_cy,
                    "comment": (None if is_ok_cy else comment),
                }
            )
            marks.append(
                {
                    "row": df_row,
                    "col": prior_col_pos,
                    "ok": is_ok_py,
                    "comment": (None if is_ok_py else comment),
                }
            )
        if not is_ok_cy or not is_ok_py:
            issues.append(comment)

    # 7) Áp công thức
    # 10 = 01 - 02
    check("10", ["01", "-02"])
    # 20 = 10 - 11
    check("20", ["10", "-11"])
    # 20 = 01 - 11
    check("20", ["01", "-02", "-11"])
    # 30 = 20 + (21 - 22) + 24 - (25 + 26)
    # check('30', ['20', '21', '-22', '24', '-25', '-26'])
    # 30 = 20 + 21 + (22 - 23) - (25 + 26) + 27
    check("30", ["20", "21", "22", "-23", "-25", "-26", "27"])
    # 40 = 31 - 32
    check("40", ["31", "-32"])
    # 50 = 30 + 40
    check("50", ["30", "40"])
    # 60 = 50 - 51 - 52
    check("60", ["50", "-51", "-52"])
    # 60 = 61 + 62, chỉ kiểm khi có đủ 61 & 62
    check("60", ["61", "62"])

    # 8) Kết luận trả về

    if not issues:
        status = "✅ Statement of income - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues)-10} dòng)" if len(issues) > 10 else ""
        status = f"❌ Statement of income - kiểm tra công thức: {len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": cross_ref_marks}


# ----------------------------
# Statement of Cash Flows casting
# ----------------------------
def _check_cash_flows(df: pd.DataFrame) -> dict:
    # 1) Tìm hàng tiêu đề có chữ 'Code'
    header_idx = _find_header_idx_has_code(df)
    if header_idx is None:
        return {
            "status": "⚠️ Cash flows - không tìm thấy cột 'Code' để kiểm tra",
            "marks": [],
            "cross_ref_marks": [],
        }

    # 2) Dựng bảng dữ liệu với header từ hàng 'Code'
    header = [str(c).strip() for c in df.iloc[header_idx].tolist()]
    tmp = df.iloc[header_idx + 1 :].copy()
    tmp.columns = header
    tmp = tmp.reset_index(drop=True)

    # 3) Xác định cột 'Code' và 2 cột số (cột hiện tại & so sánh)
    code_col = next(
        (c for c in tmp.columns if str(c).strip().lower() == _CODE_COL_NAME), None
    )
    if code_col is None:
        return {
            "status": "⚠️ Cash flows - không xác định được cột 'Code'",
            "marks": [],
            "cross_ref_marks": [],
        }

    candidate_num_cols = [c for c in tmp.columns if _HEADER_DATE_RE.search(str(c))]
    if len(candidate_num_cols) >= 2:
        cur_col, prior_col = candidate_num_cols[:2]  # giữ thứ tự xuất hiện
    else:
        cur_col, prior_col = tmp.columns[-2], tmp.columns[-1]

    # 4) Map code -> (current, prior) và lưu vị trí hàng
    data = {}
    code_rowpos = {}
    rows_cache = []  # [(ridx, code_norm, cur, prior)]
    cross_ref_marks = []
    issues = []
    marks = []

    for ridx, row in tmp.iterrows():
        code = _norm_code(row.get(code_col, ""))
        cur_val = _parse_num(row.get(cur_col, ""))
        prior_val = _parse_num(row.get(prior_col, ""))

        # Giữ cache thô để còn bắt "dòng không có code nhưng có tổng" (phục vụ Code 18)
        rows_cache.append((ridx, code, cur_val, prior_val))

        # Chỉ cộng dồn cho các mã code hợp lệ dạng số (có thể kèm suffix chữ)
        # Ví dụ: '05', '14', '21', hoặc '411A' (nếu có)
        if code and _CODE_VALID_RE.match(code):
            # ✅ Cộng dồn nếu cùng mã xuất hiện nhiều dòng (vd. nhiều dòng Code '05')
            agg_cur, agg_pr = data.get(code, (0.0, 0.0))
            data[code] = (agg_cur + cur_val, agg_pr + prior_val)
            code_rowpos.setdefault(code, ridx)
            if code == "70":
                account_name = "cash and cash equivalents"
                CY_bal = cur_val
                PY_bal = prior_val
                cross_check_with_BSPL(
                    df,
                    cross_ref_marks,
                    issues,
                    account_name,
                    CY_bal,
                    PY_bal,
                    ridx + header_idx + 1,
                    len(df.columns) - 2,
                    0,
                    -1,
                )

    # 5) Bắt trường hợp dòng "không có code nhưng có tổng" ngay trước Code 14~20 => đặt tạm là Code 18
    #    Code 18 = 08 + 09 + 10 + 11 + 12 + 13
    target_set = {str(i) for i in range(14, 21)}  # '14'..'20'
    first_block_idx = None
    for ridx, code, _, _ in rows_cache:
        if code in target_set:
            first_block_idx = ridx
            break
    if first_block_idx is not None and first_block_idx > 0:
        idx18 = first_block_idx - 1
        prev_code = rows_cache[idx18][1]
        prev_cur = rows_cache[idx18][2]
        prev_pr = rows_cache[idx18][3]
        # Nếu ngay trước là dòng không có code nhưng có số ⇒ coi đó là Code '18'
        if (prev_code == "" or prev_code is None) and (
            abs(prev_cur) + abs(prev_pr) != 0
        ):
            data["18"] = (prev_cur, prev_pr)
            code_rowpos.setdefault("18", idx18)
    try:
        cur_col_pos = header.index(cur_col)
        prior_col_pos = header.index(prior_col)
    except ValueError:
        cur_col_pos = len(header) - 2
        prior_col_pos = len(header) - 1

    def check(parent, children, label=None):
        parent_norm = _norm_code(parent)
        if parent_norm not in data:
            return
        have_any, cur_sum, prior_sum, missing = _sum_weighted(data, children)
        if not have_any:
            return
        ac_cur, ac_pr = data[parent_norm]
        dc = cur_sum - ac_cur
        dp = prior_sum - ac_pr

        is_ok_cy = abs(dc) == 0
        is_ok_py = abs(dp) == 0

        if parent_norm in code_rowpos:
            df_row = header_idx + 1 + code_rowpos[parent_norm]
            comment = (
                f"{parent_norm} = {' + '.join(children).replace('+ -',' - ')}; "
                f"Tính={cur_sum:,.0f}/{prior_sum:,.0f}; Thực tế={ac_cur:,.0f}/{ac_pr:,.0f}; Δ={dc:,.0f}/{dp:,.0f}"
                + (f"; Thiếu={','.join(missing)}" if missing else "")
            )
            marks.append(
                {
                    "row": df_row,
                    "col": cur_col_pos,
                    "ok": is_ok_cy,
                    "comment": (None if is_ok_cy else comment),
                }
            )
            marks.append(
                {
                    "row": df_row,
                    "col": prior_col_pos,
                    "ok": is_ok_py,
                    "comment": (None if is_ok_py else comment),
                }
            )
        if not is_ok_cy or not is_ok_py:
            issues.append(comment)

    # 8) Áp công thức cho Cash flows
    # 08 = 01 + 02 + 03 + 04 + 05 + 06 + 07
    check("08", ["01", "02", "03", "04", "05", "06", "07"])

    # 18 = 08 + 09 + 10 + 11 + 12 + 13 (nếu đã bắt được dòng 18, thì so khớp)
    if "18" in data:
        check("18", ["08", "09", "10", "11", "12", "13"])
    # 20 = 08 + 09 + 10 + 11 + 12 + 13 + 14 + 15 + 16 + 17
    check("20", ["08", "09", "10", "11", "12", "13", "14", "15", "16", "17"])
    # 30 = 21 + 22 + 23 + 24 + 25 + 26 + 27
    check("30", ["21", "22", "23", "24", "25", "26", "27"])
    # 40 = 31 + 32 + 33 + 34 + 35 + 36
    check("40", ["31", "32", "33", "34", "35", "36"])
    # 50 = 20 + 30 + 40
    check("50", ["20", "30", "40"])
    # 70 = 50 + 60 + 61
    check("70", ["50", "60", "61"])

    # 9) Kết luận trả về

    if not issues:
        status = "✅ Statement of cash flows - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues)-10} dòng)" if len(issues) > 10 else ""
        status = f"❌ Statement of cash flows - kiểm tra công thức: {len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": cross_ref_marks}


# -----------------------------------------------------------------------
# Income tax disclosure including DTAs and DTLs
# -----------------------------------------------------------------------
def _check_income_tax_rec_table(df: pd.DataFrame) -> dict:
    df_numeric = df.map(normalize_numeric_column)
    # Tìm dòng chứa "Accounting profit before tax" hoặc "Accounting loss before tax"
    profit_row_idx = None
    for i, row in df.iterrows():
        row_text = " ".join(str(cell).lower() for cell in row)
        if (
            "before tax" in row_text
        ):  # or "accounting loss before tax" in row_text or "accounting profit/(loss) before tax" in row_text or "accounting (loss)/profit before tax" in row_text:
            profit_row_idx = i
            account_name = "50"
            break

    if profit_row_idx is None:
        return {
            "status": "ℹ️ Không có dòng Accounting profit before tax hay Accounting loss before tax",
            "marks": [],
            "cross_ref_marks": [],
        }

    issues = []
    marks = []
    cross_ref_marks = []
    # cross check accounting profit/(loss) before tax
    CY_bal = (
        0
        if pd.isna(df_numeric.iloc[i, len(df.columns) - 2])
        else df_numeric.iloc[i, len(df.columns) - 2]
    )
    PY_bal = (
        0
        if pd.isna(df_numeric.iloc[i, len(df.columns) - 1])
        else df_numeric.iloc[i, len(df.columns) - 1]
    )
    if account_name not in BSPL_cross_check_mark:
        cross_check_with_BSPL(
            df,
            cross_ref_marks,
            issues,
            account_name,
            CY_bal,
            PY_bal,
            i,
            len(df.columns) - 2,
            0,
            -1,
        )
        BSPL_cross_check_mark.append(account_name)

    try:
        user_input = input("Vui lòng nhập thuế suất công ty (%): ")
        tax_rate = float(user_input.strip()) / 100
    except Exception:
        return {
            "status": "⚠️ Không thể đọc thuế suất từ người dùng",
            "marks": [],
            "cross_ref_marks": [],
        }

    # Bước 1: So sánh %tax rate của dòng profit với dòng tax rate cho từng cột
    profit_row = df_numeric.iloc[profit_row_idx]
    tax_rate_row_idx = None
    for i in range(profit_row_idx + 1, len(df)):
        row_text = " ".join(str(cell).lower() for cell in df.iloc[i])
        if (
            "tax at the company’s tax rate" in row_text
            or "tax at the group’s tax rate" in row_text
        ):
            tax_rate_row_idx = i
            break

    if tax_rate_row_idx is not None:
        tax_row = df_numeric.iloc[tax_rate_row_idx]
        for col in range(len(df.columns)):
            profit_val = profit_row[col]
            tax_val = tax_row[col]
            if not pd.isna(profit_val) and not pd.isna(tax_val):
                expected_tax = profit_val * tax_rate
                diff = expected_tax - tax_val
                is_ok = abs(round(diff)) == 0
                comment = f"Bước 1 - Cột {col+1}: {tax_rate*100}% lợi nhuận = {expected_tax:,.2f}, Thuế trên bảng = {tax_val:,.2f}, Sai lệch = {diff:,.2f}"
                marks.append(
                    {
                        "row": tax_rate_row_idx,
                        "col": col,
                        "ok": is_ok,
                        "comment": (None if is_ok else comment),
                    }
                )
                if not is_ok:
                    issues.append(comment)
        # Bước 2: Cộng dồn từ dòng sau dòng thuế đến dòng trống
        sum1 = [0.0] * len(df.columns)
        i = tax_rate_row_idx
        while i < len(df):
            row = df.iloc[i]
            if all(str(cell).strip() == "" for cell in row):
                break
            for col in range(len(df.columns)):
                val = df_numeric.iloc[i, col]
                if not pd.isna(val):
                    sum1[col] += val
            i += 1

        if i < len(df) - 1:
            total1_row = df_numeric.iloc[i + 1]
            for col in range(len(df.columns)):
                total_val = total1_row[col]
                if not pd.isna(total_val):
                    diff = sum1[col] - total_val
                    is_ok = abs(round(diff)) == 0
                    comment = f"Bước 2 - Cột {col+1}: Tổng chi tiết = {sum1[col]:,.2f}, Tổng 1 = {total_val:,.2f}, Sai lệch = {diff:,.2f}"
                    marks.append(
                        {
                            "row": i + 1,
                            "col": col,
                            "ok": is_ok,
                            "comment": (None if is_ok else comment),
                        }
                    )
                    if not is_ok:
                        issues.append(comment)

            # Bước 3: Cộng tiếp các dòng sau nếu có số liệu
            sum2 = sum1.copy()
            j = i + 2
            while j < len(df) - 1:
                row = df.iloc[j]
                if all(str(cell).strip() == "" for cell in row):
                    break
                for col in range(len(df.columns)):
                    val = df_numeric.iloc[j, col]
                    if not pd.isna(val):
                        sum2[col] += val
                j += 1

            if j < len(df):
                total2_row = df_numeric.iloc[len(df) - 1]
                for col in range(len(df.columns)):
                    total_val = total2_row[col]
                    if not pd.isna(total_val):
                        diff = sum2[col] - total_val
                        is_ok = abs(round(diff)) == 0
                        comment = f"Bước 3 - Cột {col+1}: Tổng cộng dồn = {sum2[col]:,.2f}, Tổng 2 = {total_val:,.2f}, Sai lệch = {diff:,.2f}"
                        marks.append(
                            {
                                "row": len(df) - 1,
                                "col": col,
                                "ok": is_ok,
                                "comment": (None if is_ok else comment),
                            }
                        )
                        if not is_ok:
                            issues.append(comment)

        account_name = "income tax"
        CY_bal = (
            0
            if pd.isna(df_numeric.iloc[len(df) - 1, len(df.columns) - 2])
            else df_numeric.iloc[len(df) - 1, len(df.columns) - 2]
        )
        PY_bal = (
            0
            if pd.isna(df_numeric.iloc[len(df) - 1, len(df.columns) - 1])
            else df_numeric.iloc[len(df) - 1, len(df.columns) - 1]
        )
        if account_name not in BSPL_cross_check_mark:
            cross_check_with_BSPL(
                df,
                cross_ref_marks,
                issues,
                account_name,
                CY_bal,
                PY_bal,
                len(df) - 1,
                len(df.columns) - 2,
                0,
                -1,
            )
            BSPL_cross_check_mark.append(account_name)

    if not issues:
        status = "✅ Reconciliation of effective tax rate - kiểm tra công thức: KHỚP (0 sai lệch)"
    else:
        preview = "; ".join(issues[:10])
        more = f" ... (+{len(issues)-10} dòng)" if len(issues) > 10 else ""
        status = f"❌ Reconciliation of effective tax rate - kiểm tra công thức: {len(issues)} sai lệch. {preview}{more}"

    return {"status": status, "marks": marks, "cross_ref_marks": cross_ref_marks}


# -----------------------------------------------------------------------
# FS casting
# -----------------------------------------------------------------------
def cross_check_with_BSPL(
    df,
    cross_ref_marks,
    issues,
    account_name,
    CY_bal,
    PY_bal,
    CY_row,
    CY_col,
    gap_row,
    gap_col,
):
    if account_name in BSPL_cross_check_cache:
        BSPL_CY_bal, BSPL_PY_bal = BSPL_cross_check_cache[account_name]
        diffCB = BSPL_CY_bal - CY_bal
        diffOB = BSPL_PY_bal - PY_bal
        is_okCB = abs(round(diffCB)) == 0
        is_okOB = abs(round(diffOB)) == 0

        if (
            (account_name in TABLES_NEED_CHECK_SEPARATELY)
            or (account_name in valid_codes)
            or (account_name in CROSS_CHECK_TABLES_FORM_1B)
            or (
                account_name
                in [
                    "50",
                    "construction in progress",
                    "investment property held for capital appreciation",
                    "long-term deferred expenses",
                    "allowance for doubtful debts",
                    "allowance for doubtful long-term debts",
                ]
            )
            or (account_name in ["151", "152"])
            or ((account_name in ["142"]) and (account_name in BSPL_cross_check_mark))
            or (
                account_name
                in [
                    "livestock producing one-time products – long-term",
                    "seasonal crops or plants producing one-time products – long-term",
                    "immature livestock producing periodic products",
                ]
            )
            or (account_name in ["deferred tax assets", "deferred tax liabilities"])
            or ("taxes" in account_name)
            or ("provisions" in account_name)
            or ("revenue from" in account_name)
            or ("revenue deductions" in account_name)
            or ("10" in account_name)
            or ("accounts receivable from customers" in account_name)
            or ("accounts payable to suppliers" in account_name)
        ):
            CY_row = CY_row - 1
            CY_col = len(df.columns)

        if account_name in ["investments in other entities", "264"]:
            if (
                gap_col == 0
                and CY_bal == PY_bal
                and account_name not in BSPL_cross_check_mark
            ):
                commentCB = f"BSPL = {BSPL_CY_bal:,.2f}, Note = {CY_bal:,.2f}, Diff = {diffCB:,.0f}"
                cross_ref_marks.append(
                    {
                        "row": CY_row,
                        "col": CY_col,
                        "ok": is_okCB,
                        "comment": (None if is_okCB else commentCB),
                    }
                )
                if not is_okCB:
                    issues.append(commentCB)
            elif (
                gap_col == 0
                and CY_bal == PY_bal
                and account_name in BSPL_cross_check_mark
            ):
                commentOB = f"BSPL = {BSPL_PY_bal:,.2f}, Note = {PY_bal:,.2f}, Diff = {diffOB:,.0f}"
                cross_ref_marks.append(
                    {
                        "row": CY_row - gap_row,
                        "col": CY_col - gap_col,
                        "ok": is_okOB,
                        "comment": (None if is_okOB else commentOB),
                    }
                )
                if not is_okOB:
                    issues.append(commentOB)

        else:
            commentCB = (
                f"BSPL = {BSPL_CY_bal:,.2f}, Note = {CY_bal:,.2f}, Diff = {diffCB:,.0f}"
            )
            cross_ref_marks.append(
                {
                    "row": CY_row,
                    "col": CY_col,
                    "ok": is_okCB,
                    "comment": (None if is_okCB else commentCB),
                }
            )
            if not is_okCB:
                issues.append(commentCB)

            commentOB = (
                f"BSPL = {BSPL_PY_bal:,.2f}, Note = {PY_bal:,.2f}, Diff = {diffOB:,.0f}"
            )
            cross_ref_marks.append(
                {
                    "row": CY_row - gap_row,
                    "col": CY_col - gap_col,
                    "ok": is_okOB,
                    "comment": (None if is_okOB else commentOB),
                }
            )
            if not is_okOB:
                issues.append(commentOB)


def check_table_total(df, heading=None):
    """Kiểm tra dòng tổng của bảng có khớp với chi tiết không và trả về trạng thái chi tiết."""
    df_numeric = df.map(normalize_numeric_column)
    heading_lower = heading.lower().strip() if heading else ""
    total_row_idx = find_total_row_index(df, heading_lower)
    last_col_idx = len(df.columns) - 1
    is_table_without_total = heading_lower in [
        name.lower() for name in TABLES_WITHOUT_TOTAL
    ]
    is_table_need_check_separately = heading_lower in [
        name.lower() for name in TABLES_NEED_CHECK_SEPARATELY
    ]  # dạng bảng có Costs, Accumulated depreciation and NBV
    is_table_with_column_total = heading_lower in [
        name.lower() for name in TABLES_NEED_COLUMN_CHECK
    ]

    if total_row_idx is None and not is_table_with_column_total:
        return {
            "status": "ℹ️ The table does not have a total row/column | Bảng không có dòng/cột tổng",
            "marks": [],
            "cross_ref_marks": [],
        }

    is_table_without_figures = False
    subset = df_numeric.iloc[
        2:total_row_idx
    ]  # ℹ️ chỗ này kiểm tra lại nên để total_row_idx hay last_col_idx???
    if subset.isna().all().all():
        is_table_without_figures = True

    try:
        if (
            is_table_without_total or is_table_without_figures
        ) and not is_table_with_column_total:  # Nếu là bảng mặc định không có tổng hoăc là bảng không có số từ dòng tổng trở lên VÀ bảng này không cần tính cột tổng
            cross_ref_marks = []
            issues = []
            if heading_lower not in CROSS_CHECK_TABLES_FORM_3 and (
                is_table_without_figures and not is_table_without_total
            ):
                # Nếu không phải dạng bảng 3 (i.e. dạng bảng đặc biệt có nhiều hơn 3 cột) + dạng bảng có dòng tổng và không có số tính từ dòng tổng trở lên, cross ref với BSPL
                # ℹ️ chỗ này kiểm tra lại có cần chèn thêm DK not in CROSS_CHECK_TABLES_FORM_3
                if (
                    not any(term in heading_lower for term in RE_PARTY_TABLE)
                    and heading_lower not in BSPL_cross_check_mark
                ):
                    account_name = heading_lower
                    CY_bal = (
                        0
                        if pd.isna(df_numeric.iloc[total_row_idx, len(df.columns) - 2])
                        else df_numeric.iloc[total_row_idx, len(df.columns) - 2]
                    )
                    PY_bal = (
                        0
                        if pd.isna(df_numeric.iloc[total_row_idx, len(df.columns) - 1])
                        else df_numeric.iloc[total_row_idx, len(df.columns) - 1]
                    )
                    cross_check_with_BSPL(
                        df,
                        cross_ref_marks,
                        issues,
                        account_name,
                        CY_bal,
                        PY_bal,
                        total_row_idx,
                        len(df.columns) - 2,
                        0,
                        -1,
                    )
                    BSPL_cross_check_mark.append(account_name)

            return {
                "status": "ℹ️ The table does not have a total row/column",
                "marks": [],
                "cross_ref_marks": cross_ref_marks,
            }

        else:
            if heading_lower == "financial position":
                return _check_financial_position(df)
            elif heading_lower == "statement of income":
                return _check_income_statement(df)
            elif heading_lower == "statement of cash flows":
                return _check_cash_flows(df)
            elif "reconciliation of effective tax rate" in heading_lower:
                return _check_income_tax_rec_table(df)
            else:
                # -----------------------------------------------------------------------
                # Cast dòng tổng trong bảng
                # -----------------------------------------------------------------------
                def find_block_sum(start_idx):
                    sum_vals = [0.0] * len(df.columns)
                    count = 0
                    i = start_idx + 1
                    while i < len(df) - 1:  # mới thêm vô -1 - check output
                        row = df.iloc[i]
                        if all(str(cell).strip() == "" for cell in row) and (
                            i - (start_idx + 1) > 0
                        ):
                            if "taxes" in heading_lower:
                                row1 = df.iloc[i + 1]
                                row1_text = " ".join(str(x).lower() for x in row1)
                                if (
                                    "other obligations" not in row1_text
                                    and "other receivables" not in row1_text
                                ):
                                    break
                            else:
                                break
                        for col in range(len(df.columns)):
                            val = df_numeric.iloc[i, col]
                            if not pd.isna(val):
                                sum_vals[col] += val
                        count += 1
                        i += 1
                    return sum_vals, count, i

                def compare_sum_with_total(sum_vals, total_row, end_row):
                    for col in range(len(df.columns)):
                        cell_value = df.iloc[end_row + 1, col]
                        if pd.isna(total_row[col]) and "-" in str(cell_value):
                            total_val = 0 if pd.isna(total_row[col]) else total_row[col]
                        else:
                            total_val = total_row[col]
                        if not pd.isna(sum_vals[col]) and not pd.isna(total_val):
                            diff = sum_vals[col] - total_val
                            is_ok = abs(round(diff)) == 0
                            comment = f"Col {col+1}: Recalculate = {sum_vals[col]:,.2f}, Note = {total_val:,.2f}, Diff = {diff:,.0f}"
                            marks.append(
                                {
                                    "row": end_row + 1,
                                    "col": col,
                                    "ok": is_ok,
                                    "comment": (None if is_ok else comment),
                                }
                            )
                            if not is_ok:
                                issues.append(comment)

                if is_table_need_check_separately:
                    # -----------------------------------------------------------------------
                    # #TABLES_NEED_CHECK_SEPARATELY included fixed asets disclosure
                    # -----------------------------------------------------------------------
                    # Tìm dòng chứa từ khóa tổng
                    cost_keywords = ["cost", "giá vốn"]
                    AD_keywords = [
                        "accumulated depreciation",
                        "accumulated amortisation",
                        "accumulated amortization",
                        "impairment losses",
                        "khấu hao lũy kế",
                        "hao mòn lũy kế",
                    ]
                    NBV_keywords = ["net book value", "giá trị còn lại"]
                    cost_row_idx = 0
                    cost_start_row_idx = 0
                    cost_total_row_idx = 0
                    AD_row_idx = 0
                    AD_start_row_idx = 0
                    AD_total_row_idx = 0
                    NBV_row_idx = 0
                    NBV_start_row_idx = 0
                    NBV_total_row_idx = 0

                    for i, row in df.iterrows():
                        if any(
                            keyword.lower() in str(cell).lower()
                            for cell in row
                            for keyword in cost_keywords
                        ):
                            cost_row_idx = i
                        if any(
                            keyword.lower() in str(cell).lower()
                            for cell in row
                            for keyword in AD_keywords
                        ):
                            AD_row_idx = i
                        if any(
                            keyword.lower() in str(cell).lower()
                            for cell in row
                            for keyword in NBV_keywords
                        ):
                            NBV_row_idx = i

                    for i, row in df.iterrows():
                        if any("closing" in str(cell).lower() for cell in row):
                            if i < AD_row_idx:
                                cost_total_row_idx = i
                            elif i < NBV_row_idx:
                                AD_total_row_idx = i
                                if (
                                    cost_total_row_idx == 0
                                    and cost_row_idx == AD_row_idx
                                ):
                                    cost_total_row_idx = i
                            else:
                                NBV_total_row_idx = i

                        if any("opening" in str(cell).lower() for cell in row):
                            if i < AD_row_idx:
                                cost_start_row_idx = i
                            elif i < NBV_row_idx:
                                AD_start_row_idx = i
                                if (
                                    cost_start_row_idx == 0
                                    and cost_row_idx == AD_row_idx
                                ):
                                    cost_start_row_idx = i
                            else:
                                NBV_start_row_idx = i

                    issues = []
                    marks = []
                    cross_ref_marks = []

                    if cost_row_idx != 0 and AD_row_idx != 0 and NBV_row_idx != 0:
                        # Case 1: Nếu là bảng movement
                        cost_detail_sum = df_numeric.iloc[
                            cost_start_row_idx : cost_total_row_idx - 1
                        ].sum(skipna=True)
                        cost_total_row = df_numeric.iloc[cost_total_row_idx].fillna(0)
                        AD_detail_sum = df_numeric.iloc[
                            AD_start_row_idx : AD_total_row_idx - 1
                        ].sum(skipna=True)
                        AD_total_row = df_numeric.iloc[AD_total_row_idx].fillna(0)
                        OB_detail_cal = df_numeric.iloc[cost_start_row_idx].fillna(
                            0
                        ) - df_numeric.iloc[AD_start_row_idx].fillna(0)
                        CB_detail_cal = cost_total_row - AD_total_row
                        OB_NBV_total_row = df_numeric.iloc[NBV_start_row_idx].fillna(0)
                        CB_NBV_total_row = df_numeric.iloc[NBV_total_row_idx].fillna(0)

                        for col in range(len(df.columns)):
                            total_val = (
                                0
                                if pd.isna(cost_total_row[col])
                                else cost_total_row[col]
                            )
                            if not pd.isna(cost_total_row[col]) and not pd.isna(
                                cost_detail_sum[col]
                            ):
                                diff = cost_detail_sum[col] - cost_total_row[col]
                                is_ok = abs(round(diff)) == 0
                                comment = f"Costs - Total row: Col - {col+1}: Recalculate={cost_detail_sum[col]:,.2f}, Note ={cost_total_row[col]:,.2f}, Diff={diff:,.2f}"
                                marks.append(
                                    {
                                        "row": cost_total_row_idx,
                                        "col": col,
                                        "ok": is_ok,
                                        "comment": (None if is_ok else comment),
                                    }
                                )
                                if not is_ok:
                                    issues.append(comment)

                        for col in range(len(df.columns)):
                            total_val = (
                                0 if pd.isna(AD_total_row[col]) else AD_total_row[col]
                            )
                            if not pd.isna(AD_total_row[col]) and not pd.isna(
                                AD_detail_sum[col]
                            ):
                                diff = AD_detail_sum[col] - AD_total_row[col]
                                is_ok = abs(round(diff)) == 0
                                comment = f"AD - Total row: Col - {col+1}: Recalculate={AD_detail_sum[col]:,.2f}, Note={AD_total_row[col]:,.2f}, Diff={diff:,.2f}"
                                marks.append(
                                    {
                                        "row": AD_total_row_idx,
                                        "col": col,
                                        "ok": is_ok,
                                        "comment": (None if is_ok else comment),
                                    }
                                )
                                if not is_ok:
                                    issues.append(comment)

                        for col in range(len(df.columns)):
                            total_val = (
                                0
                                if pd.isna(OB_NBV_total_row[col])
                                else OB_NBV_total_row[col]
                            )
                            if not pd.isna(OB_NBV_total_row[col]) and not pd.isna(
                                OB_detail_cal[col]
                            ):
                                diffOB = OB_detail_cal[col] - OB_NBV_total_row[col]
                                is_okOB = abs(round(diffOB)) == 0
                                commentOB = f"NBV - Opening balance row:  Col {col+1}: Recalculate ={OB_detail_cal[col]:,.2f}, Note ={OB_NBV_total_row[col]:,.2f}, Diff={diff:,.2f}"
                                marks.append(
                                    {
                                        "row": NBV_start_row_idx,
                                        "col": col,
                                        "ok": is_okOB,
                                        "comment": (None if is_okOB else commentOB),
                                    }
                                )
                                if not is_okOB:
                                    issues.append(commentOB)

                            total_val = (
                                0
                                if pd.isna(CB_NBV_total_row[col])
                                else CB_NBV_total_row[col]
                            )
                            if not pd.isna(CB_NBV_total_row[col]) and not pd.isna(
                                CB_detail_cal[col]
                            ):
                                diffCB = CB_detail_cal[col] - CB_NBV_total_row[col]
                                is_okCB = abs(round(diffCB)) == 0
                                commentCB = f"NBV - Closing balance row: Col {col+1}: Recalculate={CB_detail_cal[col]:,.2f}, Note={CB_NBV_total_row[col]:,.2f}, Diff={diff:,.2f}"
                                marks.append(
                                    {
                                        "row": NBV_total_row_idx,
                                        "col": col,
                                        "ok": is_okCB,
                                        "comment": (None if is_okCB else commentCB),
                                    }
                                )
                                if not is_okCB:
                                    issues.append(commentCB)

                        account_name = heading_lower

                        # NBV
                        CY_bal = (
                            0
                            if pd.isna(
                                df_numeric.iloc[NBV_total_row_idx, len(df.columns) - 1]
                            )
                            else df_numeric.iloc[NBV_total_row_idx, len(df.columns) - 1]
                        )
                        PY_bal = (
                            0
                            if pd.isna(
                                df_numeric.iloc[NBV_start_row_idx, len(df.columns) - 1]
                            )
                            else df_numeric.iloc[NBV_start_row_idx, len(df.columns) - 1]
                        )
                        if account_name not in BSPL_cross_check_mark:
                            cross_check_with_BSPL(
                                df,
                                cross_ref_marks,
                                issues,
                                account_name,
                                CY_bal,
                                PY_bal,
                                NBV_total_row_idx,
                                len(df.columns) - 1,
                                1,
                                0,
                            )
                            BSPL_cross_check_mark.append(account_name)

                        # Costs
                        CY_bal = (
                            0
                            if pd.isna(
                                df_numeric.iloc[cost_total_row_idx, len(df.columns) - 1]
                            )
                            else df_numeric.iloc[
                                cost_total_row_idx, len(df.columns) - 1
                            ]
                        )
                        PY_bal = (
                            0
                            if pd.isna(
                                df_numeric.iloc[cost_start_row_idx, len(df.columns) - 1]
                            )
                            else df_numeric.iloc[
                                cost_start_row_idx, len(df.columns) - 1
                            ]
                        )

                        if heading_lower == "tangible fixed assets":
                            account_name = "222"
                        elif heading_lower == "finance lease tangible fixed assets":
                            account_name = "225"
                        elif heading_lower == "intangible fixed assets":
                            account_name = "228"
                        elif (
                            heading_lower
                            == "mature livestock producing periodic products"
                        ):
                            account_name = "234"
                        elif heading_lower == "investment property held to earn rental":
                            if (
                                "investment property held for capital appreciation"
                                not in BSPL_cross_check_cache
                            ):
                                account_name = "241"

                        if (
                            account_name not in BSPL_cross_check_mark
                            and account_name
                            != "investment property held to earn rental"
                        ):
                            cross_check_with_BSPL(
                                df,
                                cross_ref_marks,
                                issues,
                                account_name,
                                CY_bal,
                                PY_bal,
                                cost_total_row_idx,
                                len(df.columns) - 1,
                                cost_total_row_idx - cost_start_row_idx,
                                0,
                            )
                            BSPL_cross_check_mark.append(account_name)

                        # Accumulated depreciation
                        CY_bal = (
                            0
                            if pd.isna(
                                (df_numeric.iloc[AD_total_row_idx, len(df.columns) - 1])
                                * -1
                            )
                            else (
                                df_numeric.iloc[AD_total_row_idx, len(df.columns) - 1]
                            )
                            * -1
                        )
                        PY_bal = (
                            0
                            if pd.isna(
                                (df_numeric.iloc[AD_start_row_idx, len(df.columns) - 1])
                                * -1
                            )
                            else (
                                df_numeric.iloc[AD_start_row_idx, len(df.columns) - 1]
                            )
                            * -1
                        )

                        if heading_lower == "tangible fixed assets":
                            account_name = "223"
                        elif heading_lower == "finance lease tangible fixed assets":
                            account_name = "226"
                        elif heading_lower == "intangible fixed assets":
                            account_name = "229"
                        elif (
                            heading_lower
                            == "mature livestock producing periodic products"
                        ):
                            account_name = "235"
                        elif heading_lower == "investment property held to earn rental":
                            account_name = "242"

                        if account_name not in BSPL_cross_check_mark:
                            cross_check_with_BSPL(
                                df,
                                cross_ref_marks,
                                issues,
                                account_name,
                                CY_bal,
                                PY_bal,
                                AD_total_row_idx,
                                len(df.columns) - 1,
                                AD_total_row_idx - AD_start_row_idx,
                                0,
                            )
                            BSPL_cross_check_mark.append(account_name)
                    else:
                        ##Case 2: Nếu là bảng breakdown trong các TABLES_NEED_CHECK_SEPARATELY, casting số tổng chứ ko check cross ref
                        if find_total_row_index is not None:
                            start_idx = 0
                            while start_idx < len(df):
                                row = df.iloc[start_idx]
                                row_text = " ".join(str(x).lower() for x in row)
                                if all(str(cell).strip() == "" for cell in row):
                                    break
                                start_idx += 1
                            total1 = [0.0] * len(df.columns)
                            sum1, count1, end1 = find_block_sum(start_idx)
                            if count1 > 1 and end1 < len(df) - 1:
                                total1_row = df_numeric.iloc[end1 + 1]
                                compare_sum_with_total(sum1, total1_row, end1)

                else:

                    def search_col_and_cross_ref(
                        key_word, account_name, total_row_xref
                    ):
                        CY_col = 0
                        PY_col = 0
                        for j, row in df.iterrows():
                            row_text = " ".join(str(x).lower() for x in row)
                            if key_word in row_text:
                                for col in range(len(df.columns)):
                                    if key_word in row.get(col, "").lower():
                                        if CY_col == 0:
                                            CY_col = col
                                            PY_col = CY_col
                                        else:
                                            PY_col = col
                                            break
                                break

                        if (
                            "borrowings" in account_name
                            or "bonds" in account_name
                            or "borrowing" in account_name
                            or "bond" in account_name
                        ):
                            temp_col = PY_col
                            PY_col = CY_col
                            CY_col = temp_col

                        if "shortage of assets awaiting resolution" in account_name:
                            if CY_col == 0 and PY_col == 0:
                                if len(df.columns) - 1 == 2:
                                    CY_col = len(df.columns) - 2
                                    PY_col = len(df.columns) - 1
                            elif CY_col != 0 and PY_col != 0:
                                CY_col += 1
                                PY_col += 1

                        CY_bal = (
                            0
                            if pd.isna(df_numeric.iloc[total_row_xref, CY_col])
                            else df_numeric.iloc[total_row_xref, CY_col]
                        )
                        PY_bal = (
                            0
                            if pd.isna(df_numeric.iloc[total_row_xref, PY_col])
                            else df_numeric.iloc[total_row_xref, PY_col]
                        )
                        if (
                            account_name
                            == "allowance for diminution in the value of trading securities"
                            or account_name == "124"
                            or account_name == "264"
                            or account_name == "266"
                            or account_name == "allowance for doubtful debts"
                            or account_name == "allowance for doubtful long-term debts"
                        ):
                            CY_bal = CY_bal * (-1)
                            PY_bal = PY_bal * (-1)
                        if CY_col != 0 or PY_col != 0:
                            cross_check_with_BSPL(
                                df,
                                cross_ref_marks,
                                issues,
                                account_name,
                                CY_bal,
                                PY_bal,
                                total_row_xref,
                                CY_col,
                                0,
                                CY_col - PY_col,
                            )
                        BSPL_cross_check_mark.append(account_name)

                    def search_row_and_cross_ref(account_name, col_xref):
                        CY_row = len(df) - 1  # lúc trước là subtotalrows[0]
                        PY_row = 0
                        for j, row in df.iterrows():
                            row_text = " ".join(str(x).lower() for x in row)
                            if "opening balance" in row_text:
                                PY_row = j
                                break

                        CY_bal = (
                            0
                            if pd.isna(df_numeric.iloc[CY_row, col_xref])
                            else df_numeric.iloc[CY_row, col_xref]
                        )
                        PY_bal = (
                            0
                            if pd.isna(df_numeric.iloc[PY_row, col_xref])
                            else df_numeric.iloc[PY_row, col_xref]
                        )
                        if CY_row != 0 or PY_row != 0:
                            cross_check_with_BSPL(
                                df,
                                cross_ref_marks,
                                issues,
                                account_name,
                                CY_bal,
                                PY_bal,
                                CY_row,
                                col_xref,
                                CY_row - PY_row,
                                0,
                            )
                        BSPL_cross_check_mark.append(account_name)

                    issues = []
                    marks = []
                    cross_ref_marks = []
                    if find_total_row_index is not None:
                        start_idx = 0
                        while start_idx < len(df):
                            row = df.iloc[start_idx]
                            row_text = " ".join(str(x).lower() for x in row)
                            if (
                                all(str(cell).strip() == "" for cell in row)
                                or "equity investments" in row_text
                                or "balance at" in row_text
                            ):
                                break
                            start_idx += 1

                        n_cols = len(df.columns)
                        grand_components = [0.0] * n_cols
                        end_blocks = (
                            []
                        )  # lưu end_block (dòng trống / điểm kết thúc phần detail)
                        subtotal_rows = []  # lưu dòng subtotal (end_block + 1)
                        if is_table_without_figures and start_idx + 1 >= total_row_idx:
                            subtotal_rows.append(total_row_idx)

                        while start_idx + 1 < total_row_idx:
                            sum_block, count_block, end_block = find_block_sum(
                                start_idx
                            )
                            # Nếu không có dòng detail nào thì thoát, tránh vòng lặp vô hạn
                            if count_block == 0:
                                break

                            if heading_lower == "changes in owners’ equity":
                                for col in range(len(df.columns)):
                                    sum_block[col] += grand_components[col]

                            subtotal_row = None
                            subtotal_row_idx = None

                            if count_block > 1 and end_block < len(df) - 1:
                                subtotal_row_idx = end_block + 1
                                subtotal_row = df_numeric.iloc[subtotal_row_idx]
                                end_blocks.append(end_block)
                                subtotal_rows.append(subtotal_row_idx)
                                compare_sum_with_total(
                                    sum_block, subtotal_row, end_block
                                )

                            # Xác định đóng góp của block vào grand total
                            contrib = list(sum_block)
                            if (
                                "revenue from" in heading_lower
                                or "recognised deferred tax" in heading_lower
                            ):
                                if (
                                    len(subtotal_rows) == 2
                                    and subtotal_row is not None
                                    and subtotal_row.dropna().gt(0).all()
                                ):
                                    contrib = [-v for v in sum_block]

                            for col in range(n_cols):
                                if heading_lower == "changes in owners’ equity":
                                    grand_components[col] = contrib[col]
                                else:
                                    grand_components[col] += contrib[col]

                            if (
                                "recognised deferred tax" in heading_lower
                                and len(subtotal_rows) == 2
                            ):
                                break

                            if count_block == 1:
                                subtotal_rows.append(end_block - 1)
                                new_start_idx = end_block - 1
                            else:
                                new_start_idx = end_block + 1

                            if new_start_idx >= total_row_idx:
                                break
                            start_idx = new_start_idx

                        if len(subtotal_rows) > 1:
                            if "recognised deferred tax" in heading_lower:
                                final_row = df_numeric.iloc[subtotal_rows[1] + 1]
                            elif (
                                "straight bonds and bonds convertible to a variable number of shares"
                                in heading_lower
                                or "convertible bonds" in heading_lower
                                or "preference shares" in heading_lower
                                or "long-term provisions" in heading_lower
                            ):
                                final_row = df_numeric.iloc[total_row_idx]
                                if (
                                    "convertible bonds" in heading_lower
                                    or "preference shares" in heading_lower
                                    or "long-term provisions" in heading_lower
                                ):
                                    for col in range(n_cols):
                                        grand_components[col] -= contrib[col]

                                if "long-term provisions" in heading_lower:
                                    for col in range(len(df.columns)):
                                        combined = contrib[col]
                                        final_val = (
                                            0
                                            if pd.isna(final_row[col])
                                            else final_row[col]
                                        )
                                        if not pd.isna(final_val) and not pd.isna(
                                            combined
                                        ):
                                            diff = combined - final_val
                                            is_ok = abs(round(diff)) == 0
                                            comment = f"Grand total row - Col {col+1}: Recalculate = {combined:,.2f}, Note = {final_val:,.2f}, Diff = {diff:,.0f}"
                                            marks.append(
                                                {
                                                    "row": total_row_idx,
                                                    "col": col,
                                                    "ok": is_ok,
                                                    "comment": (
                                                        None if is_ok else comment
                                                    ),
                                                }
                                            )
                                            if not is_ok:
                                                issues.append(comment)

                            else:
                                final_row = df_numeric.iloc[len(df) - 1]

                            if (
                                "held-to-maturity investments" in heading_lower
                            ):  # Case 8: Cast and Cross ref Subtotal 1, 2 và bảng không có dòng Grand Total
                                # search and cross ref costs
                                account_name = "123"
                                if account_name not in BSPL_cross_check_mark:
                                    search_col_and_cross_ref(
                                        "cost", account_name, subtotal_rows[0]
                                    )
                                account_name = "265"
                                if account_name not in BSPL_cross_check_mark:
                                    search_col_and_cross_ref(
                                        "cost", account_name, subtotal_rows[-1]
                                    )

                                # search and cross ref allowance
                                account_name = "124"
                                if account_name not in BSPL_cross_check_mark:
                                    search_col_and_cross_ref(
                                        "allowance", account_name, subtotal_rows[0]
                                    )
                                account_name = "266"
                                if account_name not in BSPL_cross_check_mark:
                                    search_col_and_cross_ref(
                                        "allowance", account_name, subtotal_rows[-1]
                                    )

                            elif (
                                "other parties under bcc contracts" in heading_lower
                            ):  # Case 8: Cast and Cross ref Subtotal 1, 2 và bảng không có dòng Grand Total
                                assets = df_numeric.iloc[subtotal_rows[0]]
                                liabilities = df_numeric.iloc[subtotal_rows[1]]

                                for col in range(len(df.columns)):
                                    ass_val = 0 if pd.isna(assets[col]) else assets[col]
                                    lia_val = (
                                        0
                                        if pd.isna(liabilities[col])
                                        else liabilities[col]
                                    )

                                    if not pd.isna(ass_val) and not pd.isna(lia_val):
                                        diff = ass_val - lia_val
                                        is_ok = abs(round(diff)) == 0
                                        comment = f"TA - TL: Col {col+1}: Total assets = {ass_val:,.2f} - Total liabilities {lia_val:,.2f} = {diff:,.0f}"
                                        marks.append(
                                            {
                                                "row": len(df),
                                                "col": col,
                                                "ok": is_ok,
                                                "comment": (None if is_ok else comment),
                                            }
                                        )
                                        if not is_ok:
                                            issues.append(comment)

                            elif "changes in owners’ equity" in heading_lower:
                                OB_row = CB_row = None
                                for j, row in df.iloc[::-1].iterrows():
                                    row_text = " ".join(str(x).lower() for x in row)
                                    if "balance at" in row_text:
                                        if CB_row is None:
                                            CB_row = j
                                            final_CBrow = df_numeric.loc[j]
                                        else:
                                            if OB_row is None:
                                                OB_row = j
                                                final_OBrow = df_numeric.loc[j]
                                    if OB_row is not None and CB_row is not None:
                                        break

                                for col in range(len(df.columns)):
                                    cell_value = str(df.iloc[0, col]).lower()
                                    account_name = " "
                                    if (
                                        "contributed capital" in cell_value
                                        or "share capital" in cell_value
                                    ):
                                        account_name = "411"
                                    elif (
                                        "capital surplus" in cell_value
                                        or "share premium" in cell_value
                                    ):
                                        account_name = "412"
                                    elif "convertible bonds" in cell_value:
                                        account_name = "413"
                                    elif "other capital" in cell_value:
                                        account_name = "414"
                                    elif "own shares" in cell_value:
                                        account_name = "415"
                                    elif "revaluation" in cell_value:
                                        account_name = "416"
                                    elif "exchange differences" in cell_value:
                                        account_name = "417"
                                    elif (
                                        "investment" in cell_value
                                        or "developement" in cell_value
                                    ):
                                        account_name = "418"
                                    elif "other equity funds" in cell_value:
                                        account_name = "419"
                                    elif (
                                        "retained profits" in cell_value
                                        or "accumulated losses" in cell_value
                                    ):
                                        account_name = "420"
                                    elif "non-controlling interest" in cell_value:
                                        account_name = "429"

                                    CY_bal = (
                                        0
                                        if pd.isna(df_numeric.iloc[CB_row, col])
                                        else df_numeric.iloc[CB_row, col]
                                    )
                                    PY_bal = (
                                        0
                                        if pd.isna(df_numeric.iloc[OB_row, col])
                                        else df_numeric.iloc[OB_row, col]
                                    )

                                    if (
                                        account_name not in BSPL_cross_check_mark
                                        and account_name != " "
                                    ):
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            CB_row,
                                            col,
                                            CB_row - OB_row,
                                            0,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                            elif (
                                "share capital" in heading_lower
                                or "business segments" in heading_lower
                            ):  # Case 9
                                pass

                            else:
                                if (
                                    "construction contracts" not in heading_lower
                                    and "bad and doubtful debts" not in heading_lower
                                ):
                                    for col in range(len(df.columns)):
                                        combined = grand_components[col]
                                        if (
                                            "straight bonds and bonds convertible to a variable number of shares"
                                            in heading_lower
                                            or "convertible bonds" in heading_lower
                                            or "preference shares" in heading_lower
                                        ):
                                            cell_value = df.iloc[total_row_idx, col]
                                            if pd.isna(final_row[col]) and "-" in str(
                                                cell_value
                                            ):
                                                final_val = 0
                                            else:
                                                final_val = final_row[col]
                                        else:
                                            final_val = (
                                                0
                                                if pd.isna(final_row[col])
                                                else final_row[col]
                                            )

                                        if not pd.isna(final_val) and not pd.isna(
                                            combined
                                        ):
                                            diff = combined - final_val
                                            is_ok = abs(round(diff)) == 0
                                            comment = f"Grand total row - Col {col+1}: Recalculate = {combined:,.2f}, Note = {final_val:,.2f}, Diff = {diff:,.0f}"

                                            if (
                                                "recognised deferred tax"
                                                in heading_lower
                                            ):
                                                marks.append(
                                                    {
                                                        "row": subtotal_rows[1] + 1,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                            elif (
                                                "straight bonds and bonds convertible to a variable number of shares"
                                                in heading_lower
                                                or "convertible bonds" in heading_lower
                                                or "preference shares" in heading_lower
                                            ):
                                                marks.append(
                                                    {
                                                        "row": total_row_idx,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                            else:
                                                marks.append(
                                                    {
                                                        "row": len(df) - 1,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )

                                            if not is_ok:
                                                issues.append(comment)

                                    if (
                                        "straight bonds and bonds convertible to a variable number of shares"
                                        in heading_lower
                                        or "convertible bonds" in heading_lower
                                        or "preference shares" in heading_lower
                                    ):
                                        ST_row = 0
                                        LT_row = 0
                                        final_STrow = [0.0] * len(df.columns)
                                        final_LTrow = [0.0] * len(df.columns)
                                        for j, row in df.iloc[
                                            total_row_idx:
                                        ].iterrows():
                                            row_text = " ".join(
                                                str(x).lower() for x in row
                                            )
                                            if "within" in row_text:
                                                ST_row = j
                                                final_STrow = df_numeric.iloc[ST_row]
                                            if "after" in row_text:
                                                LT_row = j
                                                final_LTrow = df_numeric.iloc[LT_row]

                                        for col in range(len(df.columns)):
                                            combined = grand_components[col]
                                            final_val = (
                                                final_LTrow[col] - final_STrow[col]
                                            )
                                            if not pd.isna(final_val):
                                                diff = combined - final_val
                                                is_ok = abs(round(diff)) == 0
                                                comment = f"Grand total row - Col {col+1}: Recalculate = {combined:,.2f}, Note = {final_val:,.2f}, Diff = {diff:,.0f}"
                                                marks.append(
                                                    {
                                                        "row": LT_row,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                                if not is_ok:
                                                    issues.append(comment)

                                        if (
                                            "convertible bonds" in heading_lower
                                            or "preference shares" in heading_lower
                                        ):
                                            CY_row = LT_row
                                            account_name = heading_lower
                                            if "convertible bonds" in heading_lower:
                                                col_xref = len(df.columns) - 2
                                            else:
                                                col_xref = len(df.columns) - 1
                                            PY_row = 0
                                            count = 0
                                            for j, row in df.iterrows():
                                                row_text = " ".join(
                                                    str(x).lower() for x in row
                                                )
                                                if "balance at" in row_text:
                                                    if count == 0:
                                                        PY_row = subtotal_rows[0]
                                                    count += 1

                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[CY_row, col_xref]
                                                )
                                                else df_numeric.iloc[CY_row, col_xref]
                                            )
                                            if count > 1:
                                                PY_bal = (
                                                    0
                                                    if pd.isna(
                                                        df_numeric.iloc[
                                                            PY_row, col_xref
                                                        ]
                                                    )
                                                    else df_numeric.iloc[
                                                        PY_row, col_xref
                                                    ]
                                                )
                                            if CY_row != 0 or PY_row != 0:
                                                cross_check_with_BSPL(
                                                    df,
                                                    cross_ref_marks,
                                                    issues,
                                                    account_name,
                                                    CY_bal,
                                                    PY_bal,
                                                    CY_row,
                                                    col_xref,
                                                    CY_row - PY_row,
                                                    0,
                                                )
                                            BSPL_cross_check_mark.append(account_name)

                                elif (
                                    "bad and doubtful debts" in heading_lower
                                ):  # Case 10
                                    ST_row = 0
                                    LT_row = 0
                                    mnt = True
                                    final_STrow = [0.0] * len(df.columns)
                                    final_LTrow = [0.0] * len(df.columns)
                                    for j, row in df.iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if (
                                            "cost" in row_text
                                            or "recoverable amount" in row_text
                                            or "overdue days" in row_text
                                        ):
                                            mnt = False
                                        if "short-term" in row_text:
                                            ST_row = j
                                            final_STrow = df_numeric.iloc[ST_row]
                                        if "long-term" in row_text:
                                            LT_row = j
                                            final_LTrow = df_numeric.iloc[LT_row]

                                    for col in range(len(df.columns)):
                                        combined = grand_components[col]
                                        if mnt:
                                            final_val = final_row[col]
                                        else:
                                            final_val = (
                                                final_STrow[col] + final_LTrow[col]
                                            )
                                        if not pd.isna(final_val):
                                            diff = combined - final_val
                                            is_ok = abs(round(diff)) == 0
                                            comment = f"Grand total row - Col {col+1}: Recalculate = {combined:,.2f}, Note = {final_val:,.2f}, Diff = {diff:,.0f}"
                                            marks.append(
                                                {
                                                    "row": len(df) - 1,
                                                    "col": col,
                                                    "ok": is_ok,
                                                    "comment": (
                                                        None if is_ok else comment
                                                    ),
                                                }
                                            )
                                            if not is_ok:
                                                issues.append(comment)

                                    if (
                                        ST_row != 0
                                        or "allowance for doubtful long-term debts"
                                        not in BSPL_cross_check_cache
                                    ):
                                        # search and cross ref ST
                                        account_name = "allowance for doubtful debts"
                                        if ST_row == 0:
                                            ST_row = subtotal_rows[-1]
                                        if not mnt:
                                            search_col_and_cross_ref(
                                                "allowance", account_name, ST_row
                                            )
                                        else:
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        ST_row, len(df.columns) - 2
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    ST_row, len(df.columns) - 2
                                                ]
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        ST_row, len(df.columns) - 1
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    ST_row, len(df.columns) - 1
                                                ]
                                            )
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                ST_row,
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)
                                        # search and cross ref LT
                                    if (
                                        LT_row != 0
                                        or "allowance for doubtful debts"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = (
                                            "allowance for doubtful long-term debts"
                                        )
                                        if LT_row == 0:
                                            LT_row = subtotal_rows[-1]
                                        if not mnt:
                                            search_col_and_cross_ref(
                                                "allowance", account_name, LT_row
                                            )
                                        else:
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LT_row, len(df.columns) - 2
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    LT_row, len(df.columns) - 2
                                                ]
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LT_row, len(df.columns) - 1
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    LT_row, len(df.columns) - 1
                                                ]
                                            )
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                LT_row,
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                if (
                                    heading_lower in CROSS_CHECK_TABLES_FORM_1A
                                ):  # Case 5: Cast and Cross ref Grand Total lên BS
                                    if (
                                        "recognised deferred tax" in heading_lower
                                        and subtotal_rows[1] + 3 < len(df) - 1
                                    ):
                                        LTA_row = 0
                                        LTL_row = 0
                                        final_LTArow = [0.0] * len(df.columns)
                                        final_LTLrow = [0.0] * len(df.columns)

                                        for j, row in df.iloc[
                                            subtotal_rows[1] + 1 :
                                        ].iterrows():
                                            row_text = " ".join(
                                                str(x).lower() for x in row
                                            )
                                            if "long-term assets" in row_text:
                                                LTA_row = j
                                                final_LTArow = df_numeric.iloc[LTA_row]
                                            if "long-term liabilities" in row_text:
                                                LTL_row = j
                                                final_LTLrow = df_numeric.iloc[LTL_row]

                                        final_row2 = df_numeric.iloc[len(df) - 1]
                                        for col in range(len(df.columns)):
                                            combined = (
                                                final_LTArow[col] + final_LTLrow[col]
                                            )
                                            final_val2 = final_row2[col]

                                            if not pd.isna(final_val2):
                                                diff = combined - final_val2
                                                is_ok = abs(round(diff)) == 0
                                                comment = f"Grand total row - Col {col+1}: Recalculate = {combined:,.2f}, Note = {final_val2:,.2f}, Diff = {diff:,.0f}"
                                                marks.append(
                                                    {
                                                        "row": len(df) - 1,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                                if not is_ok:
                                                    issues.append(comment)
                                        # search and cross ref Net DTA_DTL within note
                                        account_name = "Net_DTA_DTL"
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[1] + 1,
                                                    len(df.columns) - 2,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[1] + 1,
                                                len(df.columns) - 2,
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[1] + 1,
                                                    len(df.columns) - 1,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[1] + 1,
                                                len(df.columns) - 1,
                                            ]
                                        )
                                        BSPL_cross_check_cache[account_name] = (
                                            CY_bal,
                                            PY_bal,
                                        )
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    len(df) - 1, len(df.columns) - 2
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                len(df) - 1, len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    len(df) - 1, len(df.columns) - 1
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                len(df) - 1, len(df.columns) - 1
                                            ]
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                len(df) - 1,
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                        if LTA_row != 0:
                                            # search and cross ref DTA
                                            account_name = "deferred tax assets"
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LTA_row, len(df.columns) - 2
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    LTA_row, len(df.columns) - 2
                                                ]
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LTA_row, len(df.columns) - 1
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    LTA_row, len(df.columns) - 1
                                                ]
                                            )
                                            if (
                                                account_name
                                                not in BSPL_cross_check_mark
                                            ):
                                                cross_check_with_BSPL(
                                                    df,
                                                    cross_ref_marks,
                                                    issues,
                                                    account_name,
                                                    CY_bal,
                                                    PY_bal,
                                                    LTA_row,
                                                    len(df.columns) - 2,
                                                    0,
                                                    -1,
                                                )
                                                BSPL_cross_check_mark.append(
                                                    account_name
                                                )
                                            # search and cross ref DTL
                                        if LTL_row != 0:
                                            account_name = "deferred tax liabilities"
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LTL_row, len(df.columns) - 2
                                                    ]
                                                )
                                                else (
                                                    df_numeric.iloc[
                                                        LTL_row, len(df.columns) - 2
                                                    ]
                                                    * -1
                                                )
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LTL_row, len(df.columns) - 1
                                                    ]
                                                )
                                                else (
                                                    df_numeric.iloc[
                                                        LTL_row, len(df.columns) - 1
                                                    ]
                                                    * -1
                                                )
                                            )
                                            if (
                                                account_name
                                                not in BSPL_cross_check_mark
                                            ):
                                                cross_check_with_BSPL(
                                                    df,
                                                    cross_ref_marks,
                                                    issues,
                                                    account_name,
                                                    CY_bal,
                                                    PY_bal,
                                                    LTL_row,
                                                    len(df.columns) - 2,
                                                    0,
                                                    -1,
                                                )
                                                BSPL_cross_check_mark.append(
                                                    account_name
                                                )

                                    else:
                                        if "recognised deferred tax" in heading_lower:
                                            account_name = "Net_DTA_DTL"
                                        elif (
                                            heading_lower
                                            == "accrued expenses – short-term"
                                        ):
                                            account_name = "accrued expenses"
                                        elif (
                                            heading_lower
                                            == "accrued expenses – long-term"
                                        ):
                                            account_name = "long-term accrued expenses"
                                        elif (
                                            heading_lower
                                            == "deferred revenue – long-term"
                                        ):
                                            account_name = "long-term deferred revenue"
                                        else:
                                            account_name = heading_lower
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    len(df) - 1, len(df.columns) - 2
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                len(df) - 1, len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    len(df) - 1, len(df.columns) - 1
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                len(df) - 1, len(df.columns) - 1
                                            ]
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            if (
                                                "payables on construction contracts"
                                                in account_name
                                            ):
                                                CY_bal = -CY_bal
                                                PY_bal = -PY_bal
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                len(df) - 1,
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                        BSPL_cross_check_mark.append(account_name)

                                elif (
                                    heading_lower in CROSS_CHECK_TABLES_FORM_1B
                                ):  # Case 4: Cast and Cross ref subtotal 1 lên BS
                                    account_name = heading_lower
                                    CY_bal = (
                                        0
                                        if pd.isna(
                                            df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 2
                                            ]
                                        )
                                        else df_numeric.iloc[
                                            subtotal_rows[0], len(df.columns) - 2
                                        ]
                                    )
                                    PY_bal = (
                                        0
                                        if pd.isna(
                                            df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 1
                                            ]
                                        )
                                        else df_numeric.iloc[
                                            subtotal_rows[0], len(df.columns) - 1
                                        ]
                                    )
                                    if account_name not in BSPL_cross_check_mark:
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            subtotal_rows[0],
                                            len(df.columns) - 2,
                                            0,
                                            -1,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                                elif (
                                    heading_lower in CROSS_CHECK_TABLES_FORM_2
                                ):  # Case 6: Cast and Cross ref Subtotal 1, 2 and Grand Total lên BS
                                    account_name = heading_lower
                                    if account_name == "biological assets – short-term":
                                        account_name = "151"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )

                                        account_name = "152"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[1]
                                            )

                                        account_name = "153"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "allowance", account_name, len(df) - 1
                                            )
                                    elif (
                                        "biological assets – long-term" in account_name
                                    ):
                                        account_name = "livestock producing one-time products – long-term"
                                        if (
                                            account_name not in BSPL_cross_check_mark
                                            and account_name in BSPL_cross_check_cache
                                        ):
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )

                                        account_name = "seasonal crops or plants producing one-time products – long-term"
                                        if (
                                            account_name not in BSPL_cross_check_mark
                                            and account_name in BSPL_cross_check_cache
                                        ):
                                            if len(subtotal_rows) == 2:
                                                search_col_and_cross_ref(
                                                    "cost",
                                                    account_name,
                                                    subtotal_rows[0],
                                                )
                                            elif len(subtotal_rows) == 3:
                                                search_col_and_cross_ref(
                                                    "cost",
                                                    account_name,
                                                    subtotal_rows[1],
                                                )

                                        account_name = "immature livestock producing periodic products"
                                        if (
                                            account_name not in BSPL_cross_check_mark
                                            and account_name in BSPL_cross_check_cache
                                        ):
                                            if len(subtotal_rows) == 2:
                                                search_col_and_cross_ref(
                                                    "cost",
                                                    account_name,
                                                    subtotal_rows[1],
                                                )
                                            elif len(subtotal_rows) == 3:
                                                search_col_and_cross_ref(
                                                    "cost",
                                                    account_name,
                                                    subtotal_rows[2],
                                                )

                                        account_name = "allowance for impairment of biological assets – long-term"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "allowance", account_name, len(df) - 1
                                            )

                                    else:
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 2,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 1,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 1
                                            ]
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                subtotal_rows[0],
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                        account_name = "revenue deductions"
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[1],
                                                    len(df.columns) - 2,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[1], len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[1],
                                                    len(df.columns) - 1,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[1], len(df.columns) - 1
                                            ]
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                -CY_bal,
                                                -PY_bal,
                                                subtotal_rows[1],
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                        account_name = "10"
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    len(df) - 1, len(df.columns) - 2
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                len(df) - 1, len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    len(df) - 1, len(df.columns) - 1
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                len(df) - 1, len(df.columns) - 1
                                            ]
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                len(df) - 1,
                                                len(df.columns) - 2,
                                                0,
                                                -1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                elif (
                                    "accounts receivable from customers"
                                    in heading_lower
                                    and "by significant customer" in heading_lower
                                ):  # Case 9: Cast and Cross Grand Total của cột chứa kí tự đặc biệt
                                    if (
                                        "accounts receivable from customers – long-term"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = (
                                            "accounts receivable from customers"
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                len(df) - 1,
                                            )

                                    elif (
                                        "accounts receivable from customers"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = "accounts receivable from customers – long-term"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                len(df) - 1,
                                            )

                                    elif (
                                        "accounts receivable from customers"
                                        in BSPL_cross_check_cache
                                        and "accounts receivable from customers – long-term"
                                        in BSPL_cross_check_cache
                                    ):
                                        account_name = "accounts receivable from customers-combined"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                len(df) - 1,
                                            )

                                elif (
                                    "long-term borrowings" in heading_lower
                                    or "long-term bonds" in heading_lower
                                    or "long-term borrowing" in heading_lower
                                    or "long-term bond" in heading_lower
                                ):
                                    # search and cross ref costs
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "carrying", account_name, len(df) - 1
                                        )

                                elif "provisions" in heading_lower:
                                    if heading_lower == "long-term provisions":
                                        account_name = "provisions – long-term"
                                    else:
                                        account_name = "provisions – short-term"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_row_and_cross_ref(
                                            account_name, len(df.columns) - 1
                                        )

                        else:
                            if heading_lower not in CROSS_CHECK_TABLES_FORM_3:
                                # Case 3: Standard table là bảng có 3 cột, côt 2 là CY balance, cột 3 là PY balance, không có subtotal. Casting ở trên xuống dưới này là cross ref lên BS.
                                if heading_lower == "accrued expenses – short-term":
                                    account_name = "accrued expenses"
                                elif heading_lower == "accrued expenses – long-term":
                                    account_name = "long-term accrued expenses"
                                elif heading_lower == "deferred revenue – long-term":
                                    account_name = "long-term deferred revenue"
                                elif (
                                    "sales or disposals of investment" in heading_lower
                                ):
                                    account_name = "21"
                                elif "general and administration" in heading_lower:
                                    account_name = "26"
                                else:
                                    account_name = heading_lower

                                if (
                                    "accounts payable to suppliers" in heading_lower
                                    and "by significant" in heading_lower
                                ):
                                    if (
                                        "long-term accounts payable to suppliers"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = "accounts payable to suppliers"
                                    elif (
                                        "accounts payable to suppliers"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = (
                                            "long-term accounts payable to suppliers"
                                        )
                                    elif (
                                        "accounts payable to suppliers"
                                        in BSPL_cross_check_cache
                                        and "long-term accounts payable to suppliers"
                                        in BSPL_cross_check_cache
                                    ):
                                        account_name = (
                                            "accounts payable to suppliers-combined"
                                        )

                                    if account_name not in BSPL_cross_check_mark:
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 2,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 1,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 1
                                            ]
                                        )
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            subtotal_rows[0],
                                            len(df.columns) - 2,
                                            0,
                                            -1,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                                elif (
                                    "accounts payable to suppliers" in heading_lower
                                    and "by payment" in heading_lower
                                ):
                                    ST_row = 0
                                    LT_row = 0
                                    for j, row in df.iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if "short-term" in row_text:
                                            ST_row = j
                                        if "long-term" in row_text:
                                            LT_row = j
                                    # search and cross ref costs carrying amounts - ST & LT
                                    account_name = "accounts payable to suppliers"
                                    if account_name not in BSPL_cross_check_mark:
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    ST_row, len(df.columns) - 2
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                ST_row, len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    ST_row, len(df.columns) - 1
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                ST_row, len(df.columns) - 1
                                            ]
                                        )
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            ST_row,
                                            len(df.columns) - 2,
                                            0,
                                            -1,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                                    account_name = (
                                        "long-term accounts payable to suppliers"
                                    )
                                    if account_name not in BSPL_cross_check_mark:
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    LT_row, len(df.columns) - 2
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                LT_row, len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    LT_row, len(df.columns) - 1
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                LT_row, len(df.columns) - 1
                                            ]
                                        )
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            LT_row,
                                            len(df.columns) - 2,
                                            0,
                                            -1,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                                else:
                                    if account_name not in BSPL_cross_check_mark:
                                        CY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 2,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 2
                                            ]
                                        )
                                        PY_bal = (
                                            0
                                            if pd.isna(
                                                df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 1,
                                                ]
                                            )
                                            else df_numeric.iloc[
                                                subtotal_rows[0], len(df.columns) - 1
                                            ]
                                        )
                                        cross_check_with_BSPL(
                                            df,
                                            cross_ref_marks,
                                            issues,
                                            account_name,
                                            CY_bal,
                                            PY_bal,
                                            subtotal_rows[0],
                                            len(df.columns) - 2,
                                            0,
                                            -1,
                                        )
                                        BSPL_cross_check_mark.append(account_name)

                                    if account_name == "investment property":
                                        for j, row in df.iterrows():
                                            row_text = " ".join(
                                                str(x).lower() for x in row
                                            )
                                            if "earn rental" in row_text:
                                                BSPL_cross_check_cache[
                                                    "investment property held to earn rental"
                                                ] = (
                                                    df_numeric.iloc[
                                                        j, len(df.columns) - 2
                                                    ],
                                                    df_numeric.iloc[
                                                        j, len(df.columns) - 1
                                                    ],
                                                )
                                            if "capital appreciation" in row_text:
                                                BSPL_cross_check_cache[
                                                    "investment property held for capital appreciation"
                                                ] = (
                                                    df_numeric.iloc[
                                                        j, len(df.columns) - 2
                                                    ],
                                                    df_numeric.iloc[
                                                        j, len(df.columns) - 1
                                                    ],
                                                )

                            else:
                                # Case 7: Exceptional cases where we need to search for  key words and cross ref
                                if "acquisition of subsidiary" in heading_lower:
                                    NAss_acquired_row = GW_row = Con_trans_row = (
                                        Cash_acquired_row
                                    ) = Net_cash_outflow_row = 0
                                    NAss_acquired_val = GW_val = Con_trans_val = (
                                        Cash_acquired_val
                                    ) = Net_cash_outflow_val = 0

                                    for j, row in df.iloc[total_row_idx:].iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if "net" in row_text and "acquired" in row_text:
                                            NAss_acquired_row = j
                                            NAss_acquired_val = df_numeric.iloc[
                                                j, len(df.columns) - 1
                                            ]
                                        if "goodwill" in row_text:
                                            GW_row = j
                                            GW_val = df_numeric.iloc[
                                                j, len(df.columns) - 1
                                            ]
                                        if "consideration" in row_text:
                                            Con_trans_row = j
                                            Con_trans_val = df_numeric.iloc[
                                                j, len(df.columns) - 1
                                            ]
                                        if "cash acquired" in row_text:
                                            Cash_acquired_row = j
                                            Cash_acquired_val = df_numeric.iloc[
                                                j, len(df.columns) - 1
                                            ]
                                        if "cash outflow" in row_text:
                                            Net_cash_outflow_row = j
                                            Net_cash_outflow_val = df_numeric.iloc[
                                                j, len(df.columns) - 1
                                            ]

                                    diff = Con_trans_val - (NAss_acquired_val + GW_val)
                                    is_ok = abs(round(diff)) == 0
                                    comment = f"Consideration paid: Recalculate = {(NAss_acquired_val + GW_val):,.2f}, Note = {Con_trans_val:,.2f}, Diff = {diff:,.0f}"
                                    marks.append(
                                        {
                                            "row": Con_trans_row,
                                            "col": len(df.columns) - 1,
                                            "ok": is_ok,
                                            "comment": (None if is_ok else comment),
                                        }
                                    )
                                    if not is_ok:
                                        issues.append(comment)

                                    diff = Net_cash_outflow_val - (
                                        Con_trans_val + Cash_acquired_val
                                    )
                                    is_ok = abs(round(diff)) == 0
                                    comment = f"Net cash outflow: Recalculate = {(Con_trans_val + Cash_acquired_val):,.2f}, Note = {Net_cash_outflow_val:,.2f}, Diff = {diff:,.0f}"
                                    marks.append(
                                        {
                                            "row": Net_cash_outflow_row,
                                            "col": len(df.columns) - 1,
                                            "ok": is_ok,
                                            "comment": (None if is_ok else comment),
                                        }
                                    )
                                    if not is_ok:
                                        issues.append(comment)

                                if "business segments" in heading_lower:
                                    row = df.iloc[total_row_idx]
                                    row_text = " ".join(str(x).lower() for x in row)
                                    if "segment revenue" in row_text:
                                        OAs_results_row = NAT_row = 0
                                        OAs_results_val = NAT_val = [0.0] * len(
                                            df.columns
                                        )
                                        for j, row in df.iloc[
                                            total_row_idx + 1 :
                                        ].iterrows():
                                            row_text = " ".join(
                                                str(x).lower() for x in row
                                            )
                                            if "operating" in row_text:
                                                OAs_results_row = j
                                                OAs_results_val = df_numeric.iloc[j]
                                            if "after tax" in row_text:
                                                NAT_row = j
                                                NAT_val = df_numeric.iloc[j]

                                        OAs_results_sum = df_numeric.iloc[
                                            total_row_idx + 1 : OAs_results_row - 1
                                        ].sum(skipna=True)
                                        NAT_detail_sum = df_numeric.iloc[
                                            OAs_results_row : NAT_row - 1
                                        ].sum(skipna=True)

                                        for col in range(len(df.columns)):
                                            cell_value = df.iloc[OAs_results_row, col]
                                            if pd.isna(
                                                OAs_results_val[col]
                                            ) and "-" in str(cell_value):
                                                total_val = (
                                                    0
                                                    if pd.isna(OAs_results_val[col])
                                                    else OAs_results_val[col]
                                                )
                                            else:
                                                total_val = OAs_results_val[col]

                                            if not pd.isna(
                                                OAs_results_val[col]
                                            ) and not pd.isna(OAs_results_sum[col]):
                                                diff = (
                                                    OAs_results_sum[col]
                                                    - OAs_results_val[col]
                                                )
                                                is_ok = abs(round(diff)) == 0
                                                comment = f"OAs_results: Col - {col+1}: Recalculate={OAs_results_sum[col]:,.2f}, Note ={OAs_results_val[col]:,.2f}, Diff={diff:,.2f}"
                                                marks.append(
                                                    {
                                                        "row": OAs_results_row,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                                if not is_ok:
                                                    issues.append(comment)

                                        for col in range(len(df.columns)):
                                            cell_value = df.iloc[NAT_row, col]
                                            if pd.isna(NAT_val[col]) and "-" in str(
                                                cell_value
                                            ):
                                                total_val = (
                                                    0
                                                    if pd.isna(NAT_val[col])
                                                    else NAT_val[col]
                                                )
                                            else:
                                                total_val = NAT_val[col]

                                            if not pd.isna(
                                                NAT_val[col]
                                            ) and not pd.isna(NAT_detail_sum[col]):
                                                diff = (
                                                    NAT_detail_sum[col] - NAT_val[col]
                                                )
                                                is_ok = abs(round(diff)) == 0
                                                comment = f"NAT: Col - {col+1}: Recalculate={NAT_detail_sum[col]:,.2f}, Note={NAT_val[col]:,.2f}, Diff={diff:,.2f}"
                                                marks.append(
                                                    {
                                                        "row": NAT_row,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                                if not is_ok:
                                                    issues.append(comment)

                                if heading_lower == "trading securities":
                                    # search and cross ref costs
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, subtotal_rows[0]
                                        )
                                    # search and cross ref allowance
                                    account_name = "allowance for diminution in the value of trading securities"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, subtotal_rows[0]
                                        )

                                elif "held-to-maturity investments" in heading_lower:
                                    # kiểm tra thuyết minh là ngắn hay dài
                                    term = "long-term"
                                    for j, row in df.iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if "short-term" in row_text:
                                            term = "short-term"
                                            break
                                    if term == "short-term":
                                        # search and cross ref costs
                                        account_name = "123"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )
                                        # search and cross ref allowance
                                        account_name = "124"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "allowance",
                                                account_name,
                                                subtotal_rows[0],
                                            )
                                    else:
                                        # search and cross ref costs
                                        account_name = "265"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )
                                        # search and cross ref allowance
                                        account_name = "266"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "allowance",
                                                account_name,
                                                subtotal_rows[0],
                                            )

                                elif (
                                    heading_lower
                                    == "equity investments in other entity"
                                    or heading_lower
                                    == "equity investments in other entities"
                                ):
                                    # search and cross ref costs (separate) or carrying amounts (consol)
                                    account_name = "investments in other entities"
                                    # if account_name not in BSPL_cross_check_mark:
                                    search_col_and_cross_ref(
                                        "cost", account_name, subtotal_rows[0]
                                    )
                                    search_col_and_cross_ref(
                                        "carrying amounts",
                                        account_name,
                                        subtotal_rows[0],
                                    )
                                    # search and cross ref allowance
                                    account_name = "264"
                                    search_col_and_cross_ref(
                                        "allowance", account_name, subtotal_rows[0]
                                    )

                                elif (
                                    "accounts receivable from customers"
                                    in heading_lower
                                    and "by significant customer" in heading_lower
                                ):
                                    if (
                                        "accounts receivable from customers – long-term"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = (
                                            "accounts receivable from customers"
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                subtotal_rows[0],
                                            )

                                    elif (
                                        "accounts receivable from customers"
                                        not in BSPL_cross_check_cache
                                    ):
                                        account_name = "accounts receivable from customers – long-term"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                subtotal_rows[0],
                                            )

                                    elif (
                                        "accounts receivable from customers"
                                        in BSPL_cross_check_cache
                                        and "accounts receivable from customers – long-term"
                                        in BSPL_cross_check_cache
                                    ):
                                        account_name = "accounts receivable from customers-combined"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                subtotal_rows[0],
                                            )

                                elif (
                                    "accounts receivable from customers"
                                    in heading_lower
                                    and "by payment term" in heading_lower
                                ):
                                    ST_row = 0
                                    LT_row = 0
                                    for j, row in df.iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if "short-term" in row_text:
                                            ST_row = j
                                        if "long-term" in row_text:
                                            LT_row = j
                                    # search and cross ref costs carrying amounts - ST & LT
                                    account_name = "accounts receivable from customers"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "carrying amount", account_name, ST_row
                                        )
                                    account_name = (
                                        "accounts receivable from customers – long-term"
                                    )
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "carrying amount", account_name, LT_row
                                        )

                                elif (
                                    "other" in heading_lower
                                    and "receivables" in heading_lower
                                ):
                                    if "long-term" in heading_lower:
                                        account_name = "other long-term receivables"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                subtotal_rows[0],
                                            )
                                    else:
                                        if (
                                            "other short-term receivables"
                                            in BSPL_cross_check_mark
                                        ):
                                            account_name = (
                                                "other short-term receivables"
                                            )
                                        else:
                                            account_name = "other receivables"

                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "carrying amount",
                                                account_name,
                                                subtotal_rows[0],
                                            )

                                elif heading_lower == "bad and doubtful debts":
                                    account_name = "allowance for doubtful debts"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, subtotal_rows[0]
                                        )

                                elif (
                                    heading_lower
                                    == "shortage of assets awaiting resolution"
                                ):
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "quantity", account_name, subtotal_rows[0]
                                        )

                                elif heading_lower == "inventories":
                                    # search and cross ref costs
                                    account_name = "141"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, subtotal_rows[0]
                                        )
                                    # search and cross ref allowance
                                    account_name = "142"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, subtotal_rows[0]
                                        )
                                    else:
                                        search_row_and_cross_ref(account_name, 1)

                                elif heading_lower == "biological assets – short-term":
                                    account_name = "151"
                                    if account_name in BSPL_cross_check_cache:
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )
                                    else:
                                        account_name = "152"
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )

                                    account_name = "153"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, subtotal_rows[0]
                                        )

                                elif "biological assets – long-term" in heading_lower:
                                    account_name = "livestock producing one-time products – long-term"
                                    if account_name in BSPL_cross_check_cache:
                                        if account_name not in BSPL_cross_check_mark:
                                            search_col_and_cross_ref(
                                                "cost", account_name, subtotal_rows[0]
                                            )
                                    else:
                                        account_name = "seasonal crops or plants producing one-time products – long-term"
                                        if account_name in BSPL_cross_check_cache:
                                            if (
                                                account_name
                                                not in BSPL_cross_check_mark
                                            ):
                                                search_col_and_cross_ref(
                                                    "cost",
                                                    account_name,
                                                    subtotal_rows[0],
                                                )
                                        else:
                                            account_name = "mmature livestock producing periodic products"
                                            if (
                                                account_name
                                                not in BSPL_cross_check_mark
                                            ):
                                                search_col_and_cross_ref(
                                                    "cost",
                                                    account_name,
                                                    subtotal_rows[0],
                                                )

                                    account_name = "allowance for impairment of biological assets – long-term"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "allowance", account_name, subtotal_rows[0]
                                        )

                                elif heading_lower == "construction in progress":
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_row_and_cross_ref(account_name, 1)

                                elif heading_lower in [
                                    "investment property held for capital appreciation",
                                    "long-term deferred expenses",
                                ]:
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_row_and_cross_ref(
                                            account_name, len(df.columns) - 1
                                        )

                                elif heading_lower in ["long-term work in progress"]:
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, subtotal_rows[0]
                                        )

                                elif "accounts payable to suppliers" in heading_lower:
                                    # search and cross ref costs
                                    account_name = "accounts payable to suppliers"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "cost", account_name, subtotal_rows[0]
                                        )

                                elif "taxes" in heading_lower:
                                    # search and cross ref costs
                                    ST_row = 0
                                    LT_row = 0
                                    final_STrow = [0.0] * len(df.columns)
                                    final_LTrow = [0.0] * len(df.columns)
                                    for j, row in df.iloc[
                                        subtotal_rows[0] :
                                    ].iterrows():
                                        row_text = " ".join(str(x).lower() for x in row)
                                        if "within" in row_text:
                                            ST_row = j
                                            final_STrow = df_numeric.iloc[ST_row]
                                        if "after" in row_text:
                                            LT_row = j
                                            final_LTrow = df_numeric.iloc[LT_row]

                                    account_name = heading_lower
                                    if ST_row == 0 and LT_row == 0:
                                        if account_name not in BSPL_cross_check_mark:
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        subtotal_rows[0],
                                                        len(df.columns) - 1,
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    subtotal_rows[0],
                                                    len(df.columns) - 1,
                                                ]
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[subtotal_rows[0], 1]
                                                )
                                                else df_numeric.iloc[
                                                    subtotal_rows[0], 1
                                                ]
                                            )
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                subtotal_rows[0] + 1,
                                                len(df.columns) - 1,
                                                0,
                                                (len(df.columns) - 1) - 1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                    elif ST_row != 0 and LT_row != 0:
                                        for col in range(len(df.columns)):
                                            combined = grand_components[col]
                                            final_val = (
                                                final_STrow[col] + final_LTrow[col]
                                            )
                                            if not pd.isna(final_val):
                                                diff = combined - final_val
                                                is_ok = abs(round(diff)) == 0
                                                comment = f"Grand total row - Col {col+1}: Recalculate = {combined:,.2f}, Note = {final_val:,.2f}, Diff = {diff:,.0f}"
                                                marks.append(
                                                    {
                                                        "row": len(df) - 1,
                                                        "col": col,
                                                        "ok": is_ok,
                                                        "comment": (
                                                            None if is_ok else comment
                                                        ),
                                                    }
                                                )
                                                if not is_ok:
                                                    issues.append(comment)

                                        account_name = heading_lower
                                        if account_name not in BSPL_cross_check_mark:
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        ST_row, len(df.columns) - 1
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    ST_row, len(df.columns) - 1
                                                ]
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(df_numeric.iloc[ST_row, 1])
                                                else df_numeric.iloc[ST_row, 1]
                                            )
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                ST_row,
                                                len(df.columns) - 1,
                                                0,
                                                (len(df.columns) - 1) - 1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)
                                        account_name = (
                                            f"{heading_lower.strip()} – long-term"
                                        )
                                        if account_name not in BSPL_cross_check_mark:
                                            CY_bal = (
                                                0
                                                if pd.isna(
                                                    df_numeric.iloc[
                                                        LT_row, len(df.columns) - 1
                                                    ]
                                                )
                                                else df_numeric.iloc[
                                                    LT_row, len(df.columns) - 1
                                                ]
                                            )
                                            PY_bal = (
                                                0
                                                if pd.isna(df_numeric.iloc[LT_row, 1])
                                                else df_numeric.iloc[LT_row, 1]
                                            )
                                            cross_check_with_BSPL(
                                                df,
                                                cross_ref_marks,
                                                issues,
                                                account_name,
                                                CY_bal,
                                                PY_bal,
                                                LT_row,
                                                len(df.columns) - 1,
                                                0,
                                                (len(df.columns) - 1) - 1,
                                            )
                                            BSPL_cross_check_mark.append(account_name)

                                elif (
                                    "short-term borrowings" in heading_lower
                                    or "short-term bonds" in heading_lower
                                    or "short-term borrowing" in heading_lower
                                    or "short-term bond" in heading_lower
                                ):
                                    # search and cross ref costs
                                    account_name = heading_lower
                                    if account_name not in BSPL_cross_check_mark:
                                        search_col_and_cross_ref(
                                            "carrying", account_name, subtotal_rows[0]
                                        )

                                elif "provisions" in heading_lower:
                                    if heading_lower == "long-term provisions":
                                        account_name = "provisions – long-term"
                                    else:
                                        account_name = "provisions – short-term"
                                    if account_name not in BSPL_cross_check_mark:
                                        search_row_and_cross_ref(
                                            account_name, len(df.columns) - 1
                                        )

                # -----------------------------------------------------------------------
                # Cast cột tổng trong bảng, Chỉ kiểm tra cột tổng nếu heading nằm trong danh sách
                # -----------------------------------------------------------------------
                if is_table_with_column_total and last_col_idx > 1:

                    def is_borrowing_movement(df) -> bool:
                        borrowing_movement = False
                        for i, row in df.iterrows():
                            row_text = " ".join(str(cell).lower() for cell in row)
                            if "movement" in row_text:
                                borrowing_movement = True
                                break
                        return borrowing_movement

                    def compare_sum_with_total(start_col, end_col):
                        for i in range(total_row_idx + 1):  # Duyệt từng dòng chi tiết
                            row = df_numeric.iloc[i]
                            # row_sum = row.drop(labels=df.columns[last_col_idx]).sum(skipna=True)
                            row_sum = row[df.columns[start_col:end_col]].sum(
                                skipna=True
                            )
                            # col_total_val = row.iloc[end_col]
                            col_total_val = (
                                0 if pd.isna(row.iloc[end_col]) else row.iloc[end_col]
                            )
                            if not pd.isna(col_total_val) and not pd.isna(row_sum):
                                diff = row_sum - col_total_val
                                is_ok = abs(round(diff)) == 0
                                comment = f"CỘT TỔNG - Dòng {i+1}: Tính lại={row_sum:,.2f}, Trên bảng={col_total_val:,.2f}, Sai lệch={diff:,.2f}"
                                marks.append(
                                    {
                                        "row": i,
                                        "col": end_col,
                                        "ok": is_ok,
                                        "comment": (None if is_ok else comment),
                                    }
                                )
                                if not is_ok:
                                    issues.append(comment)

                    start_cidx = 0
                    end_cidx = last_col_idx
                    if "borrowings" not in heading_lower:
                        if is_table_need_check_separately:
                            if (
                                cost_start_row_idx != 0
                                and AD_start_row_idx != 0
                                and NBV_start_row_idx != 0
                            ):
                                compare_sum_with_total(start_cidx, end_cidx)
                        else:
                            if (
                                heading_lower
                                == "investment property held for capital appreciation"
                            ):
                                IP_movement = False
                                for j, row in df.iterrows():
                                    row_text = " ".join(str(x).lower() for x in row)
                                    if "opening balance" in row_text:
                                        IP_movement = True
                                        break
                                if IP_movement:
                                    compare_sum_with_total(start_cidx, end_cidx)
                            elif "finance lease liabilities" in heading_lower:
                                if len(df.columns) > 5:
                                    compare_sum_with_total(start_cidx, end_cidx - 3)
                                    compare_sum_with_total(end_cidx - 2, end_cidx)
                                else:
                                    compare_sum_with_total(start_cidx, end_cidx)
                            else:
                                compare_sum_with_total(start_cidx, end_cidx)
                    else:
                        if "borrowings" in heading_lower and is_borrowing_movement(df):
                            start_cidx = 1
                            end_cidx = last_col_idx
                            compare_sum_with_total(start_cidx, end_cidx)
                # -----------------------------------------------------------------------
                # Trả kết quả
                # -----------------------------------------------------------------------
                if not issues:
                    status = "✅ Kiểm tra công thức: KHỚP (0 sai lệch)"
                else:
                    preview = "; ".join(issues[:10])
                    more = f" ... (+{len(issues)-10} dòng)" if len(issues) > 10 else ""
                    status = f"❌ Kiểm tra công thức: {len(issues)} sai lệch. {preview}{more}"

                return {
                    "status": status,
                    "marks": marks,
                    "cross_ref_marks": cross_ref_marks,
                }

    except Exception:
        return {
            "status": "⚠️ Không xác định được dòng tổng hoặc lỗi khi kiểm tra",
            "marks": [],
            "cross_ref_marks": [],
        }


# =====================
# Excel writing routines
# =====================
def apply_status_colors(ws):
    """Tô màu trạng thái kiểm tra."""
    for row in ws.iter_rows(min_row=2, min_col=2, max_col=2):
        for cell in row:
            if "✅" in cell.value:
                cell.fill = PatternFill(start_color="C6EFCE", fill_type="solid")
            elif "❌" in cell.value:
                cell.fill = PatternFill(start_color="FFC7CE", fill_type="solid")
            elif "ℹ️" in cell.value:
                cell.fill = PatternFill(start_color="87CEEB", fill_type="solid")
            else:
                cell.fill = PatternFill(start_color="FFEB9C", fill_type="solid")


def write_summary_sheet(ws, results, sheet_positions, wb):
    """Ghi sheet tổng hợp trạng thái kiểm tra và tạo hyperlink đến từng bảng trong sheet 'FS casting'."""
    ws.title = "Tổng hợp kiểm tra"
    ws.append(["Tên bảng", "Trạng thái kiểm tra"])

    for i, (result, (heading, start_row)) in enumerate(zip(results, sheet_positions)):
        cell = ws.cell(row=i + 2, column=1, value=heading)
        cell.hyperlink = f"#'FS casting'!A{start_row}"
        cell.style = "Hyperlink"
        ws.cell(row=i + 2, column=2, value=result.get("status"))

    apply_status_colors(ws)


def write_table_sheet(wb, table_heading_pairs, results):
    """Ghi tất cả các bảng vào một sheet duy nhất 'FS casting', mỗi bảng cách nhau vài dòng."""
    sheet_name = "FS casting"
    ws = wb.create_sheet(title=sheet_name)
    sheet_positions = []  # lưu vị trí bắt đầu của từng bảng
    current_row = 1

    for i, ((table, heading), result) in enumerate(zip(table_heading_pairs, results)):
        table.columns = table.columns.map(str)
        raw_name = heading if heading else f"Bảng {i+1}"
        sheet_positions.append((raw_name, current_row))  # lưu heading và vị trí dòng

        # Ghi heading
        ws.cell(row=current_row, column=1, value=raw_name)
        current_row += 1

        # Ghi bảng
        start_row = current_row
        start_col = 1

        for row_idx, row in enumerate(
            dataframe_to_rows(table, index=False, header=True)
        ):
            for j, value in enumerate(row):
                # Kiểm tra nếu value là số (hoặc chuỗi có thể chuyển thành số)
                if row_idx == 0:
                    ws.cell(row=current_row, column=start_col + j, value=value)
                else:
                    if isinstance(value, (int, float)):
                        cell = ws.cell(
                            row=current_row, column=start_col + j, value=value
                        )
                        cell.number_format = '_(* #,##0_);_(* (#,##0);_(* "-"??_);_(@_)'
                    else:
                        # Thử chuyển đổi chuỗi thành số
                        try:
                            num_val = float(
                                str(value)
                                .replace(",", "")
                                .replace("(", "-")
                                .replace(")", "")
                            )
                            cell = ws.cell(
                                row=current_row, column=start_col + j, value=num_val
                            )
                            cell.number_format = (
                                '_(* #,##0_);_(* (#,##0);_(* "-"??_);_(@_)'
                            )
                        except ValueError:
                            ws.cell(row=current_row, column=start_col + j, value=value)
            current_row += 1

        end_row = current_row - 1
        end_col = start_col + len(table.columns) - 1

        # Tạo định dạng bảng Excel
        table_range = f"{ws.cell(row=start_row, column=start_col).coordinate}:{ws.cell(row=end_row, column=end_col).coordinate}"
        excel_table = Table(displayName=f"Table_{i+1}", ref=table_range)
        style = TableStyleInfo(
            name="TableStyleMedium9",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        excel_table.tableStyleInfo = style
        ws.add_table(excel_table)

        # Áp dụng highlight & comment
        apply_cell_marks(ws, result.get("marks", []), start_row - 1, start_col - 1)
        apply_crossref_marks(
            ws, result.get("cross_ref_marks", []), start_row, start_col - 1
        )

        # Ghi trạng thái kiểm tra
        ws.cell(row=current_row + 1, column=1, value="Trạng thái kiểm tra:")
        ws.cell(row=current_row + 1, column=2, value=result.get("status"))
        back_cell = ws.cell(row=current_row + 2, column=1, value="⬅ Quay lại Tổng hợp")
        back_cell.hyperlink = "#'Tổng hợp kiểm tra'!A1"
        back_cell.style = "Hyperlink"

        current_row += 4  # cách vài dòng trước khi ghi bảng tiếp theo

    return sheet_positions


# =====================
# Starting point here: Trích xuất bảng từ word để kiểm tra, và trả kết quả ra excel file
# =====================
def export_check_result(word_path, excel_path):
    """Quy trình chính: đọc Word, kiểm tra, ghi Excel."""

    table_heading_pairs = read_word_tables_with_headings(word_path)
    results = [
        check_table_total(table, heading) for table, heading in table_heading_pairs
    ]

    wb = Workbook()
    summary_ws = wb.active

    sheet_positions = write_table_sheet(wb, table_heading_pairs, results)
    write_summary_sheet(summary_ws, results, sheet_positions, wb)

    wb.save(excel_path)
    os.startfile(output_excel)


# Chạy chương trình
if __name__ == "__main__":

    """root = tk.Tk()
    root.withdraw()
    # Hiển thị hộp thoại chọn file
    word_file = filedialog.askopenfilename(
    title="Chọn file Word",
    filetypes=[("Word Documents", "*.docx *.DOCX")]
    )"""
    word_file = r"C:\\Users\\mnguyen3\\Downloads\\New folder (5)\\ABC Company-30Jun26VAS-EN.DOCX"
    output_excel = (
        r"C:\\Users\\mnguyen3\\Downloads\\New folder (5)\\kiem_tra_bang_chi_tiet.xlsx"
    )
    export_check_result(word_file, output_excel)
