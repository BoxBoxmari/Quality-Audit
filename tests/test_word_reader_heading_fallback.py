from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from quality_audit.io.word_reader import WordReader


def _append_minimal_sect_pr_to_paragraph(paragraph) -> None:
    """Insert w:sectPr as last child of w:p (section break inside paragraph)."""
    sect_pr = OxmlElement("w:sectPr")
    pg_sz = OxmlElement("w:pgSz")
    pg_sz.set(qn("w:w"), "11906")
    pg_sz.set(qn("w:h"), "16838")
    sect_pr.append(pg_sz)
    paragraph._element.append(sect_pr)


class TestWordReaderHeadingFallback:
    def test_heading_persists_across_adjacent_tables_without_new_paragraph(
        self, tmp_path
    ):
        """Heading should continue to next table when no explicit boundary exists."""
        doc_path = tmp_path / "adjacent_tables_heading_continuity.docx"
        doc = Document()

        heading = doc.add_paragraph("Income Statement")
        heading.style = doc.styles["Heading 1"]

        t1 = doc.add_table(rows=2, cols=2)
        t1.cell(0, 0).text = "Revenue"
        t1.cell(0, 1).text = "100"
        t1.cell(1, 0).text = "COGS"
        t1.cell(1, 1).text = "40"

        # No paragraph between tables: heading should remain in effect.
        t2 = doc.add_table(rows=2, cols=2)
        t2.cell(0, 0).text = "Operating expense"
        t2.cell(0, 1).text = "30"
        t2.cell(1, 0).text = "Profit"
        t2.cell(1, 1).text = "30"

        doc.save(str(doc_path))

        reader = WordReader()
        tables = reader.read_tables_with_headings(str(doc_path))

        assert len(tables) == 2
        _df1, h1, ctx1 = tables[0]
        _df2, h2, ctx2 = tables[1]

        assert h1 == "Income Statement"
        assert h2 == "Income Statement"
        assert ctx1.get("heading_source") == "paragraph"
        assert ctx2.get("heading_source") == "paragraph"

    def test_heading_from_table_first_row_when_no_paragraph_heading(
        self, tmp_path, monkeypatch
    ):
        """
        Table with no preceding paragraphs and descriptive first row
        should use first-row text as heading with heading_source='table_first_row'.
        """
        doc_path = tmp_path / "no_paragraph_heading.docx"
        doc = Document()

        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Descriptive heading row"
        table.cell(0, 1).text = "2024"
        table.cell(0, 2).text = "2023"
        table.cell(1, 0).text = "10"
        table.cell(1, 1).text = "100"
        table.cell(1, 2).text = "90"
        table.cell(2, 0).text = "20"
        table.cell(2, 1).text = "200"
        table.cell(2, 2).text = "180"

        doc.save(str(doc_path))

        monkeypatch.setattr(
            "quality_audit.io.word_reader.get_feature_flags",
            lambda: {"heading_fallback_from_table_first_row": True},
        )
        reader = WordReader()
        tables = reader.read_tables_with_headings(str(doc_path))

        assert len(tables) == 1
        df, heading, table_ctx = tables[0]

        assert heading == "Descriptive heading row"
        assert table_ctx.get("heading_source") == "table_first_row"

    def test_paragraph_heading_preferred_over_code_like_first_row(self, tmp_path):
        """
        When first rows contain only numeric/code-like cells, the reader should fall back
        to paragraph heading and not manufacture a heading from the first row.
        """
        doc_path = tmp_path / "paragraph_heading_preferred.docx"
        doc = Document()

        # Add a strong paragraph heading before the table
        para = doc.add_paragraph("Balance Sheet")
        para.style = doc.styles["Heading 1"]

        table = doc.add_table(rows=2, cols=3)
        table.cell(0, 0).text = "10"
        table.cell(0, 1).text = "100"
        table.cell(0, 2).text = "90"
        table.cell(1, 0).text = "20A"
        table.cell(1, 1).text = "200"
        table.cell(1, 2).text = "180"

        doc.save(str(doc_path))

        reader = WordReader()
        tables = reader.read_tables_with_headings(str(doc_path))

        assert len(tables) == 1
        df, heading, table_ctx = tables[0]

        # First-row cells are all numeric/code-like, so we must keep the paragraph heading
        assert heading == "Balance Sheet"
        assert table_ctx.get("heading_source") == "paragraph"

    def test_note_number_not_inherited_after_sect_pr_inside_paragraph(self, tmp_path):
        """
        After w:sectPr inside w:p, note context must reset so a later table
        does not inherit note_number from a prior note heading.
        """
        doc_path = tmp_path / "sect_pr_resets_note.docx"
        doc = Document()

        note_para = doc.add_paragraph("Note 10")
        note_para.style = doc.styles["Heading 1"]

        t1 = doc.add_table(rows=2, cols=2)
        t1.cell(0, 0).text = "A"
        t1.cell(0, 1).text = "1"
        t1.cell(1, 0).text = "B"
        t1.cell(1, 1).text = "2"

        doc.add_paragraph("")
        t2 = doc.add_table(rows=2, cols=2)
        t2.cell(0, 0).text = "C"
        t2.cell(0, 1).text = "3"
        t2.cell(1, 0).text = "D"
        t2.cell(1, 1).text = "4"

        doc.save(str(doc_path))

        doc_reload = Document(str(doc_path))
        # Body order: p (note), tbl, p (empty), tbl — sectPr on the empty paragraph
        _append_minimal_sect_pr_to_paragraph(doc_reload.paragraphs[1])
        doc_reload.save(str(doc_path))

        reader = WordReader()
        tables = reader.read_tables_with_headings(str(doc_path))

        assert len(tables) == 2
        _df1, _h1, ctx1 = tables[0]
        _df2, _h2, ctx2 = tables[1]

        assert ctx1.get("note_number") == "10"
        assert ctx2.get("note_number") is None
