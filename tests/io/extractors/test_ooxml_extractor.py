"""
Tests for OOXML table grid extractor: grid reconstruction, quality_score, invariant checks.
"""

import pytest
from docx import Document

from quality_audit.io.extractors.ooxml_table_grid_extractor import (
    Cell,
    ExtractionResult,
    OOXMLTableGridExtractor,
)


@pytest.fixture
def simple_docx_table():
    """Create a minimal docx Document with one table."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Account"
    table.cell(0, 1).text = "CY"
    table.cell(0, 2).text = "PY"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "1000"
    table.cell(1, 2).text = "900"
    table.cell(2, 0).text = "Total"
    table.cell(2, 1).text = "1000"
    table.cell(2, 2).text = "900"
    return table


def test_ooxml_extractor_returns_extraction_result(simple_docx_table):
    """OOXMLTableGridExtractor.extract returns ExtractionResult with grid and quality_score."""
    ext = OOXMLTableGridExtractor()
    result = ext.extract(simple_docx_table)
    assert isinstance(result, ExtractionResult)
    assert result.grid
    assert result.rows == 3
    assert result.cols == 3
    assert 0.0 <= result.quality_score <= 1.0


def test_extraction_result_is_usable_when_quality_above_threshold(simple_docx_table):
    """ExtractionResult.is_usable is True when quality_score >= 0.6 and no critical violations."""
    ext = OOXMLTableGridExtractor()
    result = ext.extract(simple_docx_table)
    assert result.quality_score >= OOXMLTableGridExtractor.QUALITY_THRESHOLD
    assert result.is_usable


def test_extraction_result_cell_list_populated(simple_docx_table):
    """Extraction yields list of Cell objects."""
    ext = OOXMLTableGridExtractor()
    result = ext.extract(simple_docx_table)
    assert isinstance(result.cells, list)
    for c in result.cells:
        assert isinstance(c, Cell)
        assert hasattr(c, "row") and hasattr(c, "col") and hasattr(c, "value")


def test_ooxml_extractor_low_quality_result_not_usable():
    """ExtractionResult with quality_score < 0.6 is not is_usable."""
    result = ExtractionResult(
        grid=[["x"]],
        quality_score=0.5,
        quality_flags=[],
        rows=1,
        cols=1,
    )
    assert not result.is_usable


def test_extraction_result_quality_flags():
    """ExtractionResult may have quality_flags (e.g. empty_grid, GRID_CORRUPTION)."""
    result = ExtractionResult(
        grid=[],
        quality_score=0.0,
        quality_flags=["empty_grid"],
    )
    assert not result.is_usable
    assert "empty_grid" in result.quality_flags


def test_ooxml_extractor_gridspan_multicolumn():
    """R2: Table with horizontal merge (gridSpan) yields grid_span_count >= 1 and placeholders empty."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(0, 2).text = "C"
    table.cell(1, 0).text = "D"
    table.cell(1, 1).text = "E"
    table.cell(1, 2).text = "F"
    table.cell(0, 0).merge(table.cell(0, 1))
    ext = OOXMLTableGridExtractor()
    result = ext.extract(table)
    assert result.grid_span_count >= 1
    assert result.rows == 2 and result.cols >= 2
    # First row has content from merged cell; at least one cell is empty (merge placeholder)
    row0 = result.grid[0] if result.grid else []
    has_content = any((c or "").strip() for c in row0)
    has_empty = any((c or "").strip() == "" for c in row0)
    assert has_content and has_empty


def test_ooxml_extractor_vmerge_restart_continue():
    """R2: Table with vertical merge (vMerge) yields vmerge_count >= 1."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "X"
    table.cell(0, 1).text = "Y"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = "Z"
    table.cell(0, 0).merge(table.cell(1, 0))
    ext = OOXMLTableGridExtractor()
    result = ext.extract(table)
    assert result.vmerge_count >= 1
    assert result.grid and result.rows >= 2


def test_ooxml_extractor_gridspan_zero_treated_as_one():
    """R2: _get_grid_span returns 1 when gridSpan=0 or missing."""
    ext = OOXMLTableGridExtractor()
    from lxml import etree

    cell_with_zero = etree.fromstring(
        "<w:tc xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
        "<w:tcPr><w:gridSpan w:val='0'/></w:tcPr><w:p/></w:tc>"
    )
    assert ext._get_grid_span(cell_with_zero) == 1
    cell_no_span = etree.fromstring(
        "<w:tc xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
        "<w:p/></w:tc>"
    )
    assert ext._get_grid_span(cell_no_span) == 1


def test_ooxml_extractor_tblgrid_missing_fallback():
    """R2: When tblGrid is missing, grid is inferred and soft_anomaly tblGrid_missing recorded."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "C"
    table.cell(1, 1).text = "D"
    tbl = table._element
    for grid in tbl.xpath(".//*[local-name()='tblGrid']"):
        grid.getparent().remove(grid)
    ext = OOXMLTableGridExtractor()
    result = ext.extract(table)
    soft = [v for v in result.invariant_violations if v.startswith("soft_anomaly:")]
    assert any("tblGrid_missing" in v for v in soft)
    assert result.grid and result.rows == 2 and result.cols == 2


def test_extraction_result_grid_span_and_vmerge_fields():
    """ExtractionResult accepts grid_span_count and vmerge_count (snake_case)."""
    result = ExtractionResult(
        grid=[["a"]],
        quality_score=0.8,
        rows=1,
        cols=1,
        grid_span_count=2,
        vmerge_count=1,
    )
    assert result.grid_span_count == 2
    assert result.vmerge_count == 1
